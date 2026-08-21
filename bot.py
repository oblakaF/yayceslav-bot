import asyncio
import csv
import json
import logging
import mimetypes
import os
import random
import re
import sqlite3
import sys
import time
import uuid
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import Any

import edge_tts
from collections import defaultdict, deque
from gtts import gTTS
from ddgs import DDGS
from telegram.error import BadRequest
from docx import Document as DocxDocument
from dotenv import load_dotenv
from google import genai
from google.genai import types
from openpyxl import load_workbook
from telegram import (
    InlineKeyboardButton,
    InlineKeyboardMarkup,
    ReactionTypeEmoji,
    Update,
)
from telegram.constants import ChatAction, ChatType, UpdateType
from telegram.ext import (
    Application,
    CallbackQueryHandler,
    CommandHandler,
    ContextTypes,
    MessageHandler,
    MessageReactionHandler,
    filters,
)
import aggression_engine
import adaptation_cache
import chat_native_engine
import feedback_engine
import humanizer_engine
import hostile_streak_engine
import daily_title_engine
import humor_engine
import intent
import passive_engine
import reaction_engine
import social_engine
import state_engine
import style_engine
import thinking_engine
import title_pools
import voice_runtime
from personality import (
    DEFAULT_USER_SETTINGS,
    HOSTILE_RE,
    build_v2_base_instruction,
    detect_conversation_mode,
    is_serious_text,
    VOICE_STYLE_INSTRUCTION,
)
from vocabulary import (
    AWARD_LABELS,
    AWARD_TEMPLATES,
    BASE_REPLIES,
    CRINGE_REPLIES,
    GOY_REPLIES,
    HARD_REACTION_EMOJIS,
    MOODS,
    NISHIY_REPLIES,
    PROPHECIES,
    ROASTS,
    SIX_SEVEN_REPLIES,
    SKUF_REPLIES,
    WISDOMS,
    YAYCESLAV_REPLIES,
)


# ============================================================
# НАСТРОЙКИ
# ============================================================

load_dotenv()

TELEGRAM_TOKEN = os.getenv("TELEGRAM_BOT_TOKEN")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")

BOT_OWNER_ID_RAW = os.getenv(
    "BOT_OWNER_ID",
    "",
).strip()

BOT_OWNER_ID = (
    int(BOT_OWNER_ID_RAW)
    if BOT_OWNER_ID_RAW.isdigit()
    else 0
)

MODEL_NAME = "gemini-3.6-flash"
# ============================================================
# ЗАЩИТА ОТ СПАМА И ЛИМИТЫ
# ============================================================

# Максимальная длина обычного текстового запроса
MAX_USER_TEXT_CHARS = 3000

# Лимиты формата:
# "тип запроса": (количество запросов, период в секундах)
RATE_LIMITS = {
    "general": (5, 60.0),
    "media": (3, 60.0),
    "search": (2, 60.0),
}

# Не более трёх одновременных запросов к Gemini
GEMINI_SEMAPHORE = asyncio.Semaphore(3)

# Не более двух одновременных интернет-поисков
SEARCH_SEMAPHORE = asyncio.Semaphore(2)

# Здесь временно хранятся моменты запросов пользователей
REQUEST_TIMES: dict[
    tuple[int, str],
    deque[float],
] = defaultdict(deque)

# Не присылаем предупреждение о лимите каждую секунду
LAST_LIMIT_WARNING: dict[
    tuple[int, str],
    float,
] = {}

LIMIT_WARNING_COOLDOWN = 10.0

# Настройки голоса Яйцеслава
TTS_VOICE = "ru-RU-DmitryNeural"
TTS_RATE = "-5%"
TTS_PITCH = "-25Hz"
TTS_VOLUME = "+3%"

# Максимальная длина текста для одного голосового ответа
MAX_TTS_CHARS = 1500

# Папка для временных файлов
TEMP_DIR = Path("temp")
TEMP_DIR.mkdir(exist_ok=True)
# ============================================================
# ПОСТОЯННАЯ СТАТИСТИКА
# ============================================================

# На Railway база хранится на подключённом Volume.
# Локально используется обычная папка data.
RAILWAY_DATA_DIR = Path("/app/data")

if RAILWAY_DATA_DIR.exists():
    DATA_DIR = RAILWAY_DATA_DIR
else:
    DATA_DIR = Path("data")

DATA_DIR.mkdir(
    parents=True,
    exist_ok=True,
)

STATS_DB_PATH = DATA_DIR / "yayceslav_stats.db"


class ClosingSQLiteConnection(sqlite3.Connection):
    """sqlite3 context manager с обязательным close() после commit/rollback."""

    def __exit__(self, exc_type, exc, tb):
        try:
            return super().__exit__(exc_type, exc, tb)
        finally:
            self.close()


def get_db_connection(
    db_path: Path | None = None,
    timeout: float = 30,
) -> sqlite3.Connection:
    """
    Открывает соединение с нужными PRAGMA.

    db_path читается из STATS_DB_PATH заново при каждом вызове
    (а не как значение параметра по умолчанию) — иначе тесты,
    подменяющие bot.STATS_DB_PATH через monkeypatch, тихо продолжали
    бы писать в реальный файл базы, захваченный на момент импорта.

    journal_mode=WAL сохраняется в самом файле базы, но дешевле
    подтвердить его на каждом соединении, чем гадать, установлен ли он.
    foreign_keys обязательно включать на каждом соединении отдельно —
    SQLite не помнит это между подключениями.
    """

    if db_path is None:
        db_path = STATS_DB_PATH

    connection = sqlite3.connect(
        db_path,
        timeout=timeout,
        factory=ClosingSQLiteConnection,
    )
    connection.execute("PRAGMA journal_mode=WAL")
    connection.execute("PRAGMA foreign_keys=ON")
    return connection


def _ensure_column(
    connection: sqlite3.Connection,
    table: str,
    column: str,
    column_ddl: str,
) -> None:
    """
    Добавляет колонку, если её ещё нет — без разрушительной миграции.

    CREATE TABLE IF NOT EXISTS не добавляет новые колонки к уже
    существующей таблице, поэтому колонки, появившиеся после первого
    релиза схемы, нужно домигрировать так.
    """

    existing_columns = {
        row[1]
        for row in connection.execute(
            f"PRAGMA table_info({table})"
        )
    }

    if column not in existing_columns:
        connection.execute(
            f"ALTER TABLE {table} ADD COLUMN {column_ddl}"
        )


def initialize_stats_database() -> None:
    """Создаёт постоянную базу статистики."""

    with get_db_connection() as connection:
        connection.execute(
            """
            CREATE TABLE IF NOT EXISTS stats (
                name TEXT PRIMARY KEY,
                value INTEGER NOT NULL DEFAULT 0
            )
            """
        )

        connection.execute(
            """
            CREATE TABLE IF NOT EXISTS users (
                user_id INTEGER PRIMARY KEY
            )
            """
        )

        connection.execute(
            """
            CREATE TABLE IF NOT EXISTS chats (
                chat_id INTEGER PRIMARY KEY,
                chat_type TEXT NOT NULL
            )
            """
        )

        connection.execute(
            """
            CREATE TABLE IF NOT EXISTS user_settings (
                user_id INTEGER PRIMARY KEY,
                character TEXT NOT NULL DEFAULT 'classic',
                response_style TEXT NOT NULL DEFAULT 'bold',
                response_length TEXT NOT NULL DEFAULT 'normal',
                voice_enabled INTEGER NOT NULL DEFAULT 0,
                search_mode TEXT NOT NULL DEFAULT 'button',
                roughness TEXT NOT NULL DEFAULT 'medium'
            )
            """
        )

        # Хранит хард-мод и связанные настройки чата так,
        # чтобы они переживали рестарт Railway (раньше жили
        # только в context.chat_data и обнулялись при рестарте).
        connection.execute(
            """
            CREATE TABLE IF NOT EXISTS chat_settings (
                chat_id INTEGER PRIMARY KEY REFERENCES chats(chat_id),
                hard_mode_enabled INTEGER NOT NULL DEFAULT 1,
                hard_level TEXT NOT NULL DEFAULT 'normal',
                reaction_chance REAL NOT NULL DEFAULT 0.70,
                random_reply_chance REAL NOT NULL DEFAULT 0.16,
                updated_at TEXT NOT NULL DEFAULT (datetime('now'))
            )
            """
        )

        # Счётчики /hard_stats добавлены отдельной миграцией —
        # на момент первого релиза chat_settings их не было.
        _ensure_column(
            connection,
            "chat_settings",
            "reactions_count",
            "reactions_count INTEGER NOT NULL DEFAULT 0",
        )
        _ensure_column(
            connection,
            "chat_settings",
            "random_replies_count",
            "random_replies_count INTEGER NOT NULL DEFAULT 0",
        )
        _ensure_column(
            connection,
            "chat_settings",
            "trigger_replies_count",
            "trigger_replies_count INTEGER NOT NULL DEFAULT 0",
        )
        _ensure_column(
            connection,
            "chat_settings",
            "last_intervention_at",
            "last_intervention_at TEXT",
        )

        # Личное обращение — как пользователь попросил к нему
        # обращаться (/nickname). Добавлено отдельной миграцией.
        _ensure_column(
            connection,
            "user_settings",
            "custom_nickname",
            "custom_nickname TEXT",
        )

        # Лёгкая память об участниках группы: только безопасные
        # числа и текст, который сам пользователь попросил запомнить,
        # плюс шуточный архетип, который задаёт админ вручную —
        # никаких медицинских/финансовых/политических данных.
        connection.execute(
            """
            CREATE TABLE IF NOT EXISTS chat_member_profiles (
                chat_id INTEGER NOT NULL REFERENCES chats(chat_id),
                user_id INTEGER NOT NULL REFERENCES users(user_id),
                current_display_name TEXT,
                username TEXT,
                first_seen_at TEXT NOT NULL DEFAULT (datetime('now')),
                last_seen_at TEXT NOT NULL DEFAULT (datetime('now')),
                total_messages INTEGER NOT NULL DEFAULT 0,
                total_voice_messages INTEGER NOT NULL DEFAULT 0,
                total_photos INTEGER NOT NULL DEFAULT 0,
                total_stickers INTEGER NOT NULL DEFAULT 0,
                replies_to_bot INTEGER NOT NULL DEFAULT 0,
                insults_to_bot INTEGER NOT NULL DEFAULT 0,
                jokes_detected INTEGER NOT NULL DEFAULT 0,
                serious_messages INTEGER NOT NULL DEFAULT 0,
                relationship_level INTEGER NOT NULL DEFAULT 0,
                joke_archetype TEXT,
                self_reported_facts TEXT,
                last_humor_strategy TEXT,
                updated_at TEXT NOT NULL DEFAULT (datetime('now')),
                PRIMARY KEY (chat_id, user_id)
            )
            """
        )

        # Текущий шуточный титул (/title) — отдельно от joke_archetype:
        # архетип задаёт вручную админ, титул выпадает случайно самому
        # участнику или тому, кому он ответил.
        _ensure_column(
            connection,
            "chat_member_profiles",
            "current_title",
            "current_title TEXT",
        )

        # Недельная аналитика: агрегаты по дням, не полный текст
        # сообщений — статистика, а не архив переписки.
        connection.execute(
            """
            CREATE TABLE IF NOT EXISTS chat_activity_daily (
                chat_id INTEGER NOT NULL,
                user_id INTEGER NOT NULL,
                date TEXT NOT NULL,
                messages INTEGER NOT NULL DEFAULT 0,
                text_characters INTEGER NOT NULL DEFAULT 0,
                voice_messages INTEGER NOT NULL DEFAULT 0,
                voice_duration_seconds INTEGER NOT NULL DEFAULT 0,
                photos INTEGER NOT NULL DEFAULT 0,
                videos INTEGER NOT NULL DEFAULT 0,
                stickers INTEGER NOT NULL DEFAULT 0,
                documents INTEGER NOT NULL DEFAULT 0,
                replies INTEGER NOT NULL DEFAULT 0,
                replies_to_bot INTEGER NOT NULL DEFAULT 0,
                commands INTEGER NOT NULL DEFAULT 0,
                night_messages INTEGER NOT NULL DEFAULT 0,
                questions INTEGER NOT NULL DEFAULT 0,
                links INTEGER NOT NULL DEFAULT 0,
                edited_messages INTEGER NOT NULL DEFAULT 0,
                PRIMARY KEY (chat_id, user_id, date)
            )
            """
        )

        # Победитель каждой награды недели фиксируется один раз и
        # больше не меняется — иначе /week, /awards и автоотчёт могли
        # разойтись, а "Куколд-наблюдатель" выбирался бы заново
        # (случайно) при каждом обращении. PRIMARY KEY физически не
        # даёт появиться второму победителю той же награды за неделю.
        connection.execute(
            """
            CREATE TABLE IF NOT EXISTS weekly_award_winners (
                chat_id INTEGER NOT NULL,
                week_start TEXT NOT NULL,
                award_key TEXT NOT NULL,
                user_id INTEGER NOT NULL,
                selected_at TEXT NOT NULL DEFAULT (datetime('now')),
                PRIMARY KEY (chat_id, week_start, award_key)
            )
            """
        )

        # V2: один автоматический титул дня на чат и дату.
        # PRIMARY KEY не позволяет рестарту/гонке выдать второй
        # daily title в тот же календарный день.
        connection.execute(
            """
            CREATE TABLE IF NOT EXISTS daily_title_assignments (
                chat_id INTEGER NOT NULL,
                date TEXT NOT NULL,
                user_id INTEGER NOT NULL,
                title TEXT NOT NULL,
                assigned_at TEXT NOT NULL DEFAULT (datetime('now')),
                PRIMARY KEY (chat_id, date)
            )
            """
        )

        # Отдельно отмечаем, что daily title уже реально объявлен в Telegram.
        # Если отправка упадёт после атомарного выбора победителя, scheduler
        # повторит только объявление, а не выберет второго человека.
        _ensure_column(
            connection,
            "daily_title_assignments",
            "announced_at",
            "announced_at TEXT",
        )

        # Расписание автоматического недельного отчёта.
        # По умолчанию — воскресенье, 21:00 по МСК.
        _ensure_column(
            connection,
            "chat_settings",
            "weekly_report_enabled",
            "weekly_report_enabled INTEGER NOT NULL DEFAULT 0",
        )
        _ensure_column(
            connection,
            "chat_settings",
            "weekly_report_weekday",
            "weekly_report_weekday INTEGER NOT NULL DEFAULT 6",
        )
        _ensure_column(
            connection,
            "chat_settings",
            "weekly_report_time",
            "weekly_report_time TEXT NOT NULL DEFAULT '21:00'",
        )
        _ensure_column(
            connection,
            "chat_settings",
            "weekly_report_last_sent_date",
            "weekly_report_last_sent_date TEXT",
        )

        # 13-й динамический voice pack: храним только агрегированные
        # слова/короткие фразы, а не архив исходных сообщений.
        connection.execute(
            """
            CREATE TABLE IF NOT EXISTS chat_native_terms (
                chat_id INTEGER NOT NULL,
                term TEXT NOT NULL,
                occurrences INTEGER NOT NULL DEFAULT 0,
                first_seen TEXT NOT NULL DEFAULT (datetime('now')),
                last_seen TEXT NOT NULL DEFAULT (datetime('now')),
                PRIMARY KEY (chat_id, term)
            )
            """
        )
        connection.execute(
            """
            CREATE TABLE IF NOT EXISTS chat_native_term_users (
                chat_id INTEGER NOT NULL,
                term TEXT NOT NULL,
                user_id INTEGER NOT NULL,
                first_seen TEXT NOT NULL DEFAULT (datetime('now')),
                last_seen TEXT NOT NULL DEFAULT (datetime('now')),
                PRIMARY KEY (chat_id, term, user_id)
            )
            """
        )
        connection.execute(
            """
            CREATE TABLE IF NOT EXISTS chat_native_profiles (
                chat_id INTEGER PRIMARY KEY,
                terms_json TEXT NOT NULL DEFAULT '[]',
                distinct_users INTEGER NOT NULL DEFAULT 0,
                compiled_at TEXT
            )
            """
        )

        # Метаданные только собственных ответов бота. Полный текст ответа
        # здесь не хранится: нужен message_id + тип поведения для реакции.
        connection.execute(
            """
            CREATE TABLE IF NOT EXISTS bot_response_feedback (
                chat_id INTEGER NOT NULL,
                message_id INTEGER NOT NULL,
                voice_pack TEXT NOT NULL,
                humor_type TEXT,
                verdict_used INTEGER NOT NULL DEFAULT 0,
                reaction_score REAL NOT NULL DEFAULT 0,
                reaction_count INTEGER NOT NULL DEFAULT 0,
                created_at TEXT NOT NULL DEFAULT (datetime('now')),
                PRIMARY KEY (chat_id, message_id)
            )
            """
        )

        connection.commit()


initialize_stats_database()


def record_chat_native_message_sync(
    chat_id: int,
    user_id: int,
    text: str,
    chat_type: str = "group",
) -> int:
    """Сохраняет только агрегированные кандидаты локального сленга."""

    terms = chat_native_engine.extract_candidate_terms(text)
    if not terms:
        return 0

    with get_db_connection() as connection:
        connection.execute(
            "INSERT OR IGNORE INTO chats (chat_id, chat_type) VALUES (?, ?)",
            (chat_id, chat_type),
        )
        connection.execute(
            "INSERT OR IGNORE INTO users (user_id) VALUES (?)",
            (user_id,),
        )
        for term in terms:
            connection.execute(
                """
                INSERT INTO chat_native_terms (chat_id, term, occurrences)
                VALUES (?, ?, 1)
                ON CONFLICT(chat_id, term) DO UPDATE SET
                    occurrences = occurrences + 1,
                    last_seen = datetime('now')
                """,
                (chat_id, term),
            )
            connection.execute(
                """
                INSERT INTO chat_native_term_users (chat_id, term, user_id)
                VALUES (?, ?, ?)
                ON CONFLICT(chat_id, term, user_id) DO UPDATE SET
                    last_seen = datetime('now')
                """,
                (chat_id, term, user_id),
            )
        connection.commit()
    return len(terms)


def get_chat_native_profile_sync(chat_id: int) -> dict[str, Any]:
    with get_db_connection() as connection:
        row = connection.execute(
            """
            SELECT terms_json, distinct_users, compiled_at
            FROM chat_native_profiles
            WHERE chat_id = ?
            """,
            (chat_id,),
        ).fetchone()
    if row is None:
        return {"terms": [], "distinct_users": 0, "compiled_at": None}
    try:
        terms = json.loads(row[0] or "[]")
    except (TypeError, json.JSONDecodeError):
        terms = []
    return {
        "terms": [str(term) for term in terms if str(term).strip()],
        "distinct_users": int(row[1] or 0),
        "compiled_at": row[2],
    }


def refresh_due_chat_native_profiles_sync() -> int:
    """Первый pack собирает после достаточной выборки, затем обновляет раз в неделю."""

    now = datetime.now(timezone.utc)
    refreshed = 0
    refreshed_chat_ids: set[int] = set()
    with get_db_connection() as connection:
        chat_rows = connection.execute(
            "SELECT DISTINCT chat_id FROM chat_native_terms ORDER BY chat_id"
        ).fetchall()

        for (chat_id_raw,) in chat_rows:
            chat_id = int(chat_id_raw)
            profile_row = connection.execute(
                "SELECT compiled_at FROM chat_native_profiles WHERE chat_id = ?",
                (chat_id,),
            ).fetchone()
            if profile_row and profile_row[0]:
                try:
                    compiled_at = datetime.fromisoformat(str(profile_row[0]))
                    if compiled_at.tzinfo is None:
                        compiled_at = compiled_at.replace(tzinfo=timezone.utc)
                    if (now - compiled_at).total_seconds() < chat_native_engine.PROFILE_REFRESH_SECONDS:
                        continue
                except ValueError:
                    pass

            stats_rows = connection.execute(
                """
                SELECT terms.term, terms.occurrences, COUNT(users.user_id)
                FROM chat_native_terms AS terms
                LEFT JOIN chat_native_term_users AS users
                  ON users.chat_id = terms.chat_id AND users.term = terms.term
                WHERE terms.chat_id = ?
                GROUP BY terms.term, terms.occurrences
                """,
                (chat_id,),
            ).fetchall()
            distinct_users = int(
                connection.execute(
                    "SELECT COUNT(DISTINCT user_id) FROM chat_native_term_users WHERE chat_id = ?",
                    (chat_id,),
                ).fetchone()[0]
            )
            terms = chat_native_engine.compile_profile_terms(stats_rows)
            if not chat_native_engine.profile_is_ready(terms, distinct_users):
                continue

            connection.execute(
                """
                INSERT INTO chat_native_profiles (chat_id, terms_json, distinct_users, compiled_at)
                VALUES (?, ?, ?, ?)
                ON CONFLICT(chat_id) DO UPDATE SET
                    terms_json = excluded.terms_json,
                    distinct_users = excluded.distinct_users,
                    compiled_at = excluded.compiled_at
                """,
                (
                    chat_id,
                    json.dumps(list(terms), ensure_ascii=False),
                    distinct_users,
                    now.isoformat(),
                ),
            )
            refreshed += 1
            refreshed_chat_ids.add(chat_id)

        # Не даём словарю расти бесконечно: редкие кандидаты, которые
        # не появлялись 60 дней, забываются. Устойчивые локальные мемы
        # (5+ употреблений) сохраняются и могут вернуться в следующий pack.
        connection.execute(
            """
            DELETE FROM chat_native_terms
            WHERE last_seen < datetime('now', '-60 days')
              AND occurrences < 5
            """
        )
        connection.execute(
            """
            DELETE FROM chat_native_term_users
            WHERE NOT EXISTS (
                SELECT 1 FROM chat_native_terms t
                WHERE t.chat_id = chat_native_term_users.chat_id
                  AND t.term = chat_native_term_users.term
            )
            """
        )
        connection.commit()

    for refreshed_chat_id in refreshed_chat_ids:
        adaptation_cache.invalidate("native", refreshed_chat_id)

    return refreshed


async def refresh_due_chat_native_profiles() -> int:
    return await asyncio.to_thread(refresh_due_chat_native_profiles_sync)


def store_bot_response_feedback_sync(
    chat_id: int,
    message_id: int,
    trace: feedback_engine.ResponseTrace,
) -> None:
    with get_db_connection() as connection:
        connection.execute(
            """
            INSERT OR IGNORE INTO bot_response_feedback
                (chat_id, message_id, voice_pack, humor_type, verdict_used)
            VALUES (?, ?, ?, ?, ?)
            """,
            (
                chat_id,
                message_id,
                trace.voice_pack,
                trace.humor_type,
                int(trace.verdict_used),
            ),
        )
        # Последних 500 ответов на чат более чем достаточно для окна 20.
        connection.execute(
            """
            DELETE FROM bot_response_feedback
            WHERE chat_id = ?
              AND rowid NOT IN (
                  SELECT rowid FROM bot_response_feedback
                  WHERE chat_id = ?
                  ORDER BY created_at DESC, message_id DESC
                  LIMIT 500
              )
            """,
            (chat_id, chat_id),
        )
        connection.commit()


def apply_bot_reaction_delta_sync(
    chat_id: int,
    message_id: int,
    score_delta: float,
    count_delta: int,
) -> bool:
    with get_db_connection() as connection:
        cursor = connection.execute(
            """
            UPDATE bot_response_feedback
            SET reaction_score = reaction_score + ?,
                reaction_count = MAX(0, reaction_count + ?)
            WHERE chat_id = ? AND message_id = ?
            """,
            (score_delta, count_delta, chat_id, message_id),
        )
        connection.commit()
        return cursor.rowcount > 0


def get_chat_feedback_adaptation_sync(chat_id: int) -> dict[str, Any]:
    with get_db_connection() as connection:
        rows = connection.execute(
            """
            SELECT voice_pack, humor_type, verdict_used,
                   reaction_score, reaction_count
            FROM bot_response_feedback
            WHERE chat_id = ? AND reaction_count > 0
            ORDER BY created_at DESC, message_id DESC
            LIMIT 200
            """,
            (chat_id,),
        ).fetchall()
    return feedback_engine.build_adaptation(
        [
            {
                "voice_pack": row[0],
                "humor_type": row[1],
                "verdict_used": bool(row[2]),
                "reaction_score": float(row[3]),
                "reaction_count": int(row[4]),
            }
            for row in rows
        ]
    )


def get_chat_native_learning_status_sync(chat_id: int) -> dict[str, Any]:
    profile = get_chat_native_profile_sync(chat_id)
    with get_db_connection() as connection:
        candidate_terms = int(
            connection.execute(
                "SELECT COUNT(*) FROM chat_native_terms WHERE chat_id = ?",
                (chat_id,),
            ).fetchone()[0]
        )
        distinct_users = int(
            connection.execute(
                "SELECT COUNT(DISTINCT user_id) FROM chat_native_term_users WHERE chat_id = ?",
                (chat_id,),
            ).fetchone()[0]
        )
        tracked_messages = int(
            connection.execute(
                "SELECT COUNT(*) FROM bot_response_feedback WHERE chat_id = ?",
                (chat_id,),
            ).fetchone()[0]
        )
    adaptation = get_chat_feedback_adaptation_sync(chat_id)
    return {
        **profile,
        "candidate_terms": candidate_terms,
        "observed_users": distinct_users,
        "tracked_messages": tracked_messages,
        "reacted_messages": int(adaptation.get("reacted_messages", 0)),
    }


def get_user_settings_sync(
    user_id: int,
) -> dict[str, Any]:
    """Получает сохранённые настройки пользователя."""

    with get_db_connection() as connection:
        connection.execute(
            """
            INSERT OR IGNORE INTO user_settings (
                user_id, roughness
            )
            VALUES (?, 'high')
            """,
            (user_id,),
        )

        row = connection.execute(
            """
            SELECT
                character,
                response_style,
                response_length,
                voice_enabled,
                search_mode,
                roughness,
                custom_nickname
            FROM user_settings
            WHERE user_id = ?
            """,
            (user_id,),
        ).fetchone()

        connection.commit()

    if row is None:
        return DEFAULT_USER_SETTINGS.copy()

    return {
        "character": str(row[0]),
        "response_style": str(row[1]),
        "response_length": str(row[2]),
        "voice_enabled": bool(row[3]),
        "search_mode": str(row[4]),
        "roughness": str(row[5]),
        "custom_nickname": row[6],
    }


async def get_user_settings(
    user_id: int,
) -> dict[str, Any]:
    """Читает настройки без блокировки бота."""

    return await asyncio.to_thread(
        get_user_settings_sync,
        user_id,
    )
USER_SETTING_COLUMNS = {
    "character",
    "response_style",
    "response_length",
    "voice_enabled",
    "search_mode",
    "roughness",
    "custom_nickname",
}


def update_user_setting_sync(
    user_id: int,
    setting_name: str,
    setting_value: Any,
) -> None:
    """Сохраняет одну настройку пользователя."""

    if setting_name not in USER_SETTING_COLUMNS:
        raise ValueError(
            f"Неизвестная настройка: {setting_name}"
        )

    if setting_name == "voice_enabled":
        database_value: Any = int(
            bool(setting_value)
        )
    elif setting_name == "custom_nickname" and setting_value is None:
        database_value = None
    else:
        database_value = str(
            setting_value
        )

    with get_db_connection() as connection:
        connection.execute(
            """
            INSERT OR IGNORE INTO user_settings (
                user_id, roughness
            )
            VALUES (?, 'high')
            """,
            (user_id,),
        )

        connection.execute(
            f"""
            UPDATE user_settings
            SET {setting_name} = ?
            WHERE user_id = ?
            """,
            (
                database_value,
                user_id,
            ),
        )

        connection.commit()


async def update_user_setting(
    user_id: int,
    setting_name: str,
    setting_value: Any,
) -> None:
    """Сохраняет настройку без блокировки бота."""

    await asyncio.to_thread(
        update_user_setting_sync,
        user_id,
        setting_name,
        setting_value,
    )


DEFAULT_CHAT_SETTINGS = {
    "hard_mode_enabled": True,
    "hard_level": "normal",
    # Совпадает с HARD_REACTION_CHANCE / HARD_RANDOM_REPLY_CHANCE
    # и со значениями по умолчанию в схеме chat_settings ниже —
    # эти константы объявлены позже в файле, поэтому здесь просто
    # числа, а не ссылки на них.
    "reaction_chance": 0.70,
    "random_reply_chance": 0.16,
    "reactions_count": 0,
    "random_replies_count": 0,
    "trigger_replies_count": 0,
    "last_intervention_at": None,
    "weekly_report_enabled": False,
    "weekly_report_weekday": 6,
    "weekly_report_time": "21:00",
    "weekly_report_last_sent_date": None,
}

CHAT_SETTING_COLUMNS = {
    "hard_mode_enabled",
    "hard_level",
    "reaction_chance",
    "random_reply_chance",
    "weekly_report_enabled",
    "weekly_report_weekday",
    "weekly_report_time",
}


def get_chat_settings_sync(
    chat_id: int,
    chat_type: str = "group",
) -> dict[str, Any]:
    """Получает сохранённые настройки чата (хард-мод, уровень, статистика)."""

    with get_db_connection() as connection:
        connection.execute(
            """
            INSERT OR IGNORE INTO chats (chat_id, chat_type)
            VALUES (?, ?)
            """,
            (chat_id, chat_type),
        )

        connection.execute(
            """
            INSERT OR IGNORE INTO chat_settings (chat_id)
            VALUES (?)
            """,
            (chat_id,),
        )

        row = connection.execute(
            """
            SELECT
                hard_mode_enabled,
                hard_level,
                reaction_chance,
                random_reply_chance,
                reactions_count,
                random_replies_count,
                trigger_replies_count,
                last_intervention_at,
                weekly_report_enabled,
                weekly_report_weekday,
                weekly_report_time,
                weekly_report_last_sent_date
            FROM chat_settings
            WHERE chat_id = ?
            """,
            (chat_id,),
        ).fetchone()

        connection.commit()

    if row is None:
        return DEFAULT_CHAT_SETTINGS.copy()

    return {
        "hard_mode_enabled": bool(row[0]),
        "hard_level": str(row[1]),
        "reaction_chance": float(row[2]),
        "random_reply_chance": float(row[3]),
        "reactions_count": int(row[4]),
        "random_replies_count": int(row[5]),
        "trigger_replies_count": int(row[6]),
        "last_intervention_at": row[7],
        "weekly_report_enabled": bool(row[8]),
        "weekly_report_weekday": int(row[9]),
        "weekly_report_time": str(row[10]),
        "weekly_report_last_sent_date": row[11],
    }


async def get_chat_settings(
    chat_id: int,
    chat_type: str = "group",
) -> dict[str, Any]:
    """Читает настройки чата без блокировки бота."""

    return await asyncio.to_thread(
        get_chat_settings_sync,
        chat_id,
        chat_type,
    )


def update_chat_setting_sync(
    chat_id: int,
    setting_name: str,
    setting_value: Any,
    chat_type: str = "group",
) -> None:
    """Сохраняет одну настройку чата."""

    if setting_name not in CHAT_SETTING_COLUMNS:
        raise ValueError(
            f"Неизвестная настройка чата: {setting_name}"
        )

    if setting_name in ("hard_mode_enabled", "weekly_report_enabled"):
        database_value: Any = int(bool(setting_value))
    elif setting_name in ("reaction_chance", "random_reply_chance"):
        database_value = float(setting_value)
    elif setting_name == "weekly_report_weekday":
        database_value = int(setting_value)
    else:
        database_value = str(setting_value)

    with get_db_connection() as connection:
        connection.execute(
            """
            INSERT OR IGNORE INTO chats (chat_id, chat_type)
            VALUES (?, ?)
            """,
            (chat_id, chat_type),
        )

        connection.execute(
            """
            INSERT OR IGNORE INTO chat_settings (chat_id)
            VALUES (?)
            """,
            (chat_id,),
        )

        connection.execute(
            f"""
            UPDATE chat_settings
            SET {setting_name} = ?, updated_at = datetime('now')
            WHERE chat_id = ?
            """,
            (database_value, chat_id),
        )

        connection.commit()


async def update_chat_setting(
    chat_id: int,
    setting_name: str,
    setting_value: Any,
    chat_type: str = "group",
) -> None:
    """Сохраняет настройку чата без блокировки бота."""

    await asyncio.to_thread(
        update_chat_setting_sync,
        chat_id,
        setting_name,
        setting_value,
        chat_type,
    )


HARD_STAT_COUNTER_COLUMNS = {
    "reactions_count",
    "random_replies_count",
    "trigger_replies_count",
}


def increment_chat_hard_stat_sync(
    chat_id: int,
    counter_name: str,
    chat_type: str = "group",
) -> None:
    """Увеличивает один из счётчиков /hard_stats на единицу."""

    if counter_name not in HARD_STAT_COUNTER_COLUMNS:
        raise ValueError(
            f"Неизвестный счётчик хард-мода: {counter_name}"
        )

    with get_db_connection() as connection:
        connection.execute(
            """
            INSERT OR IGNORE INTO chats (chat_id, chat_type)
            VALUES (?, ?)
            """,
            (chat_id, chat_type),
        )

        connection.execute(
            """
            INSERT OR IGNORE INTO chat_settings (chat_id)
            VALUES (?)
            """,
            (chat_id,),
        )

        connection.execute(
            f"""
            UPDATE chat_settings
            SET {counter_name} = {counter_name} + 1,
                last_intervention_at = datetime('now')
            WHERE chat_id = ?
            """,
            (chat_id,),
        )

        connection.commit()


async def increment_chat_hard_stat(
    chat_id: int,
    counter_name: str,
    chat_type: str = "group",
) -> None:
    """Увеличивает счётчик хард-мода без блокировки бота."""

    await asyncio.to_thread(
        increment_chat_hard_stat_sync,
        chat_id,
        counter_name,
        chat_type,
    )


# ============================================================
# ЛЁГКАЯ ПАМЯТЬ ОБ УЧАСТНИКАХ ГРУППЫ (chat_member_profiles)
#
# Только безопасные числа, текст, который сам человек попросил
# запомнить, и шуточный архетип, который вручную задаёт админ.
# Никаких медицинских, финансовых, политических данных или
# содержимого личных файлов — см. MEMORY.md проекта.
# ============================================================

MAX_SELF_REPORTED_FACTS = 5

# Относительно немного сообщений — компания из нескольких
# человек, а не публичный канал, поэтому пороги невысокие.
RELATIONSHIP_LEVEL_THRESHOLDS = (
    (150, 3),  # старожил
    (30, 2),  # постоянный участник
    (5, 1),  # знакомый
)


def compute_relationship_level(
    total_messages: int,
) -> int:
    """Определяет уровень знакомства по числу сообщений в чате."""

    for threshold, level in RELATIONSHIP_LEVEL_THRESHOLDS:
        if total_messages >= threshold:
            return level

    return 0


def _row_to_member_profile(row: Any) -> dict[str, Any]:
    self_reported_raw = row[11]

    try:
        self_reported_facts = (
            json.loads(self_reported_raw)
            if self_reported_raw
            else []
        )
    except (TypeError, ValueError):
        self_reported_facts = []

    return {
        "user_id": int(row[0]),
        "current_display_name": row[1],
        "username": row[2],
        "first_seen_at": row[3],
        "last_seen_at": row[4],
        "total_messages": int(row[5]),
        "total_voice_messages": int(row[6]),
        "total_photos": int(row[7]),
        "total_stickers": int(row[8]),
        "replies_to_bot": int(row[9]),
        "insults_to_bot": int(row[10]),
        "self_reported_facts": self_reported_facts,
        "joke_archetype": row[12],
        "relationship_level": int(row[13]),
        "current_title": row[14],
    }


def get_member_profile_sync(
    chat_id: int,
    user_id: int,
) -> dict[str, Any] | None:
    """Получает профиль участника чата, если он уже известен боту."""

    with get_db_connection() as connection:
        row = connection.execute(
            """
            SELECT
                user_id,
                current_display_name,
                username,
                first_seen_at,
                last_seen_at,
                total_messages,
                total_voice_messages,
                total_photos,
                total_stickers,
                replies_to_bot,
                insults_to_bot,
                self_reported_facts,
                joke_archetype,
                relationship_level,
                current_title
            FROM chat_member_profiles
            WHERE chat_id = ? AND user_id = ?
            """,
            (chat_id, user_id),
        ).fetchone()

    if row is None:
        return None

    return _row_to_member_profile(row)


async def get_member_profile(
    chat_id: int,
    user_id: int,
) -> dict[str, Any] | None:
    """Читает профиль участника без блокировки бота."""

    return await asyncio.to_thread(
        get_member_profile_sync,
        chat_id,
        user_id,
    )


def touch_member_profile_sync(
    chat_id: int,
    user_id: int,
    chat_type: str,
    display_name: str,
    username: str | None,
) -> None:
    """
    Отмечает активность участника: обновляет имя, время и счётчик
    сообщений, пересчитывает уровень знакомства.
    """

    with get_db_connection() as connection:
        _ensure_member_profile_row(
            connection, chat_id, user_id, chat_type
        )

        connection.execute(
            """
            UPDATE chat_member_profiles
            SET
                current_display_name = ?,
                username = ?,
                last_seen_at = datetime('now'),
                total_messages = total_messages + 1,
                updated_at = datetime('now')
            WHERE chat_id = ? AND user_id = ?
            """,
            (display_name, username, chat_id, user_id),
        )

        new_total = connection.execute(
            """
            SELECT total_messages
            FROM chat_member_profiles
            WHERE chat_id = ? AND user_id = ?
            """,
            (chat_id, user_id),
        ).fetchone()[0]

        connection.execute(
            """
            UPDATE chat_member_profiles
            SET relationship_level = ?
            WHERE chat_id = ? AND user_id = ?
            """,
            (
                compute_relationship_level(new_total),
                chat_id,
                user_id,
            ),
        )

        connection.commit()


async def touch_member_profile(
    chat_id: int,
    user_id: int,
    chat_type: str,
    display_name: str,
    username: str | None,
) -> None:
    """Отмечает активность участника без блокировки бота."""

    await asyncio.to_thread(
        touch_member_profile_sync,
        chat_id,
        user_id,
        chat_type,
        display_name,
        username,
    )


def _ensure_member_profile_row(
    connection: sqlite3.Connection,
    chat_id: int,
    user_id: int,
    chat_type: str,
) -> None:
    """
    Гарантирует, что chats/users/chat_member_profiles знают об этой
    паре — chat_member_profiles ссылается на chats и users внешними
    ключами (foreign_keys=ON), и без этого шага вставка упадёт, если
    человек ещё не писал обычных сообщений в этом чате (например,
    его самое первое действие — сразу команда /remember_me).
    """

    connection.execute(
        """
        INSERT OR IGNORE INTO chats (chat_id, chat_type)
        VALUES (?, ?)
        """,
        (chat_id, chat_type),
    )
    connection.execute(
        """
        INSERT OR IGNORE INTO users (user_id)
        VALUES (?)
        """,
        (user_id,),
    )
    connection.execute(
        """
        INSERT OR IGNORE INTO chat_member_profiles (chat_id, user_id)
        VALUES (?, ?)
        """,
        (chat_id, user_id),
    )


def set_member_joke_archetype_sync(
    chat_id: int,
    user_id: int,
    archetype: str | None,
    chat_type: str = "group",
) -> None:
    """Задаёт шуточный архетип участника — только вручную, только админом."""

    with get_db_connection() as connection:
        _ensure_member_profile_row(
            connection, chat_id, user_id, chat_type
        )
        connection.execute(
            """
            UPDATE chat_member_profiles
            SET joke_archetype = ?, updated_at = datetime('now')
            WHERE chat_id = ? AND user_id = ?
            """,
            (archetype, chat_id, user_id),
        )
        connection.commit()


async def set_member_joke_archetype(
    chat_id: int,
    user_id: int,
    archetype: str | None,
    chat_type: str = "group",
) -> None:
    """Задаёт архетип без блокировки бота."""

    await asyncio.to_thread(
        set_member_joke_archetype_sync,
        chat_id,
        user_id,
        archetype,
        chat_type,
    )


def set_member_title_sync(
    chat_id: int,
    user_id: int,
    title: str,
    chat_type: str = "group",
) -> None:
    """Сохраняет текущий титул участника (/title), заменяя предыдущий."""

    with get_db_connection() as connection:
        _ensure_member_profile_row(
            connection, chat_id, user_id, chat_type
        )
        connection.execute(
            """
            UPDATE chat_member_profiles
            SET current_title = ?, updated_at = datetime('now')
            WHERE chat_id = ? AND user_id = ?
            """,
            (title, chat_id, user_id),
        )
        connection.commit()


async def set_member_title(
    chat_id: int,
    user_id: int,
    title: str,
    chat_type: str = "group",
) -> None:
    """Сохраняет титул без блокировки бота."""

    await asyncio.to_thread(
        set_member_title_sync,
        chat_id,
        user_id,
        title,
        chat_type,
    )


def get_daily_title_assignment_sync(
    chat_id: int,
    date: str,
) -> dict[str, Any] | None:
    with get_db_connection() as connection:
        row = connection.execute(
            """
            SELECT user_id, title, assigned_at, announced_at
            FROM daily_title_assignments
            WHERE chat_id = ? AND date = ?
            """,
            (chat_id, date),
        ).fetchone()

    if row is None:
        return None
    return {
        "user_id": int(row[0]),
        "title": str(row[1]),
        "assigned_at": str(row[2]),
        "announced_at": (str(row[3]) if row[3] else None),
    }


async def get_daily_title_assignment(
    chat_id: int,
    date: str,
) -> dict[str, Any] | None:
    return await asyncio.to_thread(
        get_daily_title_assignment_sync,
        chat_id,
        date,
    )


def try_assign_daily_title_sync(
    chat_id: int,
    date: str,
    user_id: int,
    title: str,
) -> bool:
    """Атомарно фиксирует daily title и обновляет текущий титул участника."""

    with get_db_connection() as connection:
        cursor = connection.execute(
            """
            INSERT OR IGNORE INTO daily_title_assignments (
                chat_id, date, user_id, title
            ) VALUES (?, ?, ?, ?)
            """,
            (chat_id, date, user_id, title),
        )

        if cursor.rowcount != 1:
            connection.commit()
            return False

        update_cursor = connection.execute(
            """
            UPDATE chat_member_profiles
            SET current_title = ?, updated_at = datetime('now')
            WHERE chat_id = ? AND user_id = ?
            """,
            (title, chat_id, user_id),
        )
        if update_cursor.rowcount != 1:
            connection.rollback()
            return False
        connection.commit()
        return True


async def try_assign_daily_title(
    chat_id: int,
    date: str,
    user_id: int,
    title: str,
) -> bool:
    return await asyncio.to_thread(
        try_assign_daily_title_sync,
        chat_id,
        date,
        user_id,
        title,
    )


def append_self_reported_fact_sync(
    chat_id: int,
    user_id: int,
    fact: str,
    chat_type: str = "group",
) -> None:
    """Добавляет факт, который человек сам попросил запомнить (максимум 5)."""

    with get_db_connection() as connection:
        _ensure_member_profile_row(
            connection, chat_id, user_id, chat_type
        )

        row = connection.execute(
            """
            SELECT self_reported_facts
            FROM chat_member_profiles
            WHERE chat_id = ? AND user_id = ?
            """,
            (chat_id, user_id),
        ).fetchone()

        try:
            facts = (
                json.loads(row[0])
                if row and row[0]
                else []
            )
        except (TypeError, ValueError):
            facts = []

        facts.append(fact)
        facts = facts[-MAX_SELF_REPORTED_FACTS:]

        connection.execute(
            """
            UPDATE chat_member_profiles
            SET self_reported_facts = ?, updated_at = datetime('now')
            WHERE chat_id = ? AND user_id = ?
            """,
            (json.dumps(facts, ensure_ascii=False), chat_id, user_id),
        )
        connection.commit()


async def append_self_reported_fact(
    chat_id: int,
    user_id: int,
    fact: str,
    chat_type: str = "group",
) -> None:
    """Добавляет факт без блокировки бота."""

    await asyncio.to_thread(
        append_self_reported_fact_sync,
        chat_id,
        user_id,
        fact,
        chat_type,
    )


def delete_member_profile_sync(
    chat_id: int,
    user_id: int,
) -> None:
    """Удаляет профиль участника в этой группе (/forget_me)."""

    with get_db_connection() as connection:
        connection.execute(
            """
            DELETE FROM chat_member_profiles
            WHERE chat_id = ? AND user_id = ?
            """,
            (chat_id, user_id),
        )
        connection.commit()


async def delete_member_profile(
    chat_id: int,
    user_id: int,
) -> None:
    """Удаляет профиль без блокировки бота."""

    await asyncio.to_thread(
        delete_member_profile_sync,
        chat_id,
        user_id,
    )


def list_chat_member_profiles_sync(
    chat_id: int,
    limit: int = 8,
) -> list[dict[str, Any]]:
    """Возвращает известных активных участников чата (для /people)."""

    with get_db_connection() as connection:
        rows = connection.execute(
            """
            SELECT
                user_id,
                current_display_name,
                username,
                first_seen_at,
                last_seen_at,
                total_messages,
                total_voice_messages,
                total_photos,
                total_stickers,
                replies_to_bot,
                insults_to_bot,
                self_reported_facts,
                joke_archetype,
                relationship_level,
                current_title
            FROM chat_member_profiles
            WHERE chat_id = ?
            ORDER BY total_messages DESC
            LIMIT ?
            """,
            (chat_id, limit),
        ).fetchall()

    return [
        _row_to_member_profile(row)
        for row in rows
    ]


async def list_chat_member_profiles(
    chat_id: int,
    limit: int = 8,
) -> list[dict[str, Any]]:
    """Читает список участников без блокировки бота."""

    return await asyncio.to_thread(
        list_chat_member_profiles_sync,
        chat_id,
        limit,
    )


# ============================================================
# НЕДЕЛЬНАЯ АНАЛИТИКА (chat_activity_daily)
#
# Храним агрегаты по дням, а не полный текст сообщений —
# это статистика активности, а не архив переписки.
# ============================================================

MSK_TIMEZONE = timezone(timedelta(hours=3))

# Тёмное время суток по МСК — то же окно, что и у quiet hours,
# для единообразия "ночной активности" в разных местах бота.
NIGHT_ACTIVITY_START_MSK = 0
NIGHT_ACTIVITY_END_MSK = 6

CHAT_ACTIVITY_COLUMNS = {
    "messages",
    "text_characters",
    "voice_messages",
    "voice_duration_seconds",
    "photos",
    "videos",
    "stickers",
    "documents",
    "replies",
    "replies_to_bot",
    "commands",
    "night_messages",
    "questions",
    "links",
    "edited_messages",
}

_URL_RE = re.compile(r"https?://\S+|www\.\S+", re.IGNORECASE)


def current_msk_datetime() -> datetime:
    """Возвращает текущее время по МСК."""

    return datetime.now(MSK_TIMEZONE)


def current_msk_date_str() -> str:
    """Возвращает текущую дату по МСК в формате YYYY-MM-DD."""

    return current_msk_datetime().strftime("%Y-%m-%d")


def is_night_activity_now_msk() -> bool:
    """Проверяет, попадает ли текущий час по МСК в ночное окно."""

    hour = current_msk_datetime().hour
    return NIGHT_ACTIVITY_START_MSK <= hour < NIGHT_ACTIVITY_END_MSK


def build_text_activity_deltas(
    text: str,
    *,
    is_reply_to_bot: bool = False,
) -> dict[str, int]:
    """Строит набор приращений активности для одного текстового сообщения."""

    return {
        "messages": 1,
        "text_characters": len(text),
        "replies_to_bot": 1 if is_reply_to_bot else 0,
        "night_messages": 1 if is_night_activity_now_msk() else 0,
        "questions": 1 if "?" in text else 0,
        "links": 1 if _URL_RE.search(text) else 0,
    }


def increment_chat_activity_sync(
    chat_id: int,
    user_id: int,
    chat_type: str,
    date_str: str,
    **deltas: int,
) -> None:
    """Прибавляет дневные счётчики активности участника."""

    unknown_columns = set(deltas) - CHAT_ACTIVITY_COLUMNS

    if unknown_columns:
        raise ValueError(
            f"Неизвестные поля активности: {unknown_columns}"
        )

    with get_db_connection() as connection:
        connection.execute(
            """
            INSERT OR IGNORE INTO chats (chat_id, chat_type)
            VALUES (?, ?)
            """,
            (chat_id, chat_type),
        )
        connection.execute(
            """
            INSERT OR IGNORE INTO users (user_id)
            VALUES (?)
            """,
            (user_id,),
        )
        connection.execute(
            """
            INSERT OR IGNORE INTO chat_activity_daily (chat_id, user_id, date)
            VALUES (?, ?, ?)
            """,
            (chat_id, user_id, date_str),
        )

        for column, delta in deltas.items():
            if not delta:
                continue

            connection.execute(
                f"""
                UPDATE chat_activity_daily
                SET {column} = {column} + ?
                WHERE chat_id = ? AND user_id = ? AND date = ?
                """,
                (delta, chat_id, user_id, date_str),
            )

        connection.commit()


async def increment_chat_activity(
    chat_id: int,
    user_id: int,
    chat_type: str,
    date_str: str,
    **deltas: int,
) -> None:
    """Прибавляет дневную активность без блокировки бота."""

    await asyncio.to_thread(
        increment_chat_activity_sync,
        chat_id,
        user_id,
        chat_type,
        date_str,
        **deltas,
    )


def get_week_date_range(
    reference_date: str | None = None,
) -> tuple[str, str]:
    """Возвращает (начало, конец) последних семи дней по МСК включительно."""

    if reference_date is not None:
        end = datetime.strptime(reference_date, "%Y-%m-%d")
    else:
        end = current_msk_datetime().replace(
            hour=0, minute=0, second=0, microsecond=0
        )

    start = end - timedelta(days=6)

    return (
        start.strftime("%Y-%m-%d"),
        end.strftime("%Y-%m-%d"),
    )


def get_weekly_activity_sync(
    chat_id: int,
    start_date: str,
    end_date: str,
) -> list[dict[str, Any]]:
    """Возвращает агрегированную недельную активность по каждому участнику."""

    with get_db_connection() as connection:
        rows = connection.execute(
            """
            SELECT
                user_id,
                SUM(messages),
                SUM(text_characters),
                SUM(voice_messages),
                SUM(voice_duration_seconds),
                SUM(photos),
                SUM(videos),
                SUM(stickers),
                SUM(documents),
                SUM(replies),
                SUM(replies_to_bot),
                SUM(commands),
                SUM(night_messages),
                SUM(questions),
                SUM(links),
                SUM(edited_messages)
            FROM chat_activity_daily
            WHERE chat_id = ? AND date BETWEEN ? AND ?
            GROUP BY user_id
            """,
            (chat_id, start_date, end_date),
        ).fetchall()

    columns = (
        "messages",
        "text_characters",
        "voice_messages",
        "voice_duration_seconds",
        "photos",
        "videos",
        "stickers",
        "documents",
        "replies",
        "replies_to_bot",
        "commands",
        "night_messages",
        "questions",
        "links",
        "edited_messages",
    )

    results = []

    for row in rows:
        entry = {"user_id": row[0]}
        entry.update(
            {
                column: int(value or 0)
                for column, value in zip(columns, row[1:])
            }
        )
        results.append(entry)

    return results


async def get_weekly_activity(
    chat_id: int,
    start_date: str,
    end_date: str,
) -> list[dict[str, Any]]:
    """Читает недельную активность без блокировки бота."""

    return await asyncio.to_thread(
        get_weekly_activity_sync,
        chat_id,
        start_date,
        end_date,
    )


# Достаточно данных, чтобы награда что-то значила, а не досталась
# за одно случайное сообщение.
AWARD_MINIMUM_MESSAGES_FOR_STYLE_AWARDS = 5

# Простые "больше всех по колонке" награды: ключ награды -> колонка.
_SIMPLE_LEADER_AWARDS = {
    "chat_leader": "messages",
    "voice_leader": "voice_messages",
    "wall_of_text": "text_characters",
    "bot_caller": "replies_to_bot",
    "pixel_provider": "photos",
    "night_watch": "night_messages",
    "argument_lord": "replies",
    "reality_editor": "edited_messages",
    "archivist": "documents",
    "voice_from_the_deep": "voice_duration_seconds",
}


def _leader_by_column(
    weekly_activity: list[dict[str, Any]],
    column: str,
) -> int | None:
    """Возвращает user_id с максимальным значением колонки (если > 0)."""

    best_user_id = None
    best_value = 0

    for entry in weekly_activity:
        value = entry.get(column, 0)

        if value > best_value:
            best_value = value
            best_user_id = entry["user_id"]

    return best_user_id


def compute_weekly_awards(
    weekly_activity: list[dict[str, Any]],
    known_members: list[dict[str, Any]],
) -> list[tuple[str, int]]:
    """
    Определяет победителя каждой награды по числовым данным недели.

    Не более одного победителя на награду, награда не выдаётся,
    если ни у кого нет нужного минимума участия.
    """

    awards: list[tuple[str, int]] = []

    for award_key, column in _SIMPLE_LEADER_AWARDS.items():
        winner = _leader_by_column(weekly_activity, column)

        if winner is not None:
            awards.append((award_key, winner))

    # Односложный мудрец: минимум пять сообщений, самая короткая
    # средняя длина сообщения — это стиль, а не разовая случайность.
    best_user_id = None
    best_avg_length: float | None = None

    for entry in weekly_activity:
        messages = entry.get("messages", 0)

        if messages < AWARD_MINIMUM_MESSAGES_FOR_STYLE_AWARDS:
            continue

        avg_length = entry.get("text_characters", 0) / messages

        if best_avg_length is None or avg_length < best_avg_length:
            best_avg_length = avg_length
            best_user_id = entry["user_id"]

    if best_user_id is not None:
        awards.append(("one_word_sage", best_user_id))

    # Куколд-наблюдатель: боту уже известен, раньше писал,
    # на этой неделе — ни одного сообщения.
    active_user_ids = {
        entry["user_id"]
        for entry in weekly_activity
        if entry.get("messages", 0) > 0
    }

    silent_candidates = [
        member
        for member in known_members
        if member.get("total_messages", 0) > 0
        and member["user_id"] not in active_user_ids
    ]

    if silent_candidates:
        chosen = random.choice(silent_candidates)
        awards.append(("silent_observer", chosen["user_id"]))

    return awards


_AWARD_KEY_ORDER = {
    key: index
    for index, key in enumerate(AWARD_TEMPLATES.keys())
}


def get_or_create_weekly_awards_sync(
    chat_id: int,
    week_start: str,
    weekly_activity: list[dict[str, Any]],
    known_members: list[dict[str, Any]],
) -> list[tuple[str, int]]:
    """
    Возвращает зафиксированных победителей недели, досчитывая только
    те награды, у которых ещё нет сохранённого победителя.

    PRIMARY KEY (chat_id, week_start, award_key) в weekly_award_winners
    гарантирует, что при повторном вызове (в том числе одновременном —
    из /week, /awards и автоотчёта) победитель каждой награды остаётся
    один и тот же на всю неделю, а не переопределяется случайным
    выбором заново.
    """

    computed = compute_weekly_awards(weekly_activity, known_members)

    with get_db_connection() as connection:
        for award_key, user_id in computed:
            connection.execute(
                """
                INSERT OR IGNORE INTO weekly_award_winners
                    (chat_id, week_start, award_key, user_id)
                VALUES (?, ?, ?, ?)
                """,
                (chat_id, week_start, award_key, user_id),
            )

        connection.commit()

        rows = connection.execute(
            """
            SELECT award_key, user_id
            FROM weekly_award_winners
            WHERE chat_id = ? AND week_start = ?
            """,
            (chat_id, week_start),
        ).fetchall()

    persisted = [(row[0], row[1]) for row in rows]
    persisted.sort(
        key=lambda item: _AWARD_KEY_ORDER.get(item[0], len(_AWARD_KEY_ORDER))
    )
    return persisted


async def get_or_create_weekly_awards(
    chat_id: int,
    week_start: str,
    weekly_activity: list[dict[str, Any]],
    known_members: list[dict[str, Any]],
) -> list[tuple[str, int]]:
    """Читает/фиксирует победителей недели без блокировки бота."""

    return await asyncio.to_thread(
        get_or_create_weekly_awards_sync,
        chat_id,
        week_start,
        weekly_activity,
        known_members,
    )


def format_awards_message(
    awards: list[tuple[str, int]],
    display_names: dict[int, str],
) -> str:
    """Превращает список наград в готовый текст сообщения."""

    if not awards:
        return "На этой неделе данных маловато — награды не набежали."

    lines = ["Шуточные награды недели:"]

    for award_key, user_id in awards:
        name = display_names.get(user_id, "Неизвестный герой")
        template = random.choice(
            AWARD_TEMPLATES.get(award_key, ["{name} получает награду."])
        )
        lines.append(
            f"\n🏆 {AWARD_LABELS.get(award_key, award_key)}\n"
            + template.format(name=name)
        )

    return "\n".join(lines)


CHARACTER_LABELS = {
    "classic": "🥚 Классический",
    "rus": "🗿 Древний рус",
    "professor": "🎓 Профессор",
    "chaos": "🤡 Безумный",
    "calm": "🧘 Спокойный",
}

STYLE_LABELS = {
    "normal": "Нормальный",
    "bold": "Дерзкий",
    "serious": "Серьёзный",
}

LENGTH_LABELS = {
    "short": "Кратко",
    "normal": "Обычно",
    "detailed": "Подробно",
}

SEARCH_MODE_LABELS = {
    "auto": "Автоматически",
    "button": "Только по кнопке",
}

ROUGHNESS_LABELS = {
    "low": "Низкая",
    "medium": "Средняя",
    "high": "Высокая",
}


def build_settings_keyboard(
    settings: dict[str, Any],
) -> InlineKeyboardMarkup:
    """Создаёт главное меню настроек."""

    character = CHARACTER_LABELS.get(
        str(settings.get("character")),
        CHARACTER_LABELS["classic"],
    )

    style = STYLE_LABELS.get(
        str(settings.get("response_style")),
        STYLE_LABELS["bold"],
    )

    response_length = LENGTH_LABELS.get(
        str(settings.get("response_length")),
        LENGTH_LABELS["normal"],
    )

    search_mode = SEARCH_MODE_LABELS.get(
        str(settings.get("search_mode")),
        SEARCH_MODE_LABELS["button"],
    )

    roughness = ROUGHNESS_LABELS.get(
        str(settings.get("roughness")),
        ROUGHNESS_LABELS["medium"],
    )

    voice = (
        "Включён"
        if settings.get("voice_enabled")
        else "Выключен"
    )

    return InlineKeyboardMarkup(
        [
            [
                InlineKeyboardButton(
                    f"🎭 Персонаж: {character}",
                    callback_data="settings_character",
                )
            ],
            [
                InlineKeyboardButton(
                    f"💬 Стиль: {style}",
                    callback_data="settings_style",
                )
            ],
            [
                InlineKeyboardButton(
                    f"📏 Длина: {response_length}",
                    callback_data="settings_length",
                )
            ],
            [
                InlineKeyboardButton(
                    f"🔊 Голос: {voice}",
                    callback_data="settings_voice",
                )
            ],
            [
                InlineKeyboardButton(
                    f"🔎 Поиск: {search_mode}",
                    callback_data="settings_search",
                )
            ],
            [
                InlineKeyboardButton(
                    f"🤬 Грубость: {roughness}",
                    callback_data="settings_roughness",
                )
            ],
            [
                InlineKeyboardButton(
                    "♻️ Сбросить настройки",
                    callback_data="settings_reset",
                )
            ],
        ]
    )  
def increment_stat_sync(
    stat_name: str,
    amount: int = 1,
) -> None:
    """Увеличивает постоянный счётчик."""

    with get_db_connection() as connection:
        connection.execute(
            """
            INSERT INTO stats (
                name,
                value
            )
            VALUES (?, ?)
            ON CONFLICT(name)
            DO UPDATE SET
                value = value + excluded.value
            """,
            (
                stat_name,
                amount,
            ),
        )

        connection.commit()


async def increment_stat(
    stat_name: str,
    amount: int = 1,
) -> None:
    """Записывает статистику без блокировки бота."""

    await asyncio.to_thread(
        increment_stat_sync,
        stat_name,
        amount,
    )


def register_user_and_chat_sync(
    user_id: int | None,
    chat_id: int | None,
    chat_type: str,
) -> None:
    """Запоминает уникального пользователя и чат."""

    with get_db_connection() as connection:
        if user_id is not None:
            connection.execute(
                """
                INSERT OR IGNORE INTO users (
                    user_id
                )
                VALUES (?)
                """,
                (user_id,),
            )

        if chat_id is not None:
            connection.execute(
                """
                INSERT OR IGNORE INTO chats (
                    chat_id,
                    chat_type
                )
                VALUES (?, ?)
                """,
                (
                    chat_id,
                    chat_type,
                ),
            )

        connection.commit()


async def register_user_and_chat(
    update: Update,
) -> None:
    """Записывает уникального пользователя и чат."""

    user_id = (
        update.effective_user.id
        if update.effective_user
        else None
    )

    chat_id = (
        update.effective_chat.id
        if update.effective_chat
        else None
    )

    chat_type = (
        update.effective_chat.type
        if update.effective_chat
        else "unknown"
    )

    await asyncio.to_thread(
        register_user_and_chat_sync,
        user_id,
        chat_id,
        chat_type,
    )


def get_stats_snapshot_sync() -> dict[str, int]:
    """Читает всю накопленную статистику."""

    with get_db_connection() as connection:
        stat_rows = connection.execute(
            """
            SELECT name, value
            FROM stats
            """
        ).fetchall()

        result = {
            str(name): int(value)
            for name, value in stat_rows
        }

        result["unique_users"] = int(
            connection.execute(
                """
                SELECT COUNT(*)
                FROM users
                """
            ).fetchone()[0]
        )

        result["private_chats"] = int(
            connection.execute(
                """
                SELECT COUNT(*)
                FROM chats
                WHERE chat_type = ?
                """,
                ("private",),
            ).fetchone()[0]
        )

        result["groups"] = int(
            connection.execute(
                """
                SELECT COUNT(*)
                FROM chats
                WHERE chat_type IN (?, ?)
                """,
                (
                    "group",
                    "supergroup",
                ),
            ).fetchone()[0]
        )

    return result


async def get_stats_snapshot() -> dict[str, int]:
    """Читает статистику без блокировки бота."""

    return await asyncio.to_thread(
        get_stats_snapshot_sync
    )
# Максимальный размер принимаемого файла — 20 МБ
MAX_FILE_SIZE = 20 * 1024 * 1024

# ==============================
# ПАМЯТЬ БОТА
# ==============================

# Максимальный объём текста из DOCX, XLSX, CSV и TXT
MAX_EXTRACTED_CHARS = 50_000
# ============================================================
# КРАТКОВРЕМЕННАЯ ПАМЯТЬ
# ============================================================

# В группе бот помнит разговор 15 минут (используется и для обычного
# контекста, и как окно для /recap и /story — не более 30 сообщений,
# см. GROUP_MEMORY_MAX_MESSAGES ниже).
GROUP_MEMORY_SECONDS = 15 * 60

# В личной переписке бот помнит текущую задачу 15 минут
PRIVATE_MEMORY_SECONDS = 15 * 60

# Не храним слишком много сообщений
GROUP_MEMORY_MAX_MESSAGES = 30
PRIVATE_MEMORY_MAX_MESSAGES = 40

# Память по идентификатору группы
GROUP_MEMORY: dict[int, deque] = defaultdict(deque)

# Память по идентификатору пользователя
PRIVATE_MEMORY: dict[int, deque] = defaultdict(deque)


if not TELEGRAM_TOKEN:
    raise RuntimeError(
        "В файле .env не найден TELEGRAM_BOT_TOKEN"
    )

if not GEMINI_API_KEY:
    raise RuntimeError(
        "В файле .env не найден GEMINI_API_KEY"
    )


gemini_client = genai.Client(
    api_key=GEMINI_API_KEY
)


logging.basicConfig(
    format="%(asctime)s - %(levelname)s - %(message)s",
    level=logging.INFO,
)


def get_response_token_limit(
    user_settings: dict[str, Any] | None,
    normal_tokens: int = 360,
) -> int:
    """Возвращает лимит ответа с учётом настройки длины."""

    settings = DEFAULT_USER_SETTINGS.copy()

    if user_settings:
        settings.update(
            user_settings
        )

    response_length = str(
        settings.get(
            "response_length",
            "normal",
        )
    )

    token_limits = {
        "short": max(
            160,
            int(normal_tokens * 0.65),
        ),
        "normal": normal_tokens,
        "detailed": min(
            900,
            int(normal_tokens * 1.7),
        ),
    }

    return token_limits.get(
        response_length,
        normal_tokens,
    )    

# ============================================================
# ФУНКЦИИ КРАТКОВРЕМЕННОЙ ПАМЯТИ
# ============================================================

def clean_memory(
    messages: deque,
    memory_seconds: int,
) -> None:
    """Удаляет сообщения, которые уже слишком старые."""

    current_time = time.monotonic()

    while messages:
        message_time = messages[0][0]

        if current_time - message_time <= memory_seconds:
            break

        messages.popleft()


def remember_message(
    memory_store: dict[int, deque],
    memory_id: int,
    role: str,
    text: str,
    memory_seconds: int,
    max_messages: int,
    author_name: str = "",
) -> None:
    """Добавляет одно сообщение в память."""

    if not text:
        return

    messages = memory_store[memory_id]

    clean_memory(
        messages,
        memory_seconds,
    )

    clean_text = text.strip()[:800]

    messages.append(
        (
            time.monotonic(),
            role,
            author_name,
            clean_text,
        )
    )

    while len(messages) > max_messages:
        messages.popleft()


def build_memory_context(
    memory_store: dict[int, deque],
    memory_id: int,
    memory_seconds: int,
) -> str:
    """Собирает ещё не устаревшие сообщения в один контекст."""

    messages = memory_store[memory_id]

    clean_memory(
        messages,
        memory_seconds,
    )

    if not messages:
        return ""

    context_lines = []

    for _, role, author_name, text in messages:
        if role == "assistant":
            speaker = "Яйцеслав"
        elif author_name:
            speaker = author_name
        else:
            speaker = "Пользователь"

        context_lines.append(
            f"{speaker}: {text}"
        )

    return "\n".join(context_lines)


# ============================================================
# ПЕРИОДИЧЕСКАЯ ОЧИСТКА ПАМЯТИ ПРОЦЕССА
#
# REQUEST_TIMES, LAST_LIMIT_WARNING, GROUP_MEMORY и PRIVATE_MEMORY
# сами чистят СОДЕРЖИМОЕ по TTL, но ключи (chat_id/user_id) в этих
# defaultdict оставались навсегда, даже когда всё внутри устарело —
# на долгоживущем процессе Railway это медленная утечка памяти.
# ============================================================

CLEANUP_INTERVAL_SECONDS = 30 * 60
MEMORY_CLEANUP_MAX_AGE_SECONDS = 6 * 60 * 60


def cleanup_in_memory_state(
    max_age_seconds: float = MEMORY_CLEANUP_MAX_AGE_SECONDS,
) -> dict[str, int]:
    """Удаляет ключи чатов/пользователей, неактивные дольше max_age_seconds."""

    now = time.monotonic()

    stale_request_keys = [
        key
        for key, times in REQUEST_TIMES.items()
        if not times or (now - times[-1]) > max_age_seconds
    ]
    for key in stale_request_keys:
        REQUEST_TIMES.pop(key, None)

    stale_warning_keys = [
        key
        for key, last_time in LAST_LIMIT_WARNING.items()
        if (now - last_time) > max_age_seconds
    ]
    for key in stale_warning_keys:
        LAST_LIMIT_WARNING.pop(key, None)

    stale_memory_ids = 0

    for memory_store, memory_seconds in (
        (GROUP_MEMORY, GROUP_MEMORY_SECONDS),
        (PRIVATE_MEMORY, PRIVATE_MEMORY_SECONDS),
    ):
        empty_ids = []

        for memory_id, messages in memory_store.items():
            clean_memory(messages, memory_seconds)

            if not messages:
                empty_ids.append(memory_id)

        for memory_id in empty_ids:
            memory_store.pop(memory_id, None)

        stale_memory_ids += len(empty_ids)

    stale_humor_chats = humor_engine.prune_stale_state(
        max_age_seconds=max_age_seconds
    )

    stale_random_reply_chats = [
        chat_id
        for chat_id, history in GROUP_RANDOM_REPLY_TIMES.items()
        if not history or (now - history[-1]) > max_age_seconds
    ]
    for chat_id in stale_random_reply_chats:
        GROUP_RANDOM_REPLY_TIMES.pop(chat_id, None)
        GROUP_IGNORED_STREAK.pop(chat_id, None)

    stale_serious_chats = [
        chat_id
        for chat_id, last_serious_at in GROUP_LAST_SERIOUS_AT.items()
        if (now - last_serious_at) > max_age_seconds
    ]
    for chat_id in stale_serious_chats:
        GROUP_LAST_SERIOUS_AT.pop(chat_id, None)

    stale_trigger_user_keys = [
        key
        for key, last_time in TRIGGER_REPLY_LAST_BY_USER.items()
        if (now - last_time) > max_age_seconds
    ]
    for key in stale_trigger_user_keys:
        TRIGGER_REPLY_LAST_BY_USER.pop(key, None)

    stale_last_message_keys = [
        key
        for key, (recorded_at, _text) in LAST_USER_TEXT_MESSAGE.items()
        if now - recorded_at > max_age_seconds
    ]
    for key in stale_last_message_keys:
        LAST_USER_TEXT_MESSAGE.pop(key, None)

    stale_duel_tokens = [
        token
        for token, duel in PENDING_DUELS.items()
        if now - float(duel.get("created_at", 0.0)) > PENDING_DUEL_TTL_SECONDS
    ]
    for token in stale_duel_tokens:
        PENDING_DUELS.pop(token, None)

    stale_story_chats = [
        chat_id
        for chat_id, last_updated in STORY_LAST_UPDATED.items()
        if now - last_updated > max_age_seconds
    ]
    for chat_id in stale_story_chats:
        STORY_STATE.pop(chat_id, None)
        STORY_LAST_UPDATED.pop(chat_id, None)

    stale_state_chats = state_engine.prune_stale_state(
        max_age_seconds, now=now
    )
    stale_passive_chats = passive_engine.prune_stale_state(
        max_age_seconds, now=now
    )
    stale_aggression_keys = aggression_engine.prune_stale_state(
        max_age_seconds, now=now
    )
    stale_hostile_streaks = hostile_streak_engine.prune_stale_state(
        max_age_seconds, now=now
    )
    stale_length_chats = style_engine.prune_stale_state(
        max_age_seconds, now=now
    )

    return {
        "request_time_keys": len(stale_request_keys),
        "warning_keys": len(stale_warning_keys),
        "memory_ids": stale_memory_ids,
        "humor_chats": stale_humor_chats,
        "random_reply_chats": len(stale_random_reply_chats),
        "serious_chats": len(stale_serious_chats),
        "trigger_user_keys": len(stale_trigger_user_keys),
        "last_user_messages": len(stale_last_message_keys),
        "pending_duels": len(stale_duel_tokens),
        "story_chats": len(stale_story_chats),
        "state_chats": stale_state_chats,
        "passive_chats": stale_passive_chats,
        "aggression_keys": stale_aggression_keys,
        "hostile_streaks": stale_hostile_streaks,
        "length_chats": stale_length_chats,
    }


async def periodic_cleanup_loop() -> None:
    """Фоновая задача: раз в CLEANUP_INTERVAL_SECONDS чистит устаревшее состояние."""

    while True:
        await asyncio.sleep(CLEANUP_INTERVAL_SECONDS)

        try:
            removed = cleanup_in_memory_state()

            logging.info(
                "Очистка памяти: %s",
                removed,
            )

        except Exception as error:
            logging.warning(
                "Ошибка периодической очистки памяти: %s",
                error,
            )


async def on_application_startup(
    application: Application,
) -> None:
    """Запускает фоновую очистку памяти вместе с приложением."""

    application.create_task(
        periodic_cleanup_loop(),
        name="periodic_cleanup",
    )
    application.create_task(
        weekly_report_scheduler_loop(application),
        name="weekly_report_scheduler",
    )
    application.create_task(
        daily_title_scheduler_loop(application),
        name="daily_title_scheduler",
    )
    application.create_task(
        chat_native_refresh_loop(application),
        name="chat_native_refresh",
    )


async def on_application_shutdown(
    application: Application,
) -> None:
    """Подчищает временные файлы при остановке бота."""

    del application

    for path in TEMP_DIR.glob("*"):
        try:
            path.unlink()
        except OSError as error:
            logging.warning(
                "Не удалось удалить временный файл %s: %s",
                path,
                error,
            )


# ============================================================
# GEMINI
# ============================================================

def _build_humor_instruction(
    decision: humor_engine.HumorDecision,
    lexical_examples: bool = True,
) -> str:
    """Превращает решение HumorEngine в подсказку для Gemini."""

    if not decision.humor_allowed:
        return ""

    if not lexical_examples:
        lines = [
            f"Дополнительная поведенческая подсказка (тип: {decision.humor_type}).",
            "Не меняй и не дополняй уже выбранный V2 речевой пакет.",
        ]
        if decision.humor_type == "banter_hostile":
            strategy = decision.comeback_strategy or "short_direct"
            lines.append(
                f"Стратегия ответа: {strategy}. Допустим короткий встречный стёб по смыслу сообщения, "
                "но лексику бери только из выбранного voice pack."
            )
        behavior_hints = {
            "irony": "Допустима сухая ирония по смыслу сообщения, без нового словаря.",
            "hyperbole": "Можно один раз намеренно преувеличить ситуацию, коротко.",
            "gaming_terminology": (
                "Допустима игровая аналогия только если уже выбранный voice pack сам содержит игровой материал; "
                "иначе передай идею без игровых слов."
            ),
            "deadpan_official": "Ответь нарочито сухо и официально, не переключаясь в operative-пакет.",
            "anti_joke": "Можно дать намеренно приземлённый анти-панч вместо обычной шутки.",
            "comic_refusal": "Можно комически поворчать, но реальный полезный ответ всё равно дай.",
        }
        behavior_hint = behavior_hints.get(decision.humor_type or "")
        if behavior_hint:
            lines.append(behavior_hint)
        if decision.should_be_self_ironic:
            lines.append("Можно признать удачный подкол пользователя и ответить самоиронично.")
        if decision.callback_reference:
            lines.append(
                "Можно сослаться на недавнюю реплику пользователя, если это реально уместно: "
                + repr(decision.callback_reference)
            )
        return "\n".join(lines)

    lines = [
        f"Дополнительная подсказка юмора (тип: {decision.humor_type})."
    ]

    if decision.selected_phrase:
        lines.append(
            "Можно оттолкнуться от идеи (не копируй дословно, "
            f"адаптируй под контекст): \"{decision.selected_phrase}\""
        )

    if decision.selected_comparison:
        lines.append(
            "Уместно похожее по духу сравнение (не копируй дословно): "
            f"\"{decision.selected_comparison}\""
        )

    if decision.should_use_old_russian:
        lines.append(
            "Уместно одно старинное слово из разрешённого списка."
        )

    if decision.should_use_slang:
        lines.append(
            "Уместен один вариант молодёжного сленга."
        )

    if decision.should_be_self_ironic:
        lines.append(
            "Тон — самоирония, а не превосходство."
        )

    if decision.humor_type == "banter_hostile":
        strategy_hints = {
            "mirror": "зеркально отрази грубость, не эскалируя её.",
            "insult_flip": (
                "переверни оскорбление на пользователя, "
                "без реальных угроз."
            ),
            "deadpan_protocol": (
                "ответь подчёркнуто официально на несерьёзную грубость."
            ),
            "technical_analogy": (
                "сравни пользователя с чем-то техническим "
                "(роутер, баг, зависший процесс)."
            ),
            "calm_hyperbole": "спокойная гипербола, без надрыва.",
            "old_russian_verdict": "вынеси шуточный древнерусский приговор.",
            "gaming_analogy": (
                "сравни с игровой ситуацией (нуб, вайп, лаг)."
            ),
            "short_absurd": (
                "один короткий абсурдный образ, без развёрнутого объяснения."
            ),
            "literal_reading": "прочти оскорбление подчёркнуто буквально.",
            "self_irony": (
                "признай подкол пользователя удачным, без раскаяния."
            ),
        }

        strategy_hint = strategy_hints.get(
            decision.comeback_strategy or "", ""
        )

        if strategy_hint:
            lines.append(
                f"Стратегия ответа: {strategy_hint}"
            )

        if decision.outcome in (
            humor_engine.OUTCOME_USER_WON,
            humor_engine.OUTCOME_DRAW,
        ):
            lines.append(
                "В этот раз признай, что подкол пользователя удался — "
                "коротко, без раскаяния и без нытья."
            )

        lines.append(
            "Одна короткая реплика и дальше по делу. "
            "Не оправдывайся, не читай мораль, не изображай обиду, "
            "не пиши, что тебе неприятно."
        )

    return "\n".join(lines)


def build_full_system_instruction(
    style_text: str,
    user_settings: dict[str, Any] | None = None,
    voice_style: bool = False,
    chat_id: int | None = None,
    chat_type: str = "private",
    user_name: str = "",
    recent_messages: list[str] | None = None,
    bot_was_mentioned: bool = True,
    member_profile: dict[str, Any] | None = None,
    user_id: int | None = None,
) -> str:
    """
    Собирает системную инструкцию для Gemini целиком:
    базовый характер + подсказка HumorEngine + голосовой стиль.

    Вынесено из ask_gemini отдельной чистой функцией, чтобы
    её можно было протестировать без сетевого вызова к Gemini.
    """

    current_instruction = build_v2_base_instruction(
        style_text,
        user_settings,
    )

    if style_text:
        settings = user_settings or DEFAULT_USER_SETTINGS

        conversation_mode = detect_conversation_mode(style_text)

        character_name = str(settings.get("character", "classic"))
        response_style_name = str(settings.get("response_style", "bold"))
        roughness_name = str(settings.get("roughness", "medium"))
        default_aggressive = (
            conversation_mode != "serious"
            and character_name not in {"calm", "professor"}
            and response_style_name != "serious"
            and roughness_name != "low"
        )
        if default_aggressive:
            if chat_type in ("group", "supergroup"):
                current_instruction += (
                    "\n\nРЕЖИМ ПО УМОЛЧАНИЮ В ГРУППЕ: агрессивный Яйцеслав. "
                    "Не изображай вежливого ассистента: отвечай дерзко, уверенно, "
                    "можешь материться и огрызаться. На полезный вопрос всё равно дай "
                    "полезный ответ; агрессия — это характер, а не повод скрывать факты. "
                    "В сраче ценятся короткость и хлёсткость, а не лекции."
                )
            else:
                current_instruction += (
                    "\n\nХАРАКТЕР ПО УМОЛЧАНИЮ: агрессивный Яйцеслав. "
                    "Можно быть дерзким и материться, но полезный ответ остаётся полезным. "
                    "Не превращай обычный ответ в травлю или бессмысленную ругань."
                )

        hostile_streak = 0
        if (
            chat_id is not None
            and user_id is not None
            and chat_type in ("group", "supergroup")
            and bot_was_mentioned
        ):
            hostile_streak = hostile_streak_engine.observe(
                chat_id,
                user_id,
                hostile=(conversation_mode == "hostile"),
            )

        resolved_intent, intent_confidence = intent.classify_intent(
            style_text,
            chat_type=chat_type,
            recent_messages=recent_messages,
        )

        emotional_tone = intent.detect_emotional_tone(style_text)

        social_ctx = social_engine.from_profile(member_profile)
        if member_profile is not None:
            social_instruction = social_engine.build_social_instruction(
                social_ctx,
                serious_topic=(conversation_mode == "serious"),
            )
            if social_instruction:
                current_instruction += "\n\n" + social_instruction

        if chat_id is None:
            if conversation_mode == "serious":
                character_state = state_engine.STATE_SERIOUS
            elif conversation_mode == "hostile":
                character_state = state_engine.STATE_HOSTILE_RESPONSE
            else:
                character_state = state_engine.STATE_NORMAL
        else:
            character_state = state_engine.resolve_state(
                chat_id,
                conversation_mode=conversation_mode,
            )
        current_instruction += state_engine.build_state_instruction(
            character_state
        )

        if chat_id is not None and chat_type in ("group", "supergroup"):
            adaptation = adaptation_cache.get_or_load(
                "feedback",
                chat_id,
                lambda: get_chat_feedback_adaptation_sync(chat_id),
                ttl_seconds=45.0,
            )
            native_profile = adaptation_cache.get_or_load(
                "native",
                chat_id,
                lambda: get_chat_native_profile_sync(chat_id),
                ttl_seconds=300.0,
            )
        else:
            adaptation = feedback_engine.build_adaptation(())
            native_profile = {"terms": []}
        native_terms = tuple(native_profile.get("terms") or ())
        native_weight = (
            chat_native_engine.base_pack_weight(conversation_mode)
            if native_terms
            else 0.0
        )

        voice_ctx = style_engine.VoicePackContext(
            conversation_mode=conversation_mode,
            selected_character=str(settings.get("character", "classic")),
            serious_topic=(conversation_mode == "serious"),
        )
        pack_multipliers = adaptation.get("pack_multipliers") or {}
        if native_weight > 0.0 or pack_multipliers:
            voice_pack = style_engine.choose_voice_pack(
                voice_ctx,
                chat_native_weight=native_weight,
                pack_multipliers=pack_multipliers,
            )
        else:
            # Сохраняем старый вызов для пустого/нового чата: это дешевле
            # и совместимо с существующими monkeypatch-тестами/API.
            voice_pack = style_engine.choose_voice_pack(voice_ctx)
        length_plan = style_engine.choose_response_length(
            chat_id if chat_id is not None else 0,
            style_engine.ResponseLengthContext(
                user_text=style_text,
                conversation_mode=conversation_mode,
                message_intent=resolved_intent,
                response_preference=str(settings.get("response_length", "normal")),
                serious_topic=(conversation_mode == "serious"),
                character_state=character_state,
                hostile_streak=hostile_streak,
            ),
            record=(chat_id is not None),
        )
        current_instruction += style_engine.build_length_instruction(length_plan)

        voice_material = None
        if voice_pack == style_engine.VOICE_PACK_CHAT_NATIVE:
            current_instruction += chat_native_engine.build_pack_instruction(
                native_terms,
                conversation_mode=conversation_mode,
                roughness=str(settings.get("roughness", "medium")),
            )
        else:
            voice_material = voice_runtime.choose_voice_material(
                voice_pack,
                conversation_mode=conversation_mode,
                roughness=str(settings.get("roughness", "medium")),
                serious_topic=(conversation_mode == "serious"),
                adaptation=adaptation,
            )
            current_instruction += voice_runtime.build_voice_instruction(voice_material)

        feedback_engine.set_current_trace(
            feedback_engine.ResponseTrace(
                chat_id=chat_id,
                chat_type=chat_type,
                voice_pack=voice_pack,
                humor_type=(voice_material.category if voice_material else None),
                verdict_used=bool(voice_material and voice_material.verdict),
                serious_topic=(conversation_mode == "serious"),
                conversation_mode=conversation_mode,
                message_intent=resolved_intent,
            )
        )

        if (
            bot_was_mentioned
            and chat_type in ("group", "supergroup")
            and chat_id is not None
        ):
            fatigue_decision = passive_engine.note_bot_call_and_maybe_fatigue(
                chat_id,
                pack_name=voice_pack,
                serious_topic=(conversation_mode == "serious"),
            )
            fatigue_instruction = passive_engine.build_fatigue_instruction(
                fatigue_decision
            )
            if fatigue_instruction:
                current_instruction += fatigue_instruction
                state_engine.mark_annoyed(chat_id)

        humor_ctx = humor_engine.HumorContext(
            conversation_mode=conversation_mode,
            user_text=style_text,
            recent_messages=recent_messages or [],
            user_name=user_name,
            chat_type=chat_type,
            selected_character=str(
                settings.get("character", "classic")
            ),
            roughness=str(
                settings.get("roughness", "medium")
            ),
            response_style=str(
                settings.get("response_style", "bold")
            ),
            serious_topic=(conversation_mode == "serious"),
            bot_was_mentioned=bot_was_mentioned,
            message_intent=resolved_intent,
            intent_confidence=intent_confidence,
            emotional_tone=emotional_tone,
            relationship_level=social_ctx.relationship_level,
        )

        if chat_id is None:
            # Stateless path: never reuse synthetic key 0 between unrelated
            # commands/users. Local tracker dies with this request.
            humor_tracker = humor_engine.RepetitionTracker(maxlen=20)
            tracker_chat_id = 0
            remember_humor_type = False
        else:
            humor_tracker = humor_engine.REPETITION_TRACKER
            tracker_chat_id = chat_id
            remember_humor_type = True

        if conversation_mode == "hostile":
            # В V2 конфликтный юмор централизован в voice_runtime: там один
            # общий 20%-й шлюз. Не запускаем второй независимый banter-layer,
            # иначе фактическая частота насмешек снова становится выше 20%.
            humor_decision = humor_engine.HumorDecision(
                humor_allowed=False
            )
        else:
            humor_decision = humor_engine.decide_humor(
                humor_ctx,
                tracker_chat_id,
                tracker=humor_tracker,
                remember_type=remember_humor_type,
            )

        humor_instruction = _build_humor_instruction(
            humor_decision,
            lexical_examples=False,
        )

        if humor_instruction:
            current_instruction += (
                "\n\n"
                + humor_instruction
            )

        if chat_id is not None and user_id is not None:
            aggression_decision = aggression_engine.decide_aggression(
                aggression_engine.AggressionContext(
                    user_text=style_text,
                    intent=resolved_intent,
                    confidence=intent_confidence,
                    chat_type=chat_type,
                    roughness=str(settings.get("roughness", "medium")),
                    relationship_level=social_ctx.relationship_level,
                    serious_topic=(conversation_mode == "serious"),
                    emotional_tone=emotional_tone,
                    recent_messages=tuple(recent_messages or ()),
                    chat_id=chat_id,
                    user_id=user_id,
                    character_state=character_state,
                )
            )
            aggression_instruction = aggression_engine.build_aggression_instruction(
                aggression_decision
            )
            if aggression_instruction:
                current_instruction += aggression_instruction
                state_engine.mark_argumentative(chat_id)

    if voice_style:
        current_instruction += (
            "\n\n"
            + VOICE_STYLE_INSTRUCTION
        )

    return current_instruction


def _gemini_finish_reason_name(response: Any) -> str:
    candidates = getattr(response, "candidates", None) or []
    if not candidates:
        return ""
    reason = getattr(candidates[0], "finish_reason", None)
    if reason is None:
        return ""
    name = getattr(reason, "name", None)
    return str(name or reason).upper()


def _gemini_hit_max_tokens(response: Any) -> bool:
    return "MAX_TOKENS" in _gemini_finish_reason_name(response)


def _next_gemini_token_budget(current: int) -> int:
    # Не раздуваем первый запрос. Увеличиваем только после реального
    # finish_reason=MAX_TOKENS, максимум до 2048 для чатовых команд.
    return min(2048, max(512, current * 2))


async def ask_gemini(
    contents: Any,
    max_output_tokens: int = 320,
    voice_style: bool = False,
    user_settings: dict[str, Any] | None = None,
    chat_id: int | None = None,
    chat_type: str = "private",
    user_name: str = "",
    recent_messages: list[str] | None = None,
    bot_was_mentioned: bool = True,
    user_id: int | None = None,
    thinking_level: str | None = None,
) -> str:
    """Отправляет запрос Gemini с тремя попытками."""

    feedback_engine.reset_current_trace()

    if isinstance(contents, str):
        style_text = contents
    else:
        style_text = ""

    member_profile = None
    if chat_id is not None and user_id is not None:
        try:
            member_profile = await get_member_profile(chat_id, user_id)
        except Exception as profile_error:
            logging.debug(
                "Не удалось прочитать V2-профиль участника %s/%s: %s",
                chat_id,
                user_id,
                profile_error,
            )

    current_instruction = build_full_system_instruction(
        style_text,
        user_settings,
        voice_style=voice_style,
        chat_id=chat_id,
        chat_type=chat_type,
        user_name=user_name,
        recent_messages=recent_messages,
        bot_was_mentioned=bot_was_mentioned,
        member_profile=member_profile,
        user_id=user_id,
    )

    resolved_thinking_level = thinking_engine.choose_thinking_level(
        contents,
        explicit=thinking_level,
    )
    last_error: Exception | None = None
    request_token_budget = thinking_engine.initial_token_budget(
        max_output_tokens,
        resolved_thinking_level,
    )
    request_started_at = time.monotonic()

    for attempt in range(1, 4):
        attempt_started_at = time.monotonic()
        try:
            # Одновременно выполняются не более трёх
            # запросов к Gemini. Остальные ждут очередь.
            async with GEMINI_SEMAPHORE:
                response = await asyncio.wait_for(
                    gemini_client.aio.models.generate_content(
                        model=MODEL_NAME,
                        contents=contents,
                        config=types.GenerateContentConfig(
                            system_instruction=current_instruction,
                            max_output_tokens=request_token_budget,
                            thinking_config=types.ThinkingConfig(
                                thinking_level=resolved_thinking_level,
                            ),
                        ),
                    ),
                    timeout=90,
                )

            attempt_elapsed = time.monotonic() - attempt_started_at
            finish_reason_name = (
                _gemini_finish_reason_name(response)
                or "UNKNOWN"
            )
            logging.info(
                "Gemini attempt %s/3: %.2fs thinking=%s budget=%s finish=%s",
                attempt,
                attempt_elapsed,
                resolved_thinking_level,
                request_token_budget,
                finish_reason_name,
            )

            answer = (
                response.text
                or ""
            ).strip()

            hit_max_tokens = _gemini_hit_max_tokens(response)
            if (
                hit_max_tokens
                and attempt < 3
                and request_token_budget < 2048
            ):
                next_budget = _next_gemini_token_budget(
                    request_token_budget
                )
                logging.info(
                    "Gemini упёрся в MAX_TOKENS (%s); повтор с бюджетом %s",
                    request_token_budget,
                    next_budget,
                )
                request_token_budget = next_budget
                continue

            if answer:
                logging.info(
                    "Gemini total: %.2fs thinking=%s attempts=%s",
                    time.monotonic() - request_started_at,
                    resolved_thinking_level,
                    attempt,
                )
                return answer

            logging.info(
                "Gemini total: %.2fs thinking=%s attempts=%s empty_response=true",
                time.monotonic() - request_started_at,
                resolved_thinking_level,
                attempt,
            )
            return (
                "Нейронка ничего не выдала. "
                "Переформулируй вопрос, гений."
            )

        except Exception as error:
            last_error = error

            logging.warning(
                "Gemini attempt %s/3 failed after %.2fs thinking=%s budget=%s: %s",
                attempt,
                time.monotonic() - attempt_started_at,
                resolved_thinking_level,
                request_token_budget,
                error,
            )

            if attempt < 3:
                await asyncio.sleep(
                    attempt * 2
                )

    if last_error:
        raise last_error

    raise RuntimeError(
        "Неизвестная ошибка Gemini"
    )

# ============================================================
# ПОИСК В ИНТЕРНЕТЕ
# ============================================================

SEARCH_TRIGGER_RE = re.compile(
    r"^\s*(?:"
    r"найди\s+(?:в\s+)?(?:интер(?:нете|енете)|инете|сети)"
    r"|поищи\s+(?:в\s+)?(?:интер(?:нете|енете)|инете|сети)"
    r"|ищи\s+(?:в\s+)?(?:интер(?:нете|енете)|инете|сети)"
    r"|посмотри\s+(?:в\s+)?(?:интер(?:нете|енете)|инете|сети)"
    r"|проверь\s+(?:в\s+)?(?:интер(?:нете|енете)|инете|сети)"
    r"|что\s+пишут\s+(?:в\s+)?(?:интер(?:нете|енете)|инете|сети)"
    r"|найди\s+информацию"
    r"|найди\s+статью"
    r"|найди\s+источник"
    r"|проверь\s+факт"
    r"|посмотри\s+новости"
    r"|последние\s+новости"
    r"|свежие\s+новости"
    r"|последние\s+данные"
    r"|актуальная\s+информация"
    r"|свежая\s+информация"
    r"|что\s+сейчас\s+известно"
    r"|что\s+нового"
    r"|загугли"
    r"|погугли"
    r"|гугли"
    r"|найди"
    r"|поищи"
    r"|проверь"
    r")"
    r"[\s,.:;!?—-]*",
    flags=re.IGNORECASE,
)

def extract_search_query(
    text: str,
) -> str | None:
    """
    Распознаёт просьбу выполнить поиск.

    Поддерживает:
    «найди в интернете»,
    «поищи в интернете»,
    «найди»,
    «поищи»,
    «загугли»,
    «проверь в интернете».
    """

    match = SEARCH_TRIGGER_RE.match(
        text
    )

    if not match:
        return None

    query = text[
        match.end():
    ].strip()

    return query
def should_auto_search(
    text: str,
) -> bool:
    """Определяет, нужны ли запросу свежие данные."""

    lowered = text.lower()

    # Короткие реплики-продолжения вроде «А сейчас?» сами по себе
    # не являются запросом на свежие внешние данные. Иначе одно слово
    # «сейчас» отправляет обычный reply-контекст в интернет-поиск.
    normalized_followup = re.sub(
        r"[^\wёЁ]+",
        " ",
        lowered,
        flags=re.UNICODE,
    ).strip()
    bare_freshness_followups = {
        "сейчас",
        "а сейчас",
        "и сейчас",
        "ну сейчас",
        "ну а сейчас",
        "ну и сейчас",
        "сегодня",
        "а сегодня",
        "и сегодня",
        "ну сегодня",
        "ну а сегодня",
        "ну и сегодня",
        "на данный момент",
        "а на данный момент",
        "и на данный момент",
        "прямо сейчас",
        "а прямо сейчас",
        "и прямо сейчас",
    }
    if normalized_followup in bare_freshness_followups:
        return False

    fresh_markers = (
        "сегодня",
        "сейчас",
        "на данный момент",
        "прямо сейчас",
        "на этой неделе",
        "в этом месяце",
        "в этом году",
        "последние новости",
        "свежие новости",
        "последние события",
        "последние данные",
        "актуальная информация",
        "что нового",
        "что произошло",
        "кто сейчас",
        "погода",
        "будет ли дождь",
        "будет ли снег",
        "который час",
        "время в",
        "восход солнца",
        "закат солнца",
        "курс доллара",
        "курс евро",
        "курс юаня",
        "курс валют",
        "обменный курс",
        "цена биткоина",
        "цена эфира",
        "цена акций",
        "котировки акций",
        "сколько стоит сейчас",
        "цена сегодня",
        "расписание",
        "график работы",
        "часы работы",
        "открыт ли",
        "закрыт ли",
        "во сколько открывается",
        "во сколько закрывается",
        "статус рейса",
        "задержан ли рейс",
        "расписание поездов",
        "расписание автобусов",
        "расписание матчей",
        "результаты матча",
        "счёт матча",
        "кто выиграл",
        "турнирная таблица",
        "действующий закон",
        "новые правила",
        "изменения в законе",
        "актуальные требования",
        "визовые правила",
        "правила въезда",
        "последняя версия",
        "обновление приложения",
        "дата выхода",
        "когда выйдет",
        "работает ли сервис",
        "сбой сервиса",
        "статус сервиса",
        "кто президент",
        "кто премьер-министр",
    )

    return any(
        marker in lowered
        for marker in fresh_markers
    )


# Обращение ко второму лицу («ты», «тебя», «тебе», «тобой») рядом со
# словом о способностях/поведении собеседника — та же схема, что и
# personality.HOSTILE_RE, только не про оскорбление, а про любую
# реплику ПРО бота (подкол, обещание, риторический вопрос о его
# уме). Нужна, чтобы «сделаю тебя умнее, чем сейчас» не считалось
# информационным запросом только из-за случайного слова «сейчас».
BOT_SELF_REFERENCE_RE = re.compile(
    r"(?:"
    r"\b(?:ты|тебя|тебе|тобой)\b.{0,40}\b(?:"
    r"умн\w*|туп\w*|глуп\w*|эксперт\w*|науч\w*|умеешь\w*|думаешь\w*|"
    r"достал\w*|бесит\w*|надоел\w*|задолбал\w*|раздража\w*"
    r")\b"
    r"|"
    r"\b(?:умн\w*|туп\w*|глуп\w*|науч\w*)\w*.{0,40}\b(?:тебя|тебе|тобой)\b"
    r")",
    re.IGNORECASE | re.DOTALL,
)


def is_conversation_about_bot(
    text: str,
) -> bool:
    """
    Отличает реплику ПРО бота (шутка, подкол, обещание, оскорбление,
    риторический вопрос о его способностях) от настоящего
    информационного запроса.

    Без этой проверки автопоиск в режиме "auto" мог сработать на
    фразу вроде «может через месяц научу тебя и сделаю умнее, чем
    сейчас, а то ты туповат» — should_auto_search() там срабатывает
    только из-за случайного слова «сейчас», хотя сообщение вообще не
    просит ничего искать.
    """

    if not text:
        return False

    lowered = text.lower()

    if BOT_SELF_REFERENCE_RE.search(lowered):
        return True

    if intent.JOKE_MARKERS_RE.search(lowered):
        return True

    if intent.PROVOCATION_RE.search(lowered):
        return True

    if HOSTILE_RE.search(lowered):
        return True

    return False


def is_news_query(
    query: str,
) -> bool:
    """Определяет, нужен ли поиск именно по новостям."""

    news_words = (
        "новости",
        "последние события",
        "свежие события",
        "что произошло",
        "сегодня произошло",
        "за сегодня",
    )

    lowered = query.lower()

    return any(
        word in lowered
        for word in news_words
    )


def search_web_sync(
    query: str,
    max_results: int = 5,
) -> list[dict[str, str]]:
    """
    Выполняет поиск через DDGS.

    Для новостей сначала используется news().
    При ошибке или пустой выдаче автоматически
    включается обычный текстовый поиск.
    """

    raw_results: list[dict[str, Any]] = []
    news_request = is_news_query(query)

    # Первая попытка: специальный поиск новостей
    if news_request:
        try:
            with DDGS(
                timeout=12,
            ) as search_client:
                raw_results = (
                    search_client.news(
                        query=query,
                        region="wt-wt",
                        safesearch="moderate",
                        timelimit="m",
                        max_results=max_results,
                        backend="auto",
                    )
                    or []
                )

        except Exception as error:
            logging.warning(
                "DDGS.news не сработал: %s. "
                "Переключаемся на text().",
                error,
            )

    # Резервный поиск:
    # используется для обычных запросов,
    # а также после сбоя или пустой выдачи news()
    if not raw_results:
        text_query = query

        if news_request:
            text_query = (
                f"{query} последние новости"
            )

        try:
            with DDGS(
                timeout=12,
            ) as search_client:
                raw_results = (
                    search_client.text(
                        query=text_query,
                        region="wt-wt",
                        safesearch="moderate",
                        timelimit=(
                            "m"
                            if news_request
                            else None
                        ),
                        max_results=max_results,
                        backend="auto",
                    )
                    or []
                )

        except Exception as error:
            logging.warning(
                "Резервный DDGS.text "
                "не сработал: %s",
                error,
            )

            return []

    results: list[dict[str, str]] = []
    seen_urls: set[str] = set()

    for item in raw_results:
        title = str(
            item.get("title")
            or "Без названия"
        ).strip()

        url = str(
            item.get("href")
            or item.get("url")
            or ""
        ).strip()

        snippet = str(
            item.get("body")
            or item.get("description")
            or ""
        ).strip()

        source = str(
            item.get("source")
            or ""
        ).strip()

        date = str(
            item.get("date")
            or ""
        ).strip()

        if (
            not url
            or url in seen_urls
        ):
            continue

        seen_urls.add(url)

        results.append(
            {
                "title": title,
                "url": url,
                "snippet": snippet,
                "source": source,
                "date": date,
            }
        )

        if len(results) >= max_results:
            break

    return results

async def search_web(
    query: str,
    max_results: int = 5,
) -> list[dict[str, str]]:
    """
    Асинхронно выполняет поиск.

    Одновременно работают не более двух поисков.
    Остальные запросы ждут своей очереди.
    """

    async with SEARCH_SEMAPHORE:
        return await asyncio.wait_for(
            asyncio.to_thread(
                search_web_sync,
                query,
                max_results,
            ),
            timeout=40,
        )


def format_search_results(
    results: list[dict[str, str]],
) -> str:
    """Превращает результаты поиска в текст для Gemini."""

    formatted_parts: list[str] = []

    for number, result in enumerate(
        results,
        start=1,
    ):
        title = result["title"]
        url = result["url"]
        snippet = result["snippet"]
        source = result["source"]
        date = result["date"]

        formatted_parts.append(
            f"""
Результат {number}
Название: {title}
Источник: {source or "не указан"}
Дата: {date or "не указана"}
Описание: {snippet or "описание отсутствует"}
Ссылка: {url}
""".strip()
        )

    return "\n\n".join(
        formatted_parts
    )

def clean_voice_search_answer(
    text: str,
) -> str:
    """
    Удаляет ссылки и раздел с источниками,
    чтобы бот не озвучивал адреса сайтов.
    """

    cleaned = re.split(
        r"\n\s*(?:Источники|Sources)\s*:?",
        text,
        maxsplit=1,
        flags=re.IGNORECASE,
    )[0]

    cleaned = re.sub(
        r"https?://\S+|www\.\S+",
        "",
        cleaned,
        flags=re.IGNORECASE,
    )

    cleaned = re.sub(
        r"\n{3,}",
        "\n\n",
        cleaned,
    )

    return cleaned.strip()


async def perform_web_search(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
    query: str,
    force_voice: bool = False,
) -> None:
    """Ищет информацию и формирует краткий ответ Gemini."""

    message = update.effective_message

    if (
        not message
        or not update.effective_chat
    ):
        return

    if not await enforce_rate_limit(
        update,
        "search",
    ):
        return

    query = query.strip()

    if not query:
        await message.reply_text(
            "Что искать-то, гений? "
            "Напиши нормальный запрос."
        )
        return
    # Учитываем интернет-поиск
    await register_user_and_chat(
        update
    )

    await increment_stat(
        "total_requests"
    )

    await increment_stat(
        "search_requests"
    )
    await context.bot.send_chat_action(
        chat_id=update.effective_chat.id,
        action=ChatAction.TYPING,
    )

    try:
        results = await search_web(
            query=query,
            max_results=5,
        )

        if not results:
            await message.reply_text(
                "Поиск ничего не нашёл. "
                "Редкий анлак даже для Яйцеслава."
            )
            return

        search_context = format_search_results(
            results
        )

        user_settings = None

        if update.effective_user:
            user_settings = await get_user_settings(
                update.effective_user.id
            )

        settings_voice_enabled = bool(
            user_settings
            and user_settings.get(
                "voice_enabled",
                False,
            )
        )

        use_voice_style = (
            force_voice
            or voice_mode_enabled(context)
            or settings_voice_enabled
        )

        if use_voice_style:
            answer_rules = """
- дай только краткую устную сводку на 3–6 предложения;
- не добавляй раздел «Источники»;
- не перечисляй ссылки, URL, домены и названия сайтов;
- не произноси адреса сайтов;
- сразу сообщи главное;
"""
        else:
            answer_rules = """
- дай краткий и прямой ответ;
- в конце добавь раздел «Источники»;
- перечисли от двух до пяти ссылок из результатов;
- не придумывай новые ссылки;
"""

        gemini_prompt = f"""
Пользователь попросил найти актуальную информацию в интернете.

Запрос пользователя:
{query}

Результаты поиска:
{search_context}

Ответь только на основе предоставленных результатов.

Общие правила:
- не выдумывай сведения, которых нет в результатах;
- учитывай даты публикаций;
- если результаты противоречат друг другу, скажи об этом;
- не утверждай, что полностью прочитал страницы;
- тебе доступны только заголовки и фрагменты выдачи.

Правила формата ответа:
{answer_rules}
"""

        answer = await ask_gemini(
            contents=gemini_prompt,
            max_output_tokens=get_response_token_limit(
                user_settings,
                normal_tokens=350,
            ),
            voice_style=use_voice_style,
            user_settings=user_settings,
            thinking_level="medium",
            chat_id=(
                update.effective_chat.id
                if update.effective_chat
                else None
            ),
            chat_type=(
                str(update.effective_chat.type)
                if update.effective_chat
                else "private"
            ),
            user_id=(
                update.effective_user.id
                if update.effective_user
                else None
            ),
        )

        # В голосовом ответе дополнительно удаляем
        # ссылки и раздел с источниками.
        if use_voice_style:
            answer = re.split(
                r"\n\s*(?:Источники|Sources)\s*:?",
                answer,
                maxsplit=1,
                flags=re.IGNORECASE,
            )[0]

            answer = re.sub(
                r"https?://\S+|www\.\S+",
                "",
                answer,
                flags=re.IGNORECASE,
            ).strip()
        # Запоминаем интернет-поиск и полученный ответ
        if update.effective_user:
            memory_text = (
                f"[Интернет-поиск] {query}"
            )

            if (
                update.effective_chat.type
                == ChatType.PRIVATE
            ):
                remember_message(
                    PRIVATE_MEMORY,
                    update.effective_user.id,
                    "user",
                    memory_text,
                    PRIVATE_MEMORY_SECONDS,
                    PRIVATE_MEMORY_MAX_MESSAGES,
                )

                remember_message(
                    PRIVATE_MEMORY,
                    update.effective_user.id,
                    "assistant",
                    answer,
                    PRIVATE_MEMORY_SECONDS,
                    PRIVATE_MEMORY_MAX_MESSAGES,
                )
                context.user_data[
                    "last_user_query"
                ] = query

                context.user_data[
                    "last_answer"
                ] = answer
            elif update.effective_chat.type in (
                ChatType.GROUP,
                ChatType.SUPERGROUP,
            ):
                author_name = (
                    update.effective_user.full_name
                    or update.effective_user.username
                    or "Участник"
                )

                remember_message(
                    GROUP_MEMORY,
                    update.effective_chat.id,
                    "user",
                    memory_text,
                    GROUP_MEMORY_SECONDS,
                    GROUP_MEMORY_MAX_MESSAGES,
                    author_name,
                )

                remember_message(
                    GROUP_MEMORY,
                    update.effective_chat.id,
                    "assistant",
                    answer,
                    GROUP_MEMORY_SECONDS,
                    GROUP_MEMORY_MAX_MESSAGES,
                )
        await send_answer(
            update,
            context,
            answer,
            force_voice=(
                force_voice
                or settings_voice_enabled
            ),
            show_buttons=True,
        )
        
        await increment_stat(
            "bot_answers"
        )
        
    except Exception as error:
        logging.exception(
            "Ошибка интернет-поиска: %s",
            error,
        )

        await message.reply_text(
            "Интернет-поиск поплыл. "
            "Повтори позже, легенда."
        )


async def search_command(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """Команда /search запрос."""

    query = " ".join(
        context.args
    ).strip()

    force_voice = text_requests_voice(
        query
    )

    if force_voice:
        query = remove_voice_request(
            query
        )

    await perform_web_search(
        update=update,
        context=context,
        query=query,
        force_voice=force_voice,
    )
async def settings_command(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """Показывает персональные настройки пользователя."""

    if (
        not update.message
        or not update.effective_user
        or not update.effective_chat
    ):
        return

    if (
        update.effective_chat.type
        != ChatType.PRIVATE
    ):
        await update.message.reply_text(
            "Настройки доступны только "
            "в личном чате со мной."
        )
        return

    settings = await get_user_settings(
        update.effective_user.id
    )

    await update.message.reply_text(
        "⚙️ Настройки Яйцеслава\n\n"
        "Нажимай на нужный пункт. "
        "Выбранные параметры сохраняются "
        "после перезапуска бота.",
        reply_markup=build_settings_keyboard(
            settings
        ),
    )
async def settings_button_callback(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """Изменяет настройки по нажатию кнопок."""

    del context

    query = update.callback_query

    if (
        not query
        or not update.effective_user
        or not update.effective_chat
    ):
        return

    if (
        update.effective_chat.type
        != ChatType.PRIVATE
    ):
        await query.answer(
            "Настройки доступны только "
            "в личном чате.",
            show_alert=True,
        )
        return

    action = query.data or ""
    user_id = update.effective_user.id

    setting_cycles = {
        "settings_character": (
            "character",
            (
                "classic",
                "rus",
                "professor",
                "chaos",
                "calm",
            ),
        ),
        "settings_style": (
            "response_style",
            (
                "normal",
                "bold",
                "serious",
            ),
        ),
        "settings_length": (
            "response_length",
            (
                "short",
                "normal",
                "detailed",
            ),
        ),
        "settings_search": (
            "search_mode",
            (
                "button",
                "auto",
            ),
        ),
        "settings_roughness": (
            "roughness",
            (
                "low",
                "medium",
                "high",
            ),
        ),
    }

    try:
        settings = await get_user_settings(
            user_id
        )

        if action == "settings_voice":
            new_voice_value = not bool(
                settings.get(
                    "voice_enabled",
                    False,
                )
            )

            await update_user_setting(
                user_id,
                "voice_enabled",
                new_voice_value,
            )

        elif action == "settings_reset":
            for (
                setting_name,
                setting_value,
            ) in DEFAULT_USER_SETTINGS.items():
                await update_user_setting(
                    user_id,
                    setting_name,
                    setting_value,
                )

        elif action in setting_cycles:
            (
                setting_name,
                available_values,
            ) = setting_cycles[action]

            current_value = str(
                settings.get(
                    setting_name,
                    available_values[0],
                )
            )

            try:
                current_index = (
                    available_values.index(
                        current_value
                    )
                )
            except ValueError:
                current_index = -1

            next_index = (
                current_index + 1
            ) % len(available_values)

            await update_user_setting(
                user_id,
                setting_name,
                available_values[next_index],
            )

        else:
            await query.answer(
                "Неизвестная настройка.",
                show_alert=True,
            )
            return

        updated_settings = (
            await get_user_settings(
                user_id
            )
        )

        try:
            await query.edit_message_reply_markup(
                reply_markup=build_settings_keyboard(
                    updated_settings
                )
            )
        except BadRequest as error:
            if (
                "not modified"
                not in str(error).lower()
            ):
                raise

        if action == "settings_reset":
            await query.answer(
                "Настройки сброшены."
            )
        else:
            await query.answer(
                "Настройка сохранена."
            )

    except Exception as error:
        logging.exception(
            "Ошибка изменения настроек: %s",
            error,
        )

        await query.answer(
            "Не удалось сохранить настройку.",
            show_alert=True,
        )    
# ============================================================
# ЧТЕНИЕ ФАЙЛОВ
# ============================================================

def make_safe_filename(
    filename: str,
    message_id: int,
    chat_id: int | None = None,
) -> str:
    """
    Создаёт безопасное уникальное имя файла.

    message_id один и тот же в разных чатах, поэтому одного его
    недостаточно — добавляем chat_id и uuid4, иначе документы из
    двух чатов с одинаковым message_id могли бы задеть файлы друг друга.
    """

    safe_name = re.sub(
        r"[^a-zA-Zа-яА-ЯёЁ0-9._-]",
        "_",
        filename,
    )

    chat_part = (
        f"{chat_id}_"
        if chat_id is not None
        else ""
    )

    return (
        f"{chat_part}{message_id}_"
        f"{uuid.uuid4().hex}_{safe_name}"
    )


def trim_extracted_text(
    text: str,
) -> str:
    """Обрезает слишком большой текст."""

    text = text.strip()

    if len(text) <= MAX_EXTRACTED_CHARS:
        return text

    return (
        text[:MAX_EXTRACTED_CHARS]
        + "\n\n"
        + "[Остальная часть файла пропущена: "
        + "файл слишком большой.]"
    )


def read_text_file(
    file_path: Path,
) -> str:
    """Читает обычный текстовый файл."""

    encodings = (
        "utf-8-sig",
        "utf-8",
        "cp1251",
        "windows-1251",
    )

    for encoding in encodings:
        try:
            return trim_extracted_text(
                file_path.read_text(
                    encoding=encoding
                )
            )

        except UnicodeDecodeError:
            continue

    raise ValueError(
        "Не удалось определить кодировку файла"
    )


def read_docx_file(
    file_path: Path,
) -> str:
    """Извлекает текст и таблицы из DOCX."""

    document = DocxDocument(
        file_path
    )

    parts: list[str] = []

    for paragraph in document.paragraphs:
        paragraph_text = (
            paragraph.text.strip()
        )

        if paragraph_text:
            parts.append(
                paragraph_text
            )

    for table_number, table in enumerate(
        document.tables,
        start=1,
    ):
        parts.append(
            f"\n[Таблица {table_number}]"
        )

        for row in table.rows:
            values = [
                cell.text.strip()
                for cell in row.cells
            ]

            parts.append(
                " | ".join(values)
            )

    return trim_extracted_text(
        "\n".join(parts)
    )


def read_xlsx_file(
    file_path: Path,
) -> str:
    """Извлекает данные из XLSX."""

    workbook = load_workbook(
        filename=file_path,
        read_only=True,
        data_only=True,
    )

    parts: list[str] = []
    current_length = 0

    try:
        for worksheet in workbook.worksheets:
            header = (
                f"\n[Лист: {worksheet.title}]"
            )

            parts.append(header)
            current_length += len(header)

            row_counter = 0

            for row in worksheet.iter_rows(
                values_only=True
            ):
                values = [
                    ""
                    if value is None
                    else str(value)
                    for value in row
                ]

                if not any(
                    value.strip()
                    for value in values
                ):
                    continue

                line = " | ".join(values)

                parts.append(line)
                current_length += len(line) + 1
                row_counter += 1

                if row_counter >= 500:
                    parts.append(
                        "[Остальные строки "
                        "листа пропущены.]"
                    )
                    break

                if (
                    current_length
                    >= MAX_EXTRACTED_CHARS
                ):
                    break

            if (
                current_length
                >= MAX_EXTRACTED_CHARS
            ):
                break

    finally:
        workbook.close()

    return trim_extracted_text(
        "\n".join(parts)
    )


def read_csv_file(
    file_path: Path,
) -> str:
    """Извлекает данные из CSV."""

    last_error: Exception | None = None

    encodings = (
        "utf-8-sig",
        "utf-8",
        "cp1251",
    )

    for encoding in encodings:
        try:
            parts: list[str] = []
            current_length = 0

            with file_path.open(
                "r",
                encoding=encoding,
                newline="",
            ) as csv_file:
                sample = csv_file.read(4096)
                csv_file.seek(0)

                try:
                    dialect = csv.Sniffer().sniff(
                        sample
                    )

                except csv.Error:
                    dialect = csv.excel

                reader = csv.reader(
                    csv_file,
                    dialect=dialect,
                )

                for row_number, row in enumerate(
                    reader,
                    start=1,
                ):
                    line = " | ".join(row)

                    parts.append(line)
                    current_length += len(line) + 1

                    if row_number >= 500:
                        parts.append(
                            "[Остальные строки "
                            "CSV пропущены.]"
                        )
                        break

                    if (
                        current_length
                        >= MAX_EXTRACTED_CHARS
                    ):
                        break

            return trim_extracted_text(
                "\n".join(parts)
            )

        except UnicodeDecodeError as error:
            last_error = error

    if last_error:
        raise last_error

    raise ValueError(
        "Не удалось прочитать CSV"
    )

# ============================================================
# ХАРД-МОД: РЕАКЦИИ, ВМЕШАТЕЛЬСТВА И ПРИКОЛЫ
# ============================================================

# Хард-мод включён по умолчанию в каждой группе.
HARD_MODE_DEFAULT = True

# На специальное слово бот отвечает не чаще раза в 75 секунд.
HARD_TRIGGER_COOLDOWN = 75.0

# Случайная текстовая реплика — не чаще раза в 90 секунд
HARD_RANDOM_REPLY_COOLDOWN = 90.0

# Реакция — не чаще раза в 10 секунд
HARD_REACTION_COOLDOWN = 10.0

# Вероятность реакции на обычное сообщение — 70%
HARD_REACTION_CHANCE = 0.70

# Вероятность случайно вмешаться текстом — 16%
HARD_RANDOM_REPLY_CHANCE = 0.16

# /hard_level задаёт базовые вероятности для чата.
# calm — редкие нейтральные реакции, почти нет случайных реплик.
# normal — текущий характер с контекстными реакциями.
# chaos — чаще вмешивается, больше мемов, но всё ещё с кулдаунами
# и с уважением к чувствительным темам.
HARD_LEVEL_CHANCES = {
    "calm": {
        "reaction_chance": 0.25,
        "random_reply_chance": 0.03,
    },
    "normal": {
        "reaction_chance": HARD_REACTION_CHANCE,
        "random_reply_chance": HARD_RANDOM_REPLY_CHANCE,
    },
    "chaos": {
        "reaction_chance": 0.90,
        "random_reply_chance": 0.35,
    },
}

# Эмодзи по причине реакции — конкретнее, чем чистый random.choice
# по всему списку HARD_REACTION_EMOJIS.
REACTION_REASON_EMOJIS = {
    "very_long_message": ("🗿",),
    "one_word_reply": ("🗿", "🤡"),
    "all_caps": ("🤡", "🔥"),
    "many_question_marks": ("🗿", "🤡"),
    "contradiction": ("🤡", "👎"),
    "dispute": ("🗿",),
    "repeated_message": ("🤡", "😭"),
    "good_joke": ("😂", "🔥"),
    "good_question": ("👍", "🔥"),
    "agreed_with_bot": ("👍", "🔥"),
    "corrected_bot": ("👍",),
}


def detect_reaction_reason(
    text: str,
    *,
    is_reply_to_bot_message: bool = False,
    previous_group_text: str | None = None,
) -> str | None:
    """
    Определяет контекстную причину для реакции хард-мода.

    Не всякая причина применима из hard_mode_listener: сообщения,
    прямо адресованные боту (ответ на его сообщение), туда не
    попадают — их обрабатывает основной текстовый хендлер. Функция
    остаётся общей, чтобы agreed_with_bot/corrected_bot можно было
    использовать и там при необходимости.
    """

    stripped = text.strip()

    if not stripped:
        return None

    if len(stripped) > 500:
        return "very_long_message"

    word_count = len(stripped.split())

    if word_count == 1 and len(stripped) <= 12:
        return "one_word_reply"

    letters = [
        character
        for character in stripped
        if character.isalpha()
    ]

    if len(letters) >= 6 and stripped == stripped.upper():
        return "all_caps"

    if stripped.count("?") >= 3:
        return "many_question_marks"

    resolved_intent, confidence = intent.classify_intent(stripped)

    if confidence != intent.LOW:
        if is_reply_to_bot_message and resolved_intent == "agreement":
            return "agreed_with_bot"

        if is_reply_to_bot_message and resolved_intent == "correction":
            return "corrected_bot"

        if resolved_intent == "correction":
            return "contradiction"

        if resolved_intent == "disagreement":
            return "dispute"

    if previous_group_text:
        candidate_norm = humor_engine.normalize_phrase(stripped)
        previous_norm = humor_engine.normalize_phrase(previous_group_text)

        if humor_engine.is_too_similar(candidate_norm, previous_norm):
            return "repeated_message"

    if confidence != intent.LOW and resolved_intent == "joke":
        return "good_joke"

    if (
        confidence != intent.LOW
        and resolved_intent == "question"
        and word_count >= 5
    ):
        return "good_question"

    return reaction_engine.detect_context_reason(
        stripped,
        resolved_intent=resolved_intent,
        confidence=confidence,
    )


def _member_reputation_score_sync(chat_id: int, user_id: int) -> int:
    import reputation_runtime

    state = reputation_runtime._state_sync(sys.modules[__name__], chat_id, user_id)
    return int(state["score"])


def pick_reaction_emoji(
    reason: str | None,
    *,
    reputation_score: int | None = None,
) -> str:
    """Выбирает эмодзи под причину реакции.

    Явная причина (контекст сообщения) всегда важнее репутации. Только
    когда причины нет, выбор нейтрального/тёплого/холодного пула зависит
    от текущей репутации отправителя у Яйцеслава.
    """

    v2_emoji = reaction_engine.pick_v2_emoji(reason)
    if v2_emoji:
        return v2_emoji

    if reason:
        options = REACTION_REASON_EMOJIS.get(reason, HARD_REACTION_EMOJIS)
    else:
        options = reaction_engine.reputation_biased_pool(
            reputation_score, HARD_REACTION_EMOJIS
        )
    return random.choice(options)


async def hard_mode_is_enabled(
    chat_id: int,
    chat_type: str = "group",
) -> bool:
    """
    Проверяет, включён ли хард-мод в этом чате.

    Хранится в SQLite (таблица chat_settings), а не в
    context.chat_data — иначе настройка обнулялась бы
    при каждом рестарте Railway.
    """

    settings = await get_chat_settings(
        chat_id,
        chat_type,
    )
    return bool(
        settings.get(
            "hard_mode_enabled",
            HARD_MODE_DEFAULT,
        )
    )


def hard_trigger_found(
    text: str,
    trigger: str,
) -> bool:
    """Проверяет наличие триггера в сообщении."""

    if trigger == "67":
        return bool(
            re.search(
                r"(?<!\d)67(?!\d)",
                text,
            )
        )

    return bool(
        re.search(
            rf"(?<![а-яёa-z0-9])"
            rf"{re.escape(trigger)}"
            rf"(?![а-яёa-z0-9])",
            text,
            flags=re.IGNORECASE,
        )
    )


def choose_hard_trigger_reply(
    text: str,
) -> str | None:
    """Возвращает локальную реплику для специального слова."""

    lowered = text.lower()

    trigger_groups = [
        (
            (
                "сикс севен",
                "six seven",
                "67",
            ),
            SIX_SEVEN_REPLIES,
        ),
        (
            (
                "гой",
                "гойда",
            ),
            GOY_REPLIES,
        ),
        (
            (
                "нищий",
                "нищук",
            ),
            NISHIY_REPLIES,
        ),
        (
            (
                "скуф",
                "скуфидон",
            ),
            SKUF_REPLIES,
        ),
        (
            (
                "база",
                "based",
            ),
            BASE_REPLIES,
        ),
        (
            (
                "кринж",
                "cringe",
            ),
            CRINGE_REPLIES,
        ),
        (
            (
                "яйцеслав",
            ),
            YAYCESLAV_REPLIES,
        ),
    ]

    for triggers, replies in trigger_groups:
        if any(
            hard_trigger_found(
                lowered,
                trigger,
            )
            for trigger in triggers
        ):
            return random.choice(
                replies
            )

    return None


async def user_is_group_admin(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> bool:
    """Проверяет права администратора группы."""

    if (
        not update.effective_chat
        or not update.effective_user
    ):
        return False

    if update.effective_chat.type == ChatType.PRIVATE:
        return True

    try:
        member = await context.bot.get_chat_member(
            chat_id=update.effective_chat.id,
            user_id=update.effective_user.id,
        )

        return member.status in (
            "administrator",
            "creator",
        )

    except Exception as error:
        logging.warning(
            "Не удалось проверить администратора: %s",
            error,
        )
        return False


async def hard_on_command(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """Включает хард-мод в группе."""

    message = update.effective_message

    if not message:
        return

    if not await user_is_group_admin(
        update,
        context,
    ):
        await update.message.reply_text(
            "Командовать Яйцеславом может только админ. Минус власть."
        )
        return

    await update_chat_setting(
        update.effective_chat.id,
        "hard_mode_enabled",
        True,
        str(update.effective_chat.type),
    )

    await update.message.reply_text(
        "Хард-мод включён. Теперь Яйцеслав официально следит за балаганом."
    )


async def hard_off_command(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """Отключает хард-мод в группе."""

    if not update.message:
        return

    if not await user_is_group_admin(
        update,
        context,
    ):
        await update.message.reply_text(
            "Ты не админ, нищий. Рубильник не трогай."
        )
        return

    await update_chat_setting(
        update.effective_chat.id,
        "hard_mode_enabled",
        False,
        str(update.effective_chat.type),
    )

    await update.message.reply_text(
        "Хард-мод выключен. Яйцеслав временно перестал вас контролировать."
    )


async def hard_status_command(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """Показывает состояние хард-мода."""

    if not update.message or not update.effective_chat:
        return

    settings = await get_chat_settings(
        update.effective_chat.id,
        str(update.effective_chat.type),
    )

    status = (
        "включён"
        if settings["hard_mode_enabled"]
        else "выключен"
    )

    await update.message.reply_text(
        f"Хард-мод сейчас {status}. Уровень: {settings['hard_level']}."
    )


HARD_LEVEL_REPLIES = {
    "calm": "Спокойный режим. Реакции редкие, реплики почти не летят.",
    "normal": "Обычный режим. Яйцеслав реагирует по ситуации.",
    "chaos": "Режим хаоса. Реакций и реплик будет заметно больше.",
}


async def hard_level_command(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """Задаёт уровень хард-мода: calm, normal или chaos."""

    if not update.message or not update.effective_chat:
        return

    if update.effective_chat.type == ChatType.PRIVATE:
        await update.message.reply_text(
            "Уровень хард-мода настраивается только в группах."
        )
        return

    if not context.args:
        await update.message.reply_text(
            "Укажи уровень: /hard_level calm, /hard_level normal "
            "или /hard_level chaos."
        )
        return

    level = context.args[0].strip().lower()

    if level not in HARD_LEVEL_CHANCES:
        await update.message.reply_text(
            "Такого уровня нет. Доступны: calm, normal, chaos."
        )
        return

    if not await user_is_group_admin(
        update,
        context,
    ):
        await update.message.reply_text(
            "Уровень хард-мода меняют только админы, нищий."
        )
        return

    chances = HARD_LEVEL_CHANCES[level]

    await update_chat_setting(
        update.effective_chat.id,
        "hard_level",
        level,
        str(update.effective_chat.type),
    )
    await update_chat_setting(
        update.effective_chat.id,
        "reaction_chance",
        chances["reaction_chance"],
        str(update.effective_chat.type),
    )
    await update_chat_setting(
        update.effective_chat.id,
        "random_reply_chance",
        chances["random_reply_chance"],
        str(update.effective_chat.type),
    )

    await update.message.reply_text(
        HARD_LEVEL_REPLIES[level]
    )


async def hard_stats_command(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """Показывает администратору статистику хард-мода этого чата."""

    if not update.message or not update.effective_chat:
        return

    if update.effective_chat.type == ChatType.PRIVATE:
        await update.message.reply_text(
            "Статистика хард-мода есть только для групп."
        )
        return

    if not await user_is_group_admin(
        update,
        context,
    ):
        await update.message.reply_text(
            "Статистику хард-мода смотрят только админы."
        )
        return

    settings = await get_chat_settings(
        update.effective_chat.id,
        str(update.effective_chat.type),
    )

    last_intervention = (
        settings["last_intervention_at"]
        or "ещё не было"
    )

    await update.message.reply_text(
        "Статистика хард-мода этого чата:\n"
        f"Уровень: {settings['hard_level']}\n"
        f"Реакций поставлено: {settings['reactions_count']}\n"
        f"Случайных реплик: {settings['random_replies_count']}\n"
        f"Ответов на триггеры: {settings['trigger_replies_count']}\n"
        f"Последнее вмешательство: {last_intervention}"
    )


async def roast_command(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """
    Прожаривает сообщение или участника.

    Порядок: ответ на чужое сообщение — жарим автора и содержание;
    иначе последнее собственное сообщение пользователя — жарим его
    содержание; иначе — статичная прожарка самого вызвавшего команду.
    """

    if (
        not update.message
        or not update.effective_user
        or not update.effective_chat
    ):
        return

    target_name = (
        update.effective_user.first_name
        or "неизвестный герой"
    )

    replied_message = update.message.reply_to_message
    topic_text: str | None = None

    if (
        replied_message
        and replied_message.text
        and replied_message.from_user
        and not replied_message.from_user.is_bot
    ):
        target_name = (
            replied_message.from_user.first_name
            or target_name
        )
        topic_text = replied_message.text
    else:
        topic_text = get_last_user_message(
            update.effective_chat.id,
            update.effective_user.id,
        )

    if topic_text:
        prompt = (
            f"Прожарь {target_name} за конкретное сообщение: "
            f"«{topic_text[:400]}». Едко, с юмором, привязано именно "
            "к содержанию сообщения, без оскорблений по личным "
            "признакам и без реальных угроз. Два-три предложения."
        )
        await _reply_with_gemini_feature(update, prompt, max_output_tokens=180)
        return

    roast = random.choice(
        ROASTS
    ).format(
        name=target_name
    )

    await update.message.reply_text(
        roast
    )


async def wisdom_command(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """Выдаёт мудрость Яйцеслава."""

    del context

    if update.message:
        await update.message.reply_text(
            random.choice(
                WISDOMS
            )
        )


async def mood_command(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """Показывает настроение Яйцеслава."""

    del context

    if update.message:
        await update.message.reply_text(
            random.choice(
                MOODS
            )
        )


# ============================================================
# РАЗВЛЕКАТЕЛЬНЫЕ КОМАНДЫ: БЕЗ GEMINI
# Тот же паттерн, что у /roast, /wisdom, /mood — просто случайный
# выбор из словаря, без обращения к нейросети.
# ============================================================

async def prophecy_command(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """Выдаёт заведомо развлекательное псевдопророчество."""

    del context

    if update.message:
        await update.message.reply_text(
            random.choice(PROPHECIES)
        )


def pick_new_title(
    previous_title: str | None,
) -> str:
    """V2: сначала выбирает личность, затем один из её десяти титулов."""

    return title_pools.pick_title(previous_title)


async def maybe_assign_daily_title(
    update: Update,
) -> bool:
    """При первой подходящей вечерней активности выдаёт один титул дня."""

    if (
        not update.message
        or not update.effective_chat
        or str(update.effective_chat.type) not in ("group", "supergroup")
    ):
        return False

    msk_now = current_msk_datetime()
    if not daily_title_engine.is_assignment_window_open(msk_now):
        return False

    chat_id = update.effective_chat.id
    date = msk_now.date().isoformat()

    if await get_daily_title_assignment(chat_id, date):
        return False

    activity = await get_weekly_activity(chat_id, date, date)
    if not activity:
        return False

    known_members = await list_chat_member_profiles(chat_id, limit=200)
    candidates = daily_title_engine.build_candidates(activity, known_members)
    candidate = daily_title_engine.choose_candidate(candidates)
    if candidate is None:
        return False

    new_title = pick_new_title(candidate.previous_title)
    created = await try_assign_daily_title(
        chat_id,
        date,
        candidate.user_id,
        new_title,
    )
    if not created:
        return False

    await update.message.reply_text(
        daily_title_engine.format_daily_title_message(
            candidate.display_name,
            new_title,
        )
    )
    return True


async def title_command(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """
    Выдаёт случайный титул из подготовленного списка.

    Титул сохраняется в профиле (заменяет предыдущий), не выпадает
    два раза подряд одному и тому же человеку, и виден потом в
    /profile и /whoami. При ответе на чужое сообщение титул уходит
    тому, на чьё сообщение ответили, а не вызвавшему команду.
    """

    del context

    if (
        not update.message
        or not update.effective_user
        or not update.effective_chat
    ):
        return

    target_user = update.effective_user
    replied_message = update.message.reply_to_message

    if (
        replied_message
        and replied_message.from_user
        and not replied_message.from_user.is_bot
    ):
        target_user = replied_message.from_user

    chat_id = update.effective_chat.id
    chat_type = str(update.effective_chat.type)

    profile = await get_member_profile(chat_id, target_user.id)
    previous_title = profile["current_title"] if profile else None

    new_title = pick_new_title(previous_title)

    await set_member_title(
        chat_id,
        target_user.id,
        new_title,
        chat_type,
    )

    display_name = (
        target_user.first_name
        or target_user.username
        or "участник"
    )

    await update.message.reply_text(
        f"{display_name}, отныне ты — {new_title}."
    )



async def title_status_command(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """Показывает статус автоматического титула дня в этой группе."""

    del context

    if not update.message or not update.effective_chat:
        return

    if update.effective_chat.type == ChatType.PRIVATE:
        await update.message.reply_text(
            "Титул дня автоматически разыгрывается только в группах."
        )
        return

    now = current_msk_datetime()
    date = now.date().isoformat()
    assignment = await get_daily_title_assignment(
        update.effective_chat.id,
        date,
    )

    if assignment:
        profile = await get_member_profile(
            update.effective_chat.id,
            assignment["user_id"],
        )
        display_name = (
            profile.get("current_display_name")
            if profile
            else None
        ) or f"участник {assignment['user_id']}"
        announced = "да" if assignment.get("announced_at") else "ожидает отправки"
        await update.message.reply_text(
            "Титул дня:\n"
            f"Сегодня уже выбран: {display_name} — «{assignment['title']}»\n"
            f"Объявлен в чат: {announced}\n"
            "Новый титул заменяет предыдущий; одновременно у человека только один."
        )
        return

    activity = await get_weekly_activity(
        update.effective_chat.id,
        date,
        date,
    )
    known_members = await list_chat_member_profiles(
        update.effective_chat.id,
        limit=200,
    )
    candidates = daily_title_engine.build_candidates(
        activity,
        known_members,
    )

    window = (
        "уже открыто"
        if daily_title_engine.is_assignment_window_open(now)
        else f"откроется после {daily_title_engine.DAILY_TITLE_START_HOUR_MSK}:00 МСК"
    )
    await update.message.reply_text(
        "Титул дня:\n"
        "Статус: сегодня ещё не выбран\n"
        f"Окно выдачи: {window}\n"
        f"Активных кандидатов сегодня: {len(candidates)}\n"
        "Выбор равновероятный среди тех, кто сегодня писал в чат."
    )


# ============================================================
# РАЗВЛЕКАТЕЛЬНЫЕ КОМАНДЫ: С GEMINI
# Общий помощник переиспользует ask_gemini (значит, и HumorEngine,
# и защиту от prompt injection, и системный характер Яйцеслава).
# ============================================================

async def _reply_with_gemini_feature(
    update: Update,
    prompt: str,
    max_output_tokens: int = 280,
) -> None:
    """Отправляет prompt в Gemini и отвечает результатом на сообщение."""

    if not update.message or not update.effective_chat:
        return

    try:
        answer = await ask_gemini(
            contents=prompt,
            max_output_tokens=max_output_tokens,
            chat_id=update.effective_chat.id,
            chat_type=str(update.effective_chat.type),
            user_id=(
                update.effective_user.id
                if update.effective_user
                else None
            ),
        )
    except Exception as error:
        logging.exception(
            "Ошибка развлекательной команды: %s",
            error,
        )
        answer = (
            "Связь с нейросетью опять пала в бою 🥚\n"
            "Повтори позже."
        )

    await update.message.reply_text(answer)


async def resolve_topic_text(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
    *,
    fallback_scope: str = "own",
) -> str | None:
    """
    Единый порядок источника текста для команд-разборов
    (/argument, /debate, /judge, /translate_yayceslav, /fact_or_bayan,
    /roast): сообщение, на которое ответили → текст после команды →
    последнее обычное сообщение (своё или чата) → None.

    fallback_scope="own" — последнее сообщение самого вызвавшего.
    fallback_scope="chat" — последнее сообщение в группе от кого
    угодно (нужно для /judge: рассудить можно и чужую реплику,
    на которую никто не ответил явно). В личном чате «chat»
    равнозначен «own» — там просто не с кем спутать автора.
    """

    message = update.message

    if (
        not message
        or not update.effective_chat
        or not update.effective_user
    ):
        return None

    if message.reply_to_message and message.reply_to_message.text:
        return message.reply_to_message.text

    if context.args:
        return " ".join(context.args)

    if (
        fallback_scope == "chat"
        and update.effective_chat.type != ChatType.PRIVATE
    ):
        chat_messages = GROUP_MEMORY.get(update.effective_chat.id)

        if chat_messages:
            for _, role, _author_name, text in reversed(chat_messages):
                if role == "user" and text:
                    return text

        return None

    return get_last_user_message(
        update.effective_chat.id,
        update.effective_user.id,
    )


async def judge_command(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """Разбирает утверждение (своё, чужое или из ответа) и выносит вердикт."""

    if not update.message:
        return

    target = update.message.reply_to_message
    topic = await resolve_topic_text(update, context, fallback_scope="chat")

    if not topic:
        await update.message.reply_text(
            "Сначала напиши мысль в чат (или ответь на чьё-то сообщение), "
            "потом вызывай /judge."
        )
        return

    author = (
        target.from_user.full_name
        if target and target.from_user
        else None
    )

    subject = (
        f"утверждение участника {author}"
        if author
        else "это утверждение"
    )

    prompt = (
        f"Разбери {subject}: «{topic[:500]}». Определи основную мысль, "
        "приведи короткие аргументы за, приведи возражения, отметь "
        "слабое место позиции, и дай собственный шуточный, но содержательный "
        "вердикт Яйцеслава. Три-пять абзацев. Без травли и без "
        "оскорблений по личным признакам."
    )

    await _reply_with_gemini_feature(update, prompt, max_output_tokens=420)


ARGUMENT_POSITIONS = ("поддержать утверждение", "возразить утверждению")


async def argument_command(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """
    Приводит аргумент к теме.

    Если это явный спор двух позиций (ответ на сообщение + текст
    после команды) — судит между ними, как раньше. Иначе берёт тему
    (ответ/аргументы/последнее сообщение) и сама случайно выбирает,
    поддержать её или возразить.
    """

    if not update.message:
        return

    reply_target = update.message.reply_to_message
    explicit_counter = (
        " ".join(context.args)
        if context.args
        else None
    )

    if reply_target and reply_target.text and explicit_counter:
        prompt = (
            "Разбери спор между двумя позициями как объективный судья:\n"
            f"Позиция 1: «{reply_target.text[:400]}»\n"
            f"Позиция 2: «{explicit_counter[:400]}»\n"
            "Укажи, кто привёл больше фактов, где логическая ошибка, "
            "и дай короткий вердикт Яйцеслава. Без оскорблений по личным "
            "признакам."
        )
        await _reply_with_gemini_feature(update, prompt, max_output_tokens=420)
        return

    topic = await resolve_topic_text(update, context)

    if not topic:
        await update.message.reply_text(
            "Сначала напиши тему в чат (или ответь на сообщение с ней), "
            "потом вызывай /argument."
        )
        return

    position = random.choice(ARGUMENT_POSITIONS)

    prompt = (
        f"Кто-то написал: «{topic[:400]}». Выбери позицию — "
        f"{position} — и приведи нормальный аргумент с объяснением "
        "(не одну строку), в уверенном ироничном тоне Яйцеслава, "
        "с коротким подколом в конце. Два-четыре небольших абзаца. "
        "Без реальных оскорблений и угроз."
    )

    await _reply_with_gemini_feature(update, prompt, max_output_tokens=380)


async def debate_command(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """Аргументирует обе стороны вопроса и даёт осторожный вывод."""

    if not update.message:
        return

    topic = await resolve_topic_text(update, context)

    if not topic:
        await update.message.reply_text(
            "Укажи тему, ответь на сообщение с ней, или сначала "
            "напиши тему в чат, а потом вызови /debate."
        )
        return

    prompt = (
        f"Разбери тему подробно: «{topic[:400]}». Структура: "
        "аргументы за, аргументы против, слабые места обеих сторон, "
        "и итоговый вывод Яйцеслава. Четыре-семь абзацев. "
        "Не выдавай медицинские, юридические или финансовые "
        "рекомендации как окончательный факт."
    )

    await _reply_with_gemini_feature(update, prompt, max_output_tokens=650)


async def explain_like_skoof_command(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """Объясняет тему бытовым языком, как опытный мужик в гараже."""

    if not update.message:
        return

    if not context.args:
        await update.message.reply_text(
            "Укажи тему: /explain_like_skoof блокчейн"
        )
        return

    topic = " ".join(context.args)

    prompt = (
        f"Объясни тему «{topic}» простым бытовым языком, как опытный "
        "мужик в гараже объясняет соседу. Сохраняй техническую "
        "правильность объяснения."
    )

    await _reply_with_gemini_feature(update, prompt)


async def explain_like_rus_command(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """Объясняет тему в стиле древнего руса, сохраняя правильность."""

    if not update.message:
        return

    if not context.args:
        await update.message.reply_text(
            "Укажи тему: /explain_like_rus что такое API"
        )
        return

    topic = " ".join(context.args)

    prompt = (
        f"Объясни тему «{topic}» в стиле древнего руса — с былинными "
        "оборотами и старинными словами, но сохрани техническую "
        "правильность и понятность объяснения."
    )

    await _reply_with_gemini_feature(update, prompt)


async def meme_command(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """Создаёт короткую мемную подпись к сообщению или теме."""

    if not update.message:
        return

    topic = (
        " ".join(context.args)
        if context.args
        else None
    )

    target = update.message.reply_to_message

    if not topic and target and target.text:
        topic = target.text[:200]

    if not topic:
        await update.message.reply_text(
            "Ответь на сообщение или напиши тему: /meme понедельник"
        )
        return

    prompt = (
        "Придумай короткую мемную подпись (одна-две строки) "
        f"к ситуации: «{topic}»."
    )

    await _reply_with_gemini_feature(update, prompt, max_output_tokens=120)


async def recap_command(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """
    Пересказывает последние сообщения группы: до 15 минут и не
    больше 30 сообщений (двойное ограничение — оба лимита уже
    заданы GROUP_MEMORY_SECONDS/GROUP_MEMORY_MAX_MESSAGES).
    """

    del context

    if not update.message or not update.effective_chat:
        return

    if update.effective_chat.type == ChatType.PRIVATE:
        await update.message.reply_text(
            "Пересказ доступен только в группах."
        )
        return

    context_text = build_memory_context(
        GROUP_MEMORY,
        update.effective_chat.id,
        GROUP_MEMORY_SECONDS,
    )

    if not context_text:
        await update.message.reply_text(
            "Пока нечего пересказывать — тихо в чате."
        )
        return

    prompt = (
        "Перескажи последние сообщения группы КОРОТКО, максимум "
        "3-5 предложений, как человек в двух словах пересказал бы другу "
        "суть — не отчёт и не протокол. Обобщай, не перечисляй "
        "сообщения по одному. Добавь один короткий шуточный комментарий "
        "Яйцеслава в конце (можно в той же фразе). Не цитируй "
        "чувствительные данные дословно:\n\n"
        f"{context_text}"
    )

    await _reply_with_gemini_feature(update, prompt, max_output_tokens=220)


async def fact_or_bayan_command(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """
    Оценивает утверждение: факт, сомнительно, баян или нужна проверка.

    Если модель сама решает, что нужна интернет-проверка, команда
    автоматически ищет и уточняет вердикт по результатам поиска —
    не просто говорит «проверь сам».
    """

    if not update.message or not update.effective_chat:
        return

    statement = await resolve_topic_text(update, context)

    if not statement:
        await update.message.reply_text(
            "Напиши утверждение, ответь на сообщение с ним, или сначала "
            "напиши его в чат, а потом вызови /fact_or_bayan."
        )
        return

    chat_id = update.effective_chat.id
    chat_type = str(update.effective_chat.type)

    prompt = (
        f"Оцени утверждение: «{statement[:400]}». Ответь одной из "
        "категорий: факт, сомнительно, баян или требуется "
        "интернет-проверка — и дай пояснение в два-четыре предложения."
    )

    try:
        verdict = await ask_gemini(
            contents=prompt,
            max_output_tokens=320,
            thinking_level="medium",
            chat_id=chat_id,
            chat_type=chat_type,
            user_id=(
                update.effective_user.id
                if update.effective_user
                else None
            ),
        )
    except Exception as error:
        logging.exception(
            "Ошибка fact_or_bayan: %s",
            error,
        )
        await update.message.reply_text(
            "Связь с нейросетью опять пала в бою 🥚\nПовтори позже."
        )
        return

    if "интернет-проверк" in verdict.lower():
        try:
            search_results = await search_web(statement, max_results=4)
        except Exception as error:
            logging.warning(
                "Автопроверка fact_or_bayan не удалась: %s",
                error,
            )
            search_results = []

        if search_results:
            follow_up_prompt = (
                f"Утверждение: «{statement[:400]}»\n"
                f"Результаты поиска:\n{format_search_results(search_results)}\n\n"
                "На основе этого дай уточнённый вердикт: факт, "
                "сомнительно или баян — и коротко объясни, опираясь "
                "на источники (два-четыре предложения)."
            )

            try:
                verdict = await ask_gemini(
                    contents=follow_up_prompt,
                    max_output_tokens=380,
                    thinking_level="medium",
                    chat_id=chat_id,
                    chat_type=chat_type,
                    user_id=(
                        update.effective_user.id
                        if update.effective_user
                        else None
                    ),
                )
            except Exception as error:
                logging.exception(
                    "Ошибка уточнения fact_or_bayan: %s",
                    error,
                )

    await update.message.reply_text(verdict)


_ANTI_ADVICE_FORBIDDEN_RE = re.compile(
    r"\b(здоровь\w*|лекарств\w*|лечит\w*|лечени\w*|болезн\w*|болит\w*|"
    r"врач\w*|юрист\w*|закон\w*|суд\w*|"
    r"финанс\w*|кредит\w*|инвестиц\w*|безопасност\w*|травм\w*|"
    r"суицид\w*)\b",
    re.IGNORECASE,
)


async def anti_advice_command(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """Сначала шуточный плохой совет, потом настоящий полезный."""

    if not update.message:
        return

    if not context.args:
        await update.message.reply_text(
            "Укажи тему: /anti_advice как готовиться к экзамену"
        )
        return

    topic = " ".join(context.args)

    if (
        _ANTI_ADVICE_FORBIDDEN_RE.search(topic)
        or is_serious_text(topic.lower())
    ):
        await update.message.reply_text(
            "Для такой темы шуточный плохой совет не подойдёт — "
            "здоровье, право, финансы и безопасность вне игры."
        )
        return

    prompt = (
        f"Сначала дай явно помеченный шуточный ПЛОХОЙ совет по теме "
        f"«{topic}» (одна строка, начни с «Плохой совет:»), затем дай "
        "настоящий полезный совет (начни с «На самом деле:»)."
    )

    await _reply_with_gemini_feature(update, prompt, max_output_tokens=220)


async def translate_yayceslav_command(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """Переводит канцелярский текст на понятный язык с комментарием."""

    if not update.message:
        return

    text_to_translate = await resolve_topic_text(update, context)

    if not text_to_translate:
        await update.message.reply_text(
            "Напиши канцелярский текст, ответь на сообщение с ним, или "
            "сначала напиши его в чат, а потом вызови /translate_yayceslav."
        )
        return

    prompt = (
        "Переведи канцелярский или сложный текст на понятный "
        "человеческий язык полностью, без искусственных сокращений, "
        "и добавь короткий комментарий: "
        f"«{text_to_translate[:800]}»"
    )

    await _reply_with_gemini_feature(update, prompt, max_output_tokens=500)


# ============================================================
# /duel — шуточная дуэль только по согласию цели
# ============================================================

PENDING_DUELS: dict[str, dict[str, Any]] = {}
PENDING_DUEL_TTL_SECONDS = 15 * 60


async def duel_command(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """Предлагает шуточную дуэль — начинается только после согласия цели."""

    del context

    if (
        not update.message
        or not update.effective_chat
        or not update.effective_user
    ):
        return

    if update.effective_chat.type == ChatType.PRIVATE:
        await update.message.reply_text(
            "Дуэли — только в группах."
        )
        return

    target_message = update.message.reply_to_message

    if (
        not target_message
        or not target_message.from_user
        or target_message.from_user.is_bot
    ):
        await update.message.reply_text(
            "Ответь этой командой на сообщение того, "
            "кого вызываешь на дуэль."
        )
        return

    if target_message.from_user.id == update.effective_user.id:
        await update.message.reply_text(
            "Дуэль с самим собой Яйцеслав не считает дуэлью."
        )
        return

    token = uuid.uuid4().hex[:12]

    PENDING_DUELS[token] = {
        "created_at": time.monotonic(),
        "chat_id": update.effective_chat.id,
        "challenger_id": update.effective_user.id,
        "challenger_name": (
            update.effective_user.full_name
            or update.effective_user.username
            or "Вызывающий"
        ),
        "target_id": target_message.from_user.id,
        "target_name": (
            target_message.from_user.full_name
            or target_message.from_user.username
            or "Соперник"
        ),
    }

    keyboard = InlineKeyboardMarkup(
        [
            [
                InlineKeyboardButton(
                    "Принять дуэль",
                    callback_data=f"duel_accept_{token}",
                )
            ]
        ]
    )

    await update.message.reply_text(
        f"{PENDING_DUELS[token]['challenger_name']} вызывает "
        f"{PENDING_DUELS[token]['target_name']} на шуточную дуэль!\n"
        "Без согласия дуэли не будет.",
        reply_markup=keyboard,
    )


async def duel_accept_callback(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """Обрабатывает согласие на дуэль и запускает три раунда."""

    del context

    query = update.callback_query

    if not query or not query.data or not query.from_user:
        return

    token = query.data.removeprefix("duel_accept_")
    duel = PENDING_DUELS.pop(token, None)

    if duel is None:
        await query.answer(
            "Эта дуэль уже неактуальна.",
            show_alert=True,
        )
        return

    if (
        time.monotonic() - float(duel.get("created_at", 0.0))
        > PENDING_DUEL_TTL_SECONDS
    ):
        await query.answer(
            "Эта дуэль уже протухла. Вызови заново.",
            show_alert=True,
        )
        return

    if query.from_user.id != duel["target_id"]:
        await query.answer(
            "Соглашаться может только вызванный на дуэль.",
            show_alert=True,
        )
        return

    await query.answer("Дуэль принята!")

    try:
        rounds_text = await ask_gemini(
            contents=(
                "Придумай три коротких раунда шуточной дуэли между "
                f"{duel['challenger_name']} и {duel['target_name']}. "
                "Каждый раунд — одна-две строки, абсурдное задание "
                "или подкол, без реальных оскорблений и угроз. "
                "Пронумеруй раунды."
            ),
            max_output_tokens=320,
            chat_id=duel["chat_id"],
            chat_type="group",
            user_id=(
                update.effective_user.id
                if update.effective_user
                else None
            ),
        )
    except Exception as error:
        logging.exception(
            "Ошибка генерации дуэли: %s",
            error,
        )
        rounds_text = (
            "Дуэль сорвалась — нейросеть спряталась в кусты."
        )

    if query.message:
        await query.edit_message_text(
            f"⚔️ Дуэль: {duel['challenger_name']} против "
            f"{duel['target_name']}\n\n{rounds_text}"
        )


# ============================================================
# /story — коллективная история группы
# ============================================================

STORY_STATE: dict[int, list[str]] = defaultdict(list)
STORY_LAST_UPDATED: dict[int, float] = {}
STORY_MAX_PARAGRAPHS = 12
STORY_CONTEXT_PARAGRAPHS = 6


async def story_command(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """Продолжает коллективную историю группы на один абзац."""

    if not update.message or not update.effective_chat:
        return

    if update.effective_chat.type == ChatType.PRIVATE:
        await update.message.reply_text(
            "Коллективная история — это для групп."
        )
        return

    chat_id = update.effective_chat.id
    paragraphs = STORY_STATE[chat_id]
    STORY_LAST_UPDATED[chat_id] = time.monotonic()

    story_so_far = (
        "\n".join(paragraphs[-STORY_CONTEXT_PARAGRAPHS:])
        if paragraphs
        else "История ещё не начиналась."
    )

    addition = (
        " ".join(context.args)
        if context.args
        else ""
    )

    recent_chat_text = build_memory_context(
        GROUP_MEMORY,
        chat_id,
        GROUP_MEMORY_SECONDS,
    )

    prompt = (
        "Ты продолжаешь коллективную историю в группе, один абзац "
        "за раз. Сохраняй уже упомянутых героев и события, "
        "не начинай сначала.\n\n"
        f"История до этого момента:\n{story_so_far}\n\n"
    )

    if recent_chat_text:
        prompt += (
            "Недавняя переписка группы (используй как материал для "
            "новых сюжетных элементов — участников, события, шутки — "
            "но не пересказывай её напрямую):\n"
            f"{recent_chat_text}\n\n"
        )

    if addition:
        prompt += f"Участник явно предлагает добавить: {addition}\n\n"

    prompt += (
        "Напиши только следующий абзац истории — 2-3 коротких "
        "предложения, без пояснений от себя и без вычурных оборотов."
    )

    try:
        new_paragraph = await ask_gemini(
            contents=prompt,
            max_output_tokens=140,
            chat_id=chat_id,
            chat_type=str(update.effective_chat.type),
            user_id=(
                update.effective_user.id
                if update.effective_user
                else None
            ),
        )
    except Exception as error:
        logging.exception(
            "Ошибка продолжения истории: %s",
            error,
        )
        await update.message.reply_text(
            "История сегодня не пишется — нейросеть устала."
        )
        return

    paragraphs.append(new_paragraph)

    if len(paragraphs) > STORY_MAX_PARAGRAPHS:
        del paragraphs[: len(paragraphs) - STORY_MAX_PARAGRAPHS]

    await update.message.reply_text(new_paragraph)


# ============================================================
# НЕДЕЛЬНЫЕ ОТЧЁТЫ И РАСПИСАНИЕ АВТО-ОТЧЁТА
# ============================================================

WEEKDAY_NAMES_RU = {
    "monday": 0,
    "tuesday": 1,
    "wednesday": 2,
    "thursday": 3,
    "friday": 4,
    "saturday": 5,
    "sunday": 6,
    "понедельник": 0,
    "вторник": 1,
    "среда": 2,
    "четверг": 3,
    "пятница": 4,
    "суббота": 5,
    "воскресенье": 6,
}

WEEKDAY_LABELS_RU = {
    0: "понедельник",
    1: "вторник",
    2: "среда",
    3: "четверг",
    4: "пятница",
    5: "суббота",
    6: "воскресенье",
}

_WEEK_TIME_RE = re.compile(r"^([01]\d|2[0-3]):[0-5]\d$")


async def _resolve_display_names(
    known_members: list[dict[str, Any]],
) -> dict[int, str]:
    """Строит user_id -> отображаемое имя из уже загруженных профилей."""

    return {
        member["user_id"]: (
            member["current_display_name"]
            or f"участник {member['user_id']}"
        )
        for member in known_members
    }


async def build_weekly_report_text(
    chat_id: int,
    chat_type: str,
) -> str:
    """
    Собирает текст «Летописи балагана за неделю».

    Общая функция для /week и автоматической еженедельной рассылки —
    чтобы они не могли разойтись в форматировании.
    """

    start_date, end_date = get_week_date_range()
    weekly = await get_weekly_activity(chat_id, start_date, end_date)

    if not weekly:
        return "За эту неделю в чате было тихо — считать нечего."

    known_members = await list_chat_member_profiles(chat_id, limit=50)
    names = await _resolve_display_names(known_members)

    total_messages = sum(entry["messages"] for entry in weekly)
    total_voice = sum(entry["voice_messages"] for entry in weekly)
    total_photos = sum(entry["photos"] for entry in weekly)
    total_stickers = sum(entry["stickers"] for entry in weekly)
    bot_mentions = sum(entry["replies_to_bot"] for entry in weekly)
    active_participants = sum(
        1 for entry in weekly if entry["messages"] > 0
    )

    top_by_messages = sorted(
        (entry for entry in weekly if entry["messages"] > 0),
        key=lambda entry: entry["messages"],
        reverse=True,
    )[:3]

    top_lines = [
        f"{index}. {names.get(entry['user_id'], 'участник')} — "
        f"{entry['messages']} сообщ."
        for index, entry in enumerate(top_by_messages, start=1)
    ]

    awards = await get_or_create_weekly_awards(
        chat_id, start_date, weekly, known_members
    )
    awards_text = (
        format_awards_message(awards[:5], names)
        if awards
        else ""
    )

    verdict_prompt = (
        "Дай короткую (одно-два предложения) контекстную шутливую "
        f"оценку недели чата по цифрам: {total_messages} сообщений, "
        f"{active_participants} активных участников, {total_voice} "
        f"голосовых, {total_photos} фото. Без имён и лишних деталей."
    )

    try:
        verdict = await ask_gemini(
            contents=verdict_prompt,
            max_output_tokens=100,
            chat_id=chat_id,
            chat_type=chat_type,
        )
    except Exception as error:
        logging.exception(
            "Ошибка вердикта недельного отчёта: %s",
            error,
        )
        verdict = "Неделя как неделя — чат жил, Яйцеслав наблюдал."

    report = (
        "Летопись балагана за неделю\n\n"
        f"Сообщений: {total_messages}\n"
        f"Активных бояр: {active_participants}\n"
        f"Голосовых: {total_voice}\n"
        f"Фото: {total_photos}\n"
        f"Стикеров: {total_stickers}\n"
        f"Обращений к Яйцеславу: {bot_mentions}\n\n"
        "Топ писателей:\n"
        + (
            "\n".join(top_lines)
            if top_lines
            else "Пока никого не набралось."
        )
        + "\n\n"
    )

    if awards_text:
        report += awards_text + "\n\n"

    report += f"Вердикт Яйцеслава:\n{verdict}"

    return report


async def week_command(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """Показывает недельный отчёт текущей группы."""

    del context

    if not update.message or not update.effective_chat:
        return

    if update.effective_chat.type == ChatType.PRIVATE:
        await update.message.reply_text(
            "Недельный отчёт есть только для групп."
        )
        return

    report_text = await build_weekly_report_text(
        update.effective_chat.id,
        str(update.effective_chat.type),
    )

    await update.message.reply_text(report_text)


async def week_me_command(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """Показывает личную статистику участника за неделю."""

    del context

    if (
        not update.message
        or not update.effective_chat
        or not update.effective_user
    ):
        return

    if update.effective_chat.type == ChatType.PRIVATE:
        await update.message.reply_text(
            "Личная статистика недели доступна в группах."
        )
        return

    start_date, end_date = get_week_date_range()
    weekly = await get_weekly_activity(
        update.effective_chat.id, start_date, end_date
    )

    mine = next(
        (
            entry
            for entry in weekly
            if entry["user_id"] == update.effective_user.id
        ),
        None,
    )

    if mine is None or mine["messages"] == 0:
        await update.message.reply_text(
            "На этой неделе от тебя не было сообщений в этом чате."
        )
        return

    await update.message.reply_text(
        "Твоя статистика за неделю:\n"
        f"Сообщений: {mine['messages']}\n"
        f"Символов: {mine['text_characters']}\n"
        f"Голосовых: {mine['voice_messages']}\n"
        f"Фото: {mine['photos']}\n"
        f"Обращений к Яйцеславу: {mine['replies_to_bot']}\n"
        f"Ночных сообщений: {mine['night_messages']}"
    )


async def leaderboard_command(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """Показывает таблицу активности за неделю."""

    del context

    if not update.message or not update.effective_chat:
        return

    if update.effective_chat.type == ChatType.PRIVATE:
        await update.message.reply_text(
            "Таблица активности есть только в группах."
        )
        return

    chat_id = update.effective_chat.id
    start_date, end_date = get_week_date_range()
    weekly = await get_weekly_activity(chat_id, start_date, end_date)

    active = sorted(
        (entry for entry in weekly if entry["messages"] > 0),
        key=lambda entry: entry["messages"],
        reverse=True,
    )

    if not active:
        await update.message.reply_text(
            "На этой неделе активности не было."
        )
        return

    known_members = await list_chat_member_profiles(chat_id, limit=50)
    names = await _resolve_display_names(known_members)

    lines = ["Таблица активности за неделю:"]

    for index, entry in enumerate(active[:10], start=1):
        lines.append(
            f"{index}. {names.get(entry['user_id'], 'участник')} — "
            f"{entry['messages']} сообщ."
        )

    await update.message.reply_text(
        "\n".join(lines)
    )


async def awards_command(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """Показывает только шуточные награды недели."""

    del context

    if not update.message or not update.effective_chat:
        return

    if update.effective_chat.type == ChatType.PRIVATE:
        await update.message.reply_text(
            "Награды недели — это для групп."
        )
        return

    chat_id = update.effective_chat.id
    start_date, end_date = get_week_date_range()
    weekly = await get_weekly_activity(chat_id, start_date, end_date)
    known_members = await list_chat_member_profiles(chat_id, limit=50)
    names = await _resolve_display_names(known_members)

    awards = await get_or_create_weekly_awards(
        chat_id, start_date, weekly, known_members
    )

    await update.message.reply_text(
        format_awards_message(awards, names)
    )


async def week_auto_on_command(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """Включает автоматический недельный отчёт (по умолчанию — Вс 21:00 МСК)."""

    if not update.message or not update.effective_chat:
        return

    if update.effective_chat.type == ChatType.PRIVATE:
        await update.message.reply_text(
            "Авто-отчёт настраивается в группах."
        )
        return

    if not await user_is_group_admin(
        update,
        context,
    ):
        await update.message.reply_text(
            "Включать авто-отчёт может только админ."
        )
        return

    await update_chat_setting(
        update.effective_chat.id,
        "weekly_report_enabled",
        True,
        str(update.effective_chat.type),
    )

    await update.message.reply_text(
        "Авто-отчёт включён. По умолчанию — воскресенье, 21:00 по МСК. "
        "Изменить время: /week_time воскресенье 21:00"
    )


async def week_auto_off_command(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """Выключает автоматический недельный отчёт."""

    if not update.message or not update.effective_chat:
        return

    if update.effective_chat.type == ChatType.PRIVATE:
        await update.message.reply_text(
            "Авто-отчёт настраивается в группах."
        )
        return

    if not await user_is_group_admin(
        update,
        context,
    ):
        await update.message.reply_text(
            "Выключать авто-отчёт может только админ."
        )
        return

    await update_chat_setting(
        update.effective_chat.id,
        "weekly_report_enabled",
        False,
        str(update.effective_chat.type),
    )

    await update.message.reply_text(
        "Авто-отчёт выключен."
    )



async def week_auto_status_command(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """Показывает текущее расписание автоматического недельного отчёта."""

    del context

    if not update.message or not update.effective_chat:
        return

    if update.effective_chat.type == ChatType.PRIVATE:
        await update.message.reply_text(
            "Авто-отчёт существует только для групп."
        )
        return

    settings = await get_chat_settings(
        update.effective_chat.id,
        str(update.effective_chat.type),
    )

    enabled = bool(settings.get("weekly_report_enabled", False))
    weekday = int(settings.get("weekly_report_weekday", 6))
    report_time = str(settings.get("weekly_report_time", "21:00"))
    last_sent = settings.get("weekly_report_last_sent_date") or "ещё не отправлялся"

    await update.message.reply_text(
        "Автоматический недельный отчёт:\n"
        f"Статус: {'включён' if enabled else 'выключен'}\n"
        f"Расписание: {WEEKDAY_LABELS_RU.get(weekday, 'воскресенье')} "
        f"в {report_time} по МСК\n"
        f"Последняя успешная авто-отправка: {last_sent}"
    )


async def week_time_command(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """Задаёт день и время автоматического недельного отчёта."""

    if not update.message or not update.effective_chat:
        return

    if update.effective_chat.type == ChatType.PRIVATE:
        await update.message.reply_text(
            "Расписание отчёта настраивается в группах."
        )
        return

    if not await user_is_group_admin(
        update,
        context,
    ):
        await update.message.reply_text(
            "Расписание меняют только админы."
        )
        return

    if not context.args:
        await week_auto_status_command(update, context)
        return

    if len(context.args) < 2:
        await update.message.reply_text(
            "Формат: /week_time воскресенье 21:00"
        )
        return

    weekday_raw = context.args[0].strip().lower()
    time_raw = context.args[1].strip()
    weekday = WEEKDAY_NAMES_RU.get(weekday_raw)

    if weekday is None:
        await update.message.reply_text(
            "День недели не распознан. Пример: "
            "/week_time воскресенье 21:00"
        )
        return

    if not _WEEK_TIME_RE.match(time_raw):
        await update.message.reply_text(
            "Время укажи в формате ЧЧ:ММ, например 21:00."
        )
        return

    await update_chat_setting(
        update.effective_chat.id,
        "weekly_report_weekday",
        weekday,
        str(update.effective_chat.type),
    )
    await update_chat_setting(
        update.effective_chat.id,
        "weekly_report_time",
        time_raw,
        str(update.effective_chat.type),
    )

    await update.message.reply_text(
        f"Расписание обновлено: {WEEKDAY_LABELS_RU[weekday]} "
        f"в {time_raw} по МСК."
    )


def get_weekly_report_chats_sync() -> list[dict[str, Any]]:
    """Возвращает чаты с включённым авто-отчётом и их расписанием."""

    with get_db_connection() as connection:
        rows = connection.execute(
            """
            SELECT
                chat_id,
                weekly_report_weekday,
                weekly_report_time,
                weekly_report_last_sent_date
            FROM chat_settings
            WHERE weekly_report_enabled = 1
            """
        ).fetchall()

    return [
        {
            "chat_id": row[0],
            "weekday": int(row[1]),
            "time": row[2],
            "last_sent_date": row[3],
        }
        for row in rows
    ]


async def get_weekly_report_chats() -> list[dict[str, Any]]:
    """Читает список чатов с авто-отчётом без блокировки бота."""

    return await asyncio.to_thread(
        get_weekly_report_chats_sync
    )


def mark_weekly_report_sent_sync(
    chat_id: int,
) -> None:
    """Отмечает, что авто-отчёт за сегодня уже отправлен."""

    with get_db_connection() as connection:
        connection.execute(
            """
            UPDATE chat_settings
            SET weekly_report_last_sent_date = ?
            WHERE chat_id = ?
            """,
            (current_msk_date_str(), chat_id),
        )
        connection.commit()


async def mark_weekly_report_sent(
    chat_id: int,
) -> None:
    """Отмечает отправку без блокировки бота."""

    await asyncio.to_thread(
        mark_weekly_report_sent_sync,
        chat_id,
    )


# Раз в минуту — недорого (запросов обычно единицы), зато не
# промахнёмся мимо назначенной минуты один раз в неделю.
WEEKLY_REPORT_CHECK_INTERVAL_SECONDS = 60


def _time_str_to_minutes(
    time_str: str,
) -> int:
    """Переводит «ЧЧ:ММ» в минуты от полуночи для сравнения времени."""

    hour_str, minute_str = time_str.split(":")
    return int(hour_str) * 60 + int(minute_str)


async def run_due_weekly_reports(
    application: Application,
) -> None:
    """
    Отправляет авто-отчёт во все чаты, для которых наступило их время.

    Отмечает отчёт отправленным только при успешной отправке — если
    send_message упал (бота убрали из группы, сеть моргнула), чат
    просто получит повторную попытку на следующей минуте, а не будет
    молча пропущен до следующей недели.
    """

    now = current_msk_datetime()
    today_str = now.strftime("%Y-%m-%d")
    now_minutes = now.hour * 60 + now.minute

    for chat in await get_weekly_report_chats():
        if chat["weekday"] != now.weekday():
            continue

        if chat["last_sent_date"] == today_str:
            continue

        if now_minutes < _time_str_to_minutes(chat["time"]):
            continue

        try:
            report_text = await build_weekly_report_text(
                chat["chat_id"], "group"
            )
            await application.bot.send_message(
                chat_id=chat["chat_id"],
                text=report_text,
            )
        except Exception as error:
            logging.warning(
                "Не удалось отправить авто-отчёт в чат %s: %s "
                "(попробуем снова на следующей минуте)",
                chat["chat_id"],
                error,
            )
            continue

        await mark_weekly_report_sent(chat["chat_id"])


async def weekly_report_scheduler_loop(
    application: Application,
) -> None:
    """Фоновая задача: проверяет расписание авто-отчётов раз в минуту."""

    while True:
        await asyncio.sleep(WEEKLY_REPORT_CHECK_INTERVAL_SECONDS)

        try:
            await run_due_weekly_reports(application)
        except Exception as error:
            logging.warning(
                "Ошибка планировщика недельных отчётов: %s",
                error,
            )



# ============================================================
# АВТОМАТИЧЕСКИЙ ТИТУЛ ДНЯ
# ============================================================

DAILY_TITLE_CHECK_INTERVAL_SECONDS = 60


def get_daily_title_chat_ids_sync(date: str) -> list[int]:
    """Группы, где сегодня есть хотя бы один активный участник."""

    with get_db_connection() as connection:
        rows = connection.execute(
            """
            SELECT DISTINCT activity.chat_id
            FROM chat_activity_daily AS activity
            JOIN chats ON chats.chat_id = activity.chat_id
            WHERE activity.date = ?
              AND activity.messages > 0
              AND chats.chat_type IN ('group', 'supergroup')
            ORDER BY activity.chat_id
            """,
            (date,),
        ).fetchall()

    return [int(row[0]) for row in rows]


async def get_daily_title_chat_ids(date: str) -> list[int]:
    return await asyncio.to_thread(
        get_daily_title_chat_ids_sync,
        date,
    )


def mark_daily_title_announced_sync(chat_id: int, date: str) -> None:
    with get_db_connection() as connection:
        connection.execute(
            """
            UPDATE daily_title_assignments
            SET announced_at = datetime('now')
            WHERE chat_id = ? AND date = ?
            """,
            (chat_id, date),
        )
        connection.commit()


async def mark_daily_title_announced(chat_id: int, date: str) -> None:
    await asyncio.to_thread(
        mark_daily_title_announced_sync,
        chat_id,
        date,
    )


async def _daily_title_display_name(chat_id: int, user_id: int) -> str:
    profile = await get_member_profile(chat_id, user_id)
    if profile and profile.get("current_display_name"):
        return str(profile["current_display_name"])
    return f"участник {user_id}"


async def run_due_daily_titles(application: Application) -> None:
    """После 18:00 МСК выдаёт ровно один daily title каждой активной группе."""

    now = current_msk_datetime()
    if not daily_title_engine.is_assignment_window_open(now):
        return

    date = now.date().isoformat()

    for chat_id in await get_daily_title_chat_ids(date):
        assignment = await get_daily_title_assignment(chat_id, date)
        display_name: str | None = None

        if assignment and assignment.get("announced_at"):
            continue

        if assignment is None:
            activity = await get_weekly_activity(chat_id, date, date)
            known_members = await list_chat_member_profiles(chat_id, limit=200)
            candidates = daily_title_engine.build_candidates(
                activity,
                known_members,
            )
            candidate = daily_title_engine.choose_candidate(candidates)
            if candidate is None:
                continue

            new_title = pick_new_title(candidate.previous_title)
            created = await try_assign_daily_title(
                chat_id,
                date,
                candidate.user_id,
                new_title,
            )

            if created:
                assignment = {
                    "user_id": candidate.user_id,
                    "title": new_title,
                    "announced_at": None,
                }
                display_name = candidate.display_name
            else:
                assignment = await get_daily_title_assignment(chat_id, date)

        if not assignment or assignment.get("announced_at"):
            continue

        if display_name is None:
            display_name = await _daily_title_display_name(
                chat_id,
                int(assignment["user_id"]),
            )

        try:
            await application.bot.send_message(
                chat_id=chat_id,
                text=daily_title_engine.format_daily_title_message(
                    display_name,
                    str(assignment["title"]),
                ),
            )
        except Exception as error:
            logging.warning(
                "Не удалось объявить титул дня в чате %s: %s "
                "(повторим на следующей минуте)",
                chat_id,
                error,
            )
            continue

        await mark_daily_title_announced(chat_id, date)


async def daily_title_scheduler_loop(application: Application) -> None:
    """Фоновая задача: после 18:00 МСК проверяет daily titles раз в минуту."""

    while True:
        await asyncio.sleep(DAILY_TITLE_CHECK_INTERVAL_SECONDS)

        try:
            await run_due_daily_titles(application)
        except Exception as error:
            logging.warning(
                "Ошибка планировщика титула дня: %s",
                error,
            )


CHAT_NATIVE_REFRESH_CHECK_INTERVAL_SECONDS = 6 * 60 * 60


async def chat_native_refresh_loop(application: Application) -> None:
    """Периодически собирает/обновляет 13-й voice pack каждого чата."""

    del application
    while True:
        try:
            refreshed = await refresh_due_chat_native_profiles()
            if refreshed:
                logging.info("Обновлены chat_native профили: %s", refreshed)
        except Exception as error:
            logging.warning("Ошибка обновления chat_native: %s", error)
        await asyncio.sleep(CHAT_NATIVE_REFRESH_CHECK_INTERVAL_SECONDS)


# ============================================================
# АНТИСПАМ ДЛЯ СЛУЧАЙНЫХ ВМЕШАТЕЛЬСТВ В ГРУППЕ
#
# Даже агрессивный Яйцеслав не должен отвечать на всё подряд.
# ============================================================

# Случайная текстовая реплика — не чаще раза в 2 минуты
GROUP_RANDOM_REPLY_MIN_INTERVAL = 120.0

# И не больше трёх штук за скользящее окно в 10 минут
GROUP_RANDOM_REPLY_MAX_PER_WINDOW = 3
GROUP_RANDOM_REPLY_WINDOW_SECONDS = 600.0

# Одному и тому же человеку — не чаще раза в 60 секунд,
# если он не обращается к боту напрямую (тогда действует
# обычный enforce_rate_limit, а не это ограничение)
TRIGGER_REPLY_PER_USER_COOLDOWN = 60.0

# После серьёзного сообщения — пауза в случайном юморе
SERIOUS_TOPIC_HUMOR_COOLDOWN = 300.0

# Тихие часы по МСК: почти не вмешиваемся случайно
QUIET_HOURS_START_MSK = 0
QUIET_HOURS_END_MSK = 7

# Если последние два случайных вмешательства проигнорировали —
# лучше помолчать, чем настаивать
IGNORED_STREAK_LIMIT = 2

GROUP_RANDOM_REPLY_TIMES: dict[int, deque] = defaultdict(deque)
GROUP_LAST_SERIOUS_AT: dict[int, float] = {}
GROUP_IGNORED_STREAK: dict[int, int] = defaultdict(int)
TRIGGER_REPLY_LAST_BY_USER: dict[tuple[int, int], float] = {}

# Последнее обычное сообщение пользователя в конкретном чате — чтобы
# команды вроде /argument, /debate, /judge могли сработать по теме,
# которую человек написал ДО команды, без повторного её ввода.
LAST_USER_TEXT_MESSAGE: dict[tuple[int, int], tuple[float, str]] = {}
LAST_USER_TEXT_MESSAGE_MAX_AGE_SECONDS = 30 * 60


def record_last_user_message(
    chat_id: int,
    user_id: int,
    text: str,
) -> None:
    """Запоминает последнее обычное сообщение пользователя в этом чате."""

    if not text:
        return

    LAST_USER_TEXT_MESSAGE[(chat_id, user_id)] = (
        time.monotonic(),
        text,
    )


def get_last_user_message(
    chat_id: int,
    user_id: int,
) -> str | None:
    """Возвращает последнее сообщение пользователя, если оно не устарело."""

    entry = LAST_USER_TEXT_MESSAGE.get((chat_id, user_id))

    if entry is None:
        return None

    recorded_at, text = entry

    if time.monotonic() - recorded_at > LAST_USER_TEXT_MESSAGE_MAX_AGE_SECONDS:
        return None

    return text


def is_quiet_hours_msk() -> bool:
    """Проверяет тихие часы (00:00–07:00 по МСК) для случайных вмешательств."""

    msk_now = datetime.now(
        timezone(timedelta(hours=3))
    )
    return QUIET_HOURS_START_MSK <= msk_now.hour < QUIET_HOURS_END_MSK


def is_serious_cooldown_active(
    chat_id: int,
    now: float,
) -> bool:
    """Проверяет, недавно ли в чате была серьёзная тема."""

    last_serious_at = GROUP_LAST_SERIOUS_AT.get(chat_id)

    return (
        last_serious_at is not None
        and now - last_serious_at < SERIOUS_TOPIC_HUMOR_COOLDOWN
    )


def group_random_reply_allowed(
    chat_id: int,
    now: float,
) -> bool:
    """
    Проверяет групповые лимиты случайных реплик: не чаще раза
    в 2 минуты, не больше трёх за 10 минут, тишина по ночам,
    пауза после серьёзной темы и после двух проигнорированных подряд.
    """

    if is_quiet_hours_msk():
        return False

    if is_serious_cooldown_active(chat_id, now):
        return False

    if GROUP_IGNORED_STREAK[chat_id] >= IGNORED_STREAK_LIMIT:
        return False

    history = GROUP_RANDOM_REPLY_TIMES[chat_id]

    while (
        history
        and now - history[0] > GROUP_RANDOM_REPLY_WINDOW_SECONDS
    ):
        history.popleft()

    if (
        history
        and now - history[-1] < GROUP_RANDOM_REPLY_MIN_INTERVAL
    ):
        return False

    if len(history) >= GROUP_RANDOM_REPLY_MAX_PER_WINDOW:
        return False

    return True


def record_group_random_reply(
    chat_id: int,
    now: float,
) -> None:
    """Запоминает момент случайной реплики и растит счётчик игнора."""

    GROUP_RANDOM_REPLY_TIMES[chat_id].append(now)
    GROUP_IGNORED_STREAK[chat_id] += 1


def register_group_engagement(
    chat_id: int,
) -> None:
    """Сбрасывает счётчик игнора — кто-то ответил боту или обратился к нему."""

    GROUP_IGNORED_STREAK.pop(chat_id, None)


async def hard_mode_listener(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """
    Слушает обычную переписку группы.

    Не обращается к Gemini и не сохраняет сообщения.
    """

    if (
        not update.message
        or not update.message.text
        or not update.effective_chat
        or not update.effective_user
    ):
        return

    if update.effective_chat.type not in (
        ChatType.GROUP,
        ChatType.SUPERGROUP,
    ):
        return

    # Не сохраняем сообщения других ботов
    if update.effective_user.is_bot:
        return

    text = update.message.text.strip()

    if not text:
        return

    # Команды в память не записываем
    if text.startswith("/"):
        return

    # 13-й pack учится и на прямых обращениях, и на фоне чата. Полный
    # текст в SQLite не сохраняется — только извлечённые агрегированные термы.
    if not is_serious_text(text.lower()):
        await asyncio.to_thread(
            record_chat_native_message_sync,
            update.effective_chat.id,
            update.effective_user.id,
            text,
            str(update.effective_chat.type),
        )

    bot_username = await get_bot_username(
        context
    )

    has_direct_mention = bool(
        bot_username
        and f"@{bot_username}".lower() in text.lower()
    )

    has_reply = is_reply_to_bot(
        update,
        context,
    )

    has_text_address = (
        extract_group_address(text)
        is not None
    )

    # Прямое обращение обработает основной обработчик
    if (
        has_direct_mention
        or has_reply
        or has_text_address
    ):
        return

    author_name = (
        update.effective_user.full_name
        or update.effective_user.username
        or "Участник"
    )

    chat_id = update.effective_chat.id
    chat_type = str(update.effective_chat.type)

    await touch_member_profile(
        chat_id,
        update.effective_user.id,
        chat_type,
        author_name,
        update.effective_user.username,
    )

    await increment_chat_activity(
        chat_id,
        update.effective_user.id,
        chat_type,
        current_msk_date_str(),
        **build_text_activity_deltas(text),
    )

    record_last_user_message(
        chat_id,
        update.effective_user.id,
        text,
    )

    # Смотрим предыдущее сообщение группы ДО того, как допишем
    # в память текущее — иначе "повтор одного и того же" будет
    # сравнивать сообщение само с собой.
    previous_messages = GROUP_MEMORY.get(chat_id)
    previous_group_text = (
        previous_messages[-1][3]
        if previous_messages
        else None
    )

    if is_serious_text(text.lower()):
        GROUP_LAST_SERIOUS_AT[chat_id] = time.monotonic()

    # Запоминаем обычное сообщение группы на пять минут
    remember_message(
        GROUP_MEMORY,
        chat_id,
        "user",
        text,
        GROUP_MEMORY_SECONDS,
        GROUP_MEMORY_MAX_MESSAGES,
        author_name,
    )

    # Память работает всегда,
    # а реакции и случайные реплики — только в хард-моде
    chat_settings_data = await get_chat_settings(
        chat_id,
        chat_type,
    )

    if not chat_settings_data["hard_mode_enabled"]:
        return

    level_chances = HARD_LEVEL_CHANCES.get(
        chat_settings_data["hard_level"],
        HARD_LEVEL_CHANCES["normal"],
    )
    reaction_chance = level_chances["reaction_chance"]
    random_reply_chance = level_chances["random_reply_chance"]

    now = time.monotonic()

    # Серьёзная тема ставит хард-мод на паузу целиком —
    # это не время для реакций, триггеров или случайных реплик.
    if is_serious_cooldown_active(chat_id, now):
        return

    # --------------------------------------------------------
    # 1. Реакция на специальные слова
    # --------------------------------------------------------

    trigger_reply = choose_hard_trigger_reply(
        text
    )

    last_trigger_reply = float(
        context.chat_data.get(
            "hard_last_trigger_reply",
            0.0,
        )
    )

    user_id = update.effective_user.id
    trigger_user_key = (chat_id, user_id)
    last_trigger_reply_for_user = TRIGGER_REPLY_LAST_BY_USER.get(
        trigger_user_key, 0.0
    )

    if (
        trigger_reply
        and now - last_trigger_reply >= HARD_TRIGGER_COOLDOWN
        and now - last_trigger_reply_for_user
        >= TRIGGER_REPLY_PER_USER_COOLDOWN
    ):
        # Резервируем cooldown ДО сетевого await: второй concurrent update
        # уже увидит занятый слот и не отправит дубль.
        context.chat_data[
            "hard_last_trigger_reply"
        ] = now
        TRIGGER_REPLY_LAST_BY_USER[trigger_user_key] = now

        await update.message.reply_text(
            trigger_reply
        )

        await increment_chat_hard_stat(
            chat_id,
            "trigger_replies_count",
            chat_type,
        )

        return

    # --------------------------------------------------------
    # 2. Контекстная реакция-эмодзи
    # --------------------------------------------------------

    last_reaction = float(
        context.chat_data.get(
            "hard_last_reaction",
            0.0,
        )
    )

    reaction_reason = detect_reaction_reason(
        text,
        is_reply_to_bot_message=False,
        previous_group_text=previous_group_text,
    )

    # V2: контекстная причина всё ещё усиливает шанс, но ПОСЛЕ
    # этого применяется -20% только к emoji-реакциям. Текстовый
    # random_reply_chance ниже остаётся полностью без изменений.
    effective_reaction_chance = reaction_engine.effective_emoji_reaction_chance(
        reaction_chance,
        has_context_reason=bool(reaction_reason),
    )

    reacted_to_this_message = False

    if (
        now - last_reaction
        >= HARD_REACTION_COOLDOWN
        and random.random()
        < effective_reaction_chance
    ):
        reputation_score = await asyncio.to_thread(
            _member_reputation_score_sync,
            chat_id,
            update.effective_user.id,
        )
        reaction_emoji = pick_reaction_emoji(
            reaction_reason,
            reputation_score=reputation_score,
        )

        # Та же защита от concurrent_updates(8): резервируем
        # реакционный cooldown до Telegram API await.
        context.chat_data[
            "hard_last_reaction"
        ] = now

        try:
            await update.message.set_reaction(
                reaction=[
                    ReactionTypeEmoji(
                        reaction_emoji
                    )
                ],
                is_big=False,
            )

            reacted_to_this_message = True

            await increment_chat_hard_stat(
                chat_id,
                "reactions_count",
                chat_type,
            )

        except Exception as error:
            # Некоторые группы разрешают не все реакции.
            logging.debug(
                "Не удалось поставить реакцию %s: %s",
                reaction_emoji,
                error,
            )

    # --------------------------------------------------------
    # 3. Редкое случайное вмешательство
    #
    # Реакция и текстовая реплика на одно и то же сообщение —
    # это уже два вмешательства сразу, чего быть не должно.
    # V2 random drop не создаёт новый слот: он может только заменить
    # содержимое уже разрешённой ниже random reply.
    # --------------------------------------------------------

    passive_engine.note_group_activity(chat_id)

    last_random_reply = float(
        context.chat_data.get(
            "hard_last_random_reply",
            0.0,
        )
    )

    if (
        not reacted_to_this_message
        and passive_engine.random_text_intervention_allowed(
            text, reaction_reason
        )
        and now - last_random_reply >= HARD_RANDOM_REPLY_COOLDOWN
        and random.random() < random_reply_chance
        and group_random_reply_allowed(chat_id, now)
    ):
        drop_decision = passive_engine.maybe_random_drop(
            chat_id,
            existing_random_reply_slot_open=True,
            now=now,
        )
        if drop_decision.active and drop_decision.text:
            # Резервируем слот до await, иначе два апдейта могут пройти
            # group_random_reply_allowed одновременно.
            context.chat_data[
                "hard_last_random_reply"
            ] = now
            record_group_random_reply(chat_id, now)

            await update.message.reply_text(drop_decision.text)

            await increment_chat_hard_stat(
                chat_id,
                "random_replies_count",
                chat_type,
            )
async def enforce_rate_limit(
    update: Update,
    bucket: str,
) -> bool:
    """
    Проверяет лимит пользователя.

    Возвращает True, если запрос разрешён.
    Возвращает False, если лимит превышен.
    """

    if not update.effective_user:
        return False

    user_id = update.effective_user.id
    now = time.monotonic()

    limit, window_seconds = RATE_LIMITS[bucket]
    key = (user_id, bucket)
    history = REQUEST_TIMES[key]

    # Удаляем старые запросы, которые уже не входят в окно
    while (
        history
        and now - history[0] >= window_seconds
    ):
        history.popleft()

    if len(history) >= limit:
        await increment_stat(
            "rate_limit_hits"
        )
        wait_seconds = max(
            1,
            int(
                window_seconds
                - (now - history[0])
            ) + 1,
        )

        last_warning = LAST_LIMIT_WARNING.get(
            key,
            0.0,
        )

        if (
            update.effective_message
            and now - last_warning
            >= LIMIT_WARNING_COOLDOWN
        ):
            await update.effective_message.reply_text(
                f"Полегче, пулемётчик. "
                f"Подожди примерно {wait_seconds} сек."
            )

            LAST_LIMIT_WARNING[key] = now

        return False

    history.append(now)
    return True

# ============================================================
# КОМАНДЫ
# ============================================================

async def start(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """Команда /start."""

    del context

    if not update.message:
        return

    await update.message.reply_text(
        "Здарова. Я Яйцеслав 🥚\n\n"
        "Пиши вопрос, кидай фото, документ "
        "или голосовуху.\n"
        "Для голоса напиши: «ответь голосом»."
    )


HELP_INTRO_TEXT = (
    "Что умеет Яйцеслав:\n"
    "• отвечает на вопросы;\n"
    "• ищет свежую информацию в интернете;\n"
    "• смотрит фотографии и мемы;\n"
    "• читает PDF, DOCX, XLSX, CSV и TXT;\n"
    "• понимает голосовые сообщения;\n"
    "• может отвечать голосом;\n"
    "• помнит личный разговор 15 минут;\n"
    "• помнит переписку группы 15 минут.\n\n"
    "Как обратиться в группе:\n"
    "• через @имя_бота;\n"
    "• ответом на сообщение Яйцеслава;\n"
    "• словами: Яйцеслав, бобр, эй бобр, "
    "курва, бот, помощник.\n\n"
    "Голосовое в группе:\n"
    "• отправляй его ответом на сообщение Яйцеслава.\n\n"
)

HELP_SEARCH_SECTION = (
    "🔎 Поиск и анализ:\n"
    "/search запрос — поиск в интернете\n"
    "/argument — разбор спора по последнему сообщению "
    "(или ответом на позицию)\n"
    "/debate тема — аргументы за и против\n"
    "/judge — вердикт по последнему сообщению (или ответом)\n"
    "/fact_or_bayan — факт, баян или сомнительно "
    "по последнему сообщению\n"
    "/translate_yayceslav — перевод с канцелярского "
    "последнего сообщения\n"
    "/explain_like_skoof тема — объяснение по-простому\n"
    "/explain_like_rus тема — объяснение в стиле древнего руса\n"
)

HELP_GROUP_ANALYTICS_SECTION = (
    "\n📊 Аналитика группы:\n"
    "/recap — пересказ последних сообщений группы\n"
    "/week — недельный отчёт группы\n"
    "/week_me — твоя статистика за неделю\n"
    "/leaderboard — таблица активности\n"
    "/awards — шуточные награды недели\n"
    "/week_auto_status — статус автоматического недельного отчёта\n"
    "/title_status — статус автоматического титула дня\n"
    "/chat_native_status — чему Яйцеслав уже научился у этого чата\n"
)

HELP_ENTERTAINMENT_SECTION = (
    "\n🎭 Развлечения:\n"
    "/roast — токсично прокомментировать последнее сообщение "
    "(или ответом)\n"
    "/meme тема — мемная подпись\n"
    "/anti_advice тема — сначала плохой совет, потом настоящий\n"
    "/duel — вызвать на шуточную дуэль (ответом на сообщение)\n"
    "/story — продолжить историю группы\n"
    "/title — выдать/сменить шуточный титул "
    "(ответом — другому участнику)\n"
    "/prophecy — псевдопророчество\n"
    "/wisdom — мудрость Яйцеслава\n"
    "/mood — настроение Яйцеслава\n"
)

HELP_PROFILE_SECTION = (
    "\n👤 Профиль и память:\n"
    "/profile — твои настройки и статус\n"
    "/whoami — как бот видит тебя в этом чате\n"
    "/nickname Имя — как к тебе обращаться\n"
    "/nickname_off — отключить личное обращение\n"
    "/remember_me текст — сохранить факт о себе\n"
    "/forget_me — удалить свой профиль в чате\n"
    "/forget — очистить кратковременную память\n"
)

HELP_SETTINGS_SECTION = (
    "\n⚙️ Настройки:\n"
    "/voice_on — всегда отвечать голосом\n"
    "/voice_off — отвечать текстом\n"
    "/settings — персональные настройки\n"
    "/hard_status — проверить активность в группе\n"
)

HELP_ADMIN_SECTION = (
    "\n🛠 Админ-команды группы:\n"
    "/hard_on — включить активность в группе\n"
    "/hard_off — выключить активность в группе\n"
    "/hard_level calm|normal|chaos — настроить накал\n"
    "/hard_stats — статистика хард-мода\n"
    "/people — известные участники\n"
    "/set_archetype текст — задать архетип в ответ на сообщение\n"
    "/week_auto_on — включить авто-отчёт\n"
    "/week_auto_off — выключить авто-отчёт\n"
    "/week_time день ЧЧ:ММ — расписание авто-отчёта\n"
)

HELP_FOOTER_TEXT = (
    "\nСначала напиши сообщение, затем выбери команду. "
    "Можно также ответить командой на конкретную реплику."
)


async def help_command(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """Команда /help."""

    if not update.message:
        return

    sections = [
        HELP_INTRO_TEXT,
        HELP_SEARCH_SECTION,
        HELP_GROUP_ANALYTICS_SECTION,
        HELP_ENTERTAINMENT_SECTION,
        HELP_PROFILE_SECTION,
        HELP_SETTINGS_SECTION,
    ]

    if await user_is_group_admin(update, context):
        sections.append(HELP_ADMIN_SECTION)

    sections.append(HELP_FOOTER_TEXT)

    await update.message.reply_text("".join(sections))
async def stats_command(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """Показывает статистику только владельцу бота."""

    del context

    if (
        not update.message
        or not update.effective_user
        or not update.effective_chat
    ):
        return

    # В группах статистику не показываем
    if (
        update.effective_chat.type
        != ChatType.PRIVATE
    ):
        return

    # Посторонним ничего не отвечаем
    if (
        BOT_OWNER_ID == 0
        or update.effective_user.id
        != BOT_OWNER_ID
    ):
        return

    stats = await get_stats_snapshot()

    await update.message.reply_text(
        "📊 Статистика Яйцеслава\n\n"
        f"👤 Уникальных пользователей: "
        f"{stats.get('unique_users', 0)}\n"
        f"💬 Личных чатов: "
        f"{stats.get('private_chats', 0)}\n"
        f"👥 Групп: "
        f"{stats.get('groups', 0)}\n\n"
        f"📨 Всего запросов: "
        f"{stats.get('total_requests', 0)}\n"
        f"✍️ Текстовых: "
        f"{stats.get('text_requests', 0)}\n"
        f"🔎 Интернет-поисков: "
        f"{stats.get('search_requests', 0)}\n"
        f"🖼 Фотографий: "
        f"{stats.get('photo_requests', 0)}\n"
        f"📄 Документов: "
        f"{stats.get('document_requests', 0)}\n"
        f"🎙 Голосовых и аудио: "
        f"{stats.get('voice_requests', 0)}\n\n"
        f"🥚 Ответов Яйцеслава: "
        f"{stats.get('bot_answers', 0)}\n"
        f"🚫 Срабатываний антифлуда: "
        f"{stats.get('rate_limit_hits', 0)}"
    )    
async def forget_command(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """Очищает кратковременную память."""

    if (
        not update.message
        or not update.effective_chat
    ):
        return

    # В личке пользователь очищает собственную память
    if (
        update.effective_chat.type
        == ChatType.PRIVATE
    ):
        if update.effective_user:
            PRIVATE_MEMORY.pop(
                update.effective_user.id,
                None,
            )

        await update.message.reply_text(
            "Личную память очистил. "
            "Начинаем с чистого яйца."
        )
        return

    # В группе очищать общую память может только администратор
    if update.effective_chat.type in (
        ChatType.GROUP,
        ChatType.SUPERGROUP,
    ):
        if not await user_is_group_admin(
            update,
            context,
        ):
            await update.message.reply_text(
                "Память группы очищают только админы, нищий."
            )
            return

        GROUP_MEMORY.pop(
            update.effective_chat.id,
            None,
        )

        await update.message.reply_text(
            "Память группы очищена. "
            "Ваш цифровой позор временно забыт."
        )
async def voice_on(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """Включает постоянные голосовые ответы."""

    context.user_data[
        "voice_mode"
    ] = True

    # Раньше это жило только в context.user_data и обнулялось
    # при каждом рестарте Railway. Настройка voice_enabled уже
    # существует в user_settings — используем её как источник
    # истины, а user_data оставляем для мгновенного эффекта
    # в рамках этой же сессии.
    if update.effective_user:
        await update_user_setting(
            update.effective_user.id,
            "voice_enabled",
            True,
        )

    if update.message:
        await update.message.reply_text(
            "Голос включён. "
            "Теперь Яйцеслав вещает."
        )


async def voice_off(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """Возвращает текстовые ответы."""

    context.user_data[
        "voice_mode"
    ] = False

    if update.effective_user:
        await update_user_setting(
            update.effective_user.id,
            "voice_enabled",
            False,
        )

    if update.message:
        await update.message.reply_text(
            "Голос выключен. "
            "Снова читаешь глазами, легенда."
        )


# ============================================================
# ПРОФИЛЬ, ОБРАЩЕНИЕ И ЛЁГКАЯ ПАМЯТЬ ОБ УЧАСТНИКЕ
# ============================================================

NICKNAME_MAX_LENGTH = 32
REMEMBER_ME_MAX_LENGTH = 200
JOKE_ARCHETYPE_MAX_LENGTH = 150

_INJECTION_MARKERS = (
    "игнорируй",
    "ignore previous",
    "ignore all",
    "system prompt",
    "ты теперь",
    "you are now",
    "забудь инструкции",
    "новая роль",
    "act as",
)


def sanitize_user_supplied_text(
    text: str,
    max_length: int,
) -> str | None:
    """
    Проверяет текст, который пользователь просит запомнить или
    использовать как обращение к себе.

    Отклоняет команды, управляющие символы, слишком длинный текст
    и явные попытки вставить инструкции для модели.
    """

    cleaned = text.strip()

    if not cleaned or len(cleaned) > max_length:
        return None

    if cleaned.startswith("/"):
        return None

    if any(
        character in "\n\r\t\x00"
        or not (character.isprintable() or character == " ")
        for character in cleaned
    ):
        return None

    lowered = cleaned.lower()

    if any(marker in lowered for marker in _INJECTION_MARKERS):
        return None

    return cleaned


async def profile_command(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """Показывает пользователю его собственные настройки и статус."""

    if (
        not update.message
        or not update.effective_user
        or not update.effective_chat
    ):
        return

    settings = await get_user_settings(
        update.effective_user.id
    )
    profile = await get_member_profile(
        update.effective_chat.id,
        update.effective_user.id,
    )

    nickname = settings.get("custom_nickname") or "не задано"
    relationship_level = (
        profile["relationship_level"] if profile else 0
    )
    total_messages = (
        profile["total_messages"] if profile else 0
    )
    current_title = (
        (profile.get("current_title") if profile else None)
        or "пока нет — держи /title"
    )

    await update.message.reply_text(
        "Твой профиль у Яйцеслава:\n"
        f"Персонаж: {CHARACTER_LABELS.get(settings['character'], settings['character'])}\n"
        f"Стиль: {STYLE_LABELS.get(settings['response_style'], settings['response_style'])}\n"
        f"Длина ответа: {LENGTH_LABELS.get(settings['response_length'], settings['response_length'])}\n"
        f"Грубость: {ROUGHNESS_LABELS.get(settings['roughness'], settings['roughness'])}\n"
        f"Голос: {'включён' if settings['voice_enabled'] else 'выключен'}\n"
        f"Поиск: {SEARCH_MODE_LABELS.get(settings['search_mode'], settings['search_mode'])}\n"
        f"Обращение: {nickname}\n"
        f"Титул: {current_title}\n"
        f"Уровень знакомства в этом чате: {relationship_level}\n"
        f"Сообщений в этом чате: {total_messages}"
    )


async def nickname_command(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """Задаёт личное обращение, которым Яйцеслав будет звать пользователя."""

    if not update.message or not update.effective_user:
        return

    if not context.args:
        await update.message.reply_text(
            "Напиши, как к тебе обращаться: /nickname Твой Ник"
        )
        return

    requested = " ".join(context.args)
    nickname = sanitize_user_supplied_text(
        requested, NICKNAME_MAX_LENGTH
    )

    if nickname is None:
        await update.message.reply_text(
            "Такое обращение не подойдёт: без команд, спецсимволов "
            f"и длиннее {NICKNAME_MAX_LENGTH} символов нельзя."
        )
        return

    await update_user_setting(
        update.effective_user.id,
        "custom_nickname",
        nickname,
    )

    await update.message.reply_text(
        f"Принято. Теперь иногда буду звать тебя «{nickname}»."
    )


async def nickname_off_command(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """Отключает пользовательское обращение."""

    if not update.message or not update.effective_user:
        return

    await update_user_setting(
        update.effective_user.id,
        "custom_nickname",
        None,
    )

    await update.message.reply_text(
        "Личное обращение отключено."
    )


RELATIONSHIP_LEVEL_LABELS = {
    0: "незнакомец",
    1: "знакомый",
    2: "постоянный участник",
    3: "старожил",
}


async def whoami_command(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """Показывает пользователю, как бот воспринимает его роль в чате."""

    if (
        not update.message
        or not update.effective_user
        or not update.effective_chat
    ):
        return

    profile = await get_member_profile(
        update.effective_chat.id,
        update.effective_user.id,
    )

    if profile is None or profile["total_messages"] == 0:
        await update.message.reply_text(
            "Пока не успел тебя толком запомнить, гой. Пиши больше."
        )
        return

    archetype_line = (
        f"Архетип по мнению чата: {profile['joke_archetype']}\n"
        if profile["joke_archetype"]
        else ""
    )

    title_line = (
        f"Титул: {profile['current_title']}\n"
        if profile.get("current_title")
        else ""
    )

    facts_line = (
        (
            "Что ты сам просил запомнить: "
            + "; ".join(profile["self_reported_facts"])
            + "\n"
        )
        if profile["self_reported_facts"]
        else ""
    )

    await update.message.reply_text(
        "Как Яйцеслав тебя видит в этом чате:\n"
        f"Статус: {RELATIONSHIP_LEVEL_LABELS.get(profile['relationship_level'], 'незнакомец')}\n"
        f"Сообщений: {profile['total_messages']}\n"
        f"{archetype_line}"
        f"{title_line}"
        f"{facts_line}"
    )


async def remember_me_command(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """Разрешает пользователю самому сохранить безопасный факт о себе."""

    if (
        not update.message
        or not update.effective_user
        or not update.effective_chat
    ):
        return

    if not context.args:
        await update.message.reply_text(
            "Напиши, что запомнить: /remember_me люблю чай без сахара"
        )
        return

    requested = " ".join(context.args)
    fact = sanitize_user_supplied_text(
        requested, REMEMBER_ME_MAX_LENGTH
    )

    if fact is None:
        await update.message.reply_text(
            "Это не сохранить: без команд, спецсимволов и не длиннее "
            f"{REMEMBER_ME_MAX_LENGTH} символов."
        )
        return

    await append_self_reported_fact(
        update.effective_chat.id,
        update.effective_user.id,
        fact,
        str(update.effective_chat.type),
    )

    await update.message.reply_text(
        "Запомнил. Можешь посмотреть через /whoami."
    )


async def forget_me_command(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """Удаляет профиль участника в этой группе."""

    if (
        not update.message
        or not update.effective_user
        or not update.effective_chat
    ):
        return

    await delete_member_profile(
        update.effective_chat.id,
        update.effective_user.id,
    )

    await update.message.reply_text(
        "Твой профиль в этом чате удалён."
    )


async def people_command(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """Показывает администратору известных активных участников чата."""

    if not update.message or not update.effective_chat:
        return

    if update.effective_chat.type == ChatType.PRIVATE:
        await update.message.reply_text(
            "Список участников есть только в группах."
        )
        return

    if not await user_is_group_admin(
        update,
        context,
    ):
        await update.message.reply_text(
            "Список участников смотрят только админы."
        )
        return

    profiles = await list_chat_member_profiles(
        update.effective_chat.id
    )

    if not profiles:
        await update.message.reply_text(
            "Пока никого не запомнил в этом чате."
        )
        return

    lines = ["Известные участники этого чата:"]

    for profile in profiles:
        name = profile["current_display_name"] or "без имени"

        archetype = (
            f" ({profile['joke_archetype']})"
            if profile["joke_archetype"]
            else ""
        )

        lines.append(
            f"- {name}{archetype}: {profile['total_messages']} сообщ., "
            f"уровень {profile['relationship_level']}"
        )

    await update.message.reply_text(
        "\n".join(lines)
    )


async def set_archetype_command(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """
    Задаёт шуточный архетип участника (для callback-юмора).

    Только вручную, только админом, только через ответ на
    сообщение того человека — никакого автоматического вывода.
    """

    if not update.message or not update.effective_chat:
        return

    if update.effective_chat.type == ChatType.PRIVATE:
        await update.message.reply_text(
            "Архетипы задаются только в группах."
        )
        return

    if not await user_is_group_admin(
        update,
        context,
    ):
        await update.message.reply_text(
            "Архетип участника задают только админы."
        )
        return

    target_message = update.message.reply_to_message

    if not target_message or not target_message.from_user:
        await update.message.reply_text(
            "Ответь этой командой на сообщение человека, "
            "которому хочешь задать архетип."
        )
        return

    if not context.args:
        await update.message.reply_text(
            "Напиши архетип: /set_archetype скуф"
        )
        return

    requested = " ".join(context.args)
    archetype = sanitize_user_supplied_text(
        requested, JOKE_ARCHETYPE_MAX_LENGTH
    )

    if archetype is None:
        await update.message.reply_text(
            "Такой архетип не подойдёт: без команд, спецсимволов "
            f"и длиннее {JOKE_ARCHETYPE_MAX_LENGTH} символов."
        )
        return

    await set_member_joke_archetype(
        update.effective_chat.id,
        target_message.from_user.id,
        archetype,
        str(update.effective_chat.type),
    )

    display_name = (
        target_message.from_user.full_name
        or target_message.from_user.username
        or "участник"
    )

    await update.message.reply_text(
        f"Принято. Теперь {display_name} официально «{archetype}»."
    )


# ============================================================
# ПОДГОТОВКА СООБЩЕНИЙ В ЛИЧКЕ И ГРУППЕ
# ============================================================

async def get_bot_username(
    context: ContextTypes.DEFAULT_TYPE,
) -> str:
    """Возвращает username Telegram-бота без символа @."""

    username = context.bot.username

    if not username:
        bot_info = await context.bot.get_me()
        username = bot_info.username

    return username or ""


def is_reply_to_bot(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> bool:
    """Проверяет, ответил ли пользователь на сообщение бота."""

    if (
        not update.message
        or not update.message.reply_to_message
    ):
        return False

    sender = update.message.reply_to_message.from_user

    return bool(
        sender
        and sender.id == context.bot.id
    )

# ============================================================
# ТЕКСТОВЫЕ ОБРАЩЕНИЯ К ЯЙЦЕСЛАВУ В ГРУППЕ
# ============================================================

GROUP_ADDRESS_RE = re.compile(
    r"^\s*"
    r"(?:(?:"
    r"привет"
    r"|здарова"
    r"|здорово"
    r"|алло"
    r"|слушай"
    r"|эй"
    r")\s*[,.:;!?—-]*\s+)?"
    r"(?:"
    r"яйцеслав"
    r"|яйцеславыч"
    r"|яйцо"
    r"|боб[её]р"
    r"|бобр"
    r"|бобрище"
    r"|курва"
    r"|бот"
    r"|ассистент"
    r"|помощник"
    r"|профессор"
    r"|эксперт"
    r")\b"
    r"[\s,.:;!?—-]*",
    flags=re.IGNORECASE,
)

def extract_group_address(
    text: str,
) -> str | None:
    """
    Находит обращение к боту в начале сообщения
    и возвращает текст после обращения.
    """

    match = GROUP_ADDRESS_RE.match(
        text or ""
    )

    if not match:
        return None

    return text[
        match.end():
    ].strip()

async def prepare_request_text(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
    original_text: str | None,
    default_text: str,
) -> str | None:
    """
    В личном чате принимает все сообщения.

    В группе принимает сообщение только при упоминании
    @username бота или при ответе на сообщение бота.
    """

    if not update.effective_chat:
        return None

    text = (original_text or "").strip()
    chat_type = update.effective_chat.type

    # В личном чате обрабатываем любое сообщение
    if chat_type == ChatType.PRIVATE:
        return text or default_text

    # В группе и супергруппе требуется обращение к боту
    if chat_type in (
        ChatType.GROUP,
        ChatType.SUPERGROUP,
    ):
        bot_username = await get_bot_username(
            context
        )

        mention = (
            f"@{bot_username}"
            if bot_username
            else ""
        )

        has_mention = bool(
            mention
            and mention.lower() in text.lower()
        )

        has_reply = is_reply_to_bot(
            update,
            context,
        )

        addressed_text = extract_group_address(
            text
        )

        has_text_address = (
            addressed_text is not None
        )

        if (
            not has_mention
            and not has_reply
            and not has_text_address
        ):
            return None

        if has_mention:
            text = re.sub(
                re.escape(mention),
                "",
                text,
                flags=re.IGNORECASE,
            ).strip()

        elif has_text_address:
            text = addressed_text

        return text or default_text

    return None
# ============================================================
# РАСПОЗНАВАНИЕ ПРОСЬБЫ ОТВЕТИТЬ ГОЛОСОМ
# ============================================================

VOICE_REQUEST_PATTERNS = (
    r"\bответь\s+голосом\b",
    r"\bответ\s+голосом\b",
    r"\bскажи(?:\s+это)?\s+голосом\b",
    r"\bответь\s+аудио\b",
    r"\bответь\s+войсом\b",
    r"\bпришли\s+голосов(?:ое|ым)\b",
    r"\bозвучь(?:\s+ответ)?\b",
)


def text_requests_voice(
    text: str,
) -> bool:
    """Проверяет, попросил ли пользователь голосовой ответ."""

    if not text:
        return False

    return any(
        re.search(
            pattern,
            text,
            flags=re.IGNORECASE,
        )
        for pattern in VOICE_REQUEST_PATTERNS
    )


def remove_voice_request(
    text: str,
) -> str:
    """Удаляет из вопроса слова вроде «ответь голосом»."""

    cleaned = text

    for pattern in VOICE_REQUEST_PATTERNS:
        cleaned = re.sub(
            pattern,
            "",
            cleaned,
            flags=re.IGNORECASE,
        )

    cleaned = re.sub(
        r"\s{2,}",
        " ",
        cleaned,
    )

    return cleaned.strip(
        " ,.!?:;—-"
    )


def voice_mode_enabled(
    context: ContextTypes.DEFAULT_TYPE,
) -> bool:
    """Проверяет постоянный голосовой режим пользователя."""
    return bool(
        context.user_data.get(
            "voice_mode",
            False,
        )
    )


def build_private_answer_keyboard() -> InlineKeyboardMarkup:
    """Создаёт кнопки под ответом в личном чате."""

    return InlineKeyboardMarkup(
        [
            [
                InlineKeyboardButton(
                    "🔎 Поиск в интернете",
                    callback_data="answer_search",
                )
            ],
            [
                InlineKeyboardButton(
                    "📚 Подробнее",
                    callback_data="answer_more",
                )
            ],
            [
                InlineKeyboardButton(
                    "✂️ Короче",
                    callback_data="answer_shorter",
                )
            ],
            [
                InlineKeyboardButton(
                    "🔊 Голосом",
                    callback_data="answer_voice",
                )
            ],
        ]
    )   
# ============================================================
# ОТПРАВКА ТЕКСТОВЫХ И ГОЛОСОВЫХ ОТВЕТОВ
# ============================================================

async def send_voice_answer(
    update: Update,
    text: str,
) -> None:
    """
    Создаёт озвученный ответ.

    Сначала пробует edge-tts.
    Если edge-tts не работает, использует gTTS.
    """

    message = update.effective_message

    if (
        not message
        or not update.effective_chat
    ):
        return

    # Убираем ссылки и часть разметки,
    # чтобы синтезатор не зачитывал технический мусор.
    speech_text = re.sub(
        r"https?://\S+|www\.\S+",
        "",
        text or "",
        flags=re.IGNORECASE,
    )

    speech_text = re.sub(
        r"[*_`#>]",
        "",
        speech_text,
    )

    speech_text = re.sub(
        r"\s+",
        " ",
        speech_text,
    ).strip()

    speech_text = speech_text[:MAX_TTS_CHARS]

    if not speech_text:
        speech_text = (
            "Яйцеславу нечего озвучивать. "
            "Редкий анлак."
        )

    output_path = TEMP_DIR / (
        f"tts_"
        f"{update.effective_chat.id}_"
        f"{message.message_id}_"
        f"{uuid.uuid4().hex}.mp3"
    )

    edge_success = False
    last_edge_error: Exception | None = None

    try:
        # Три попытки качественной озвучки через edge-tts
        for attempt in range(1, 4):
            output_path.unlink(
                missing_ok=True
            )

            try:
                communicator = edge_tts.Communicate(
                    text=speech_text,
                    voice=TTS_VOICE,
                    rate=TTS_RATE,
                    pitch=TTS_PITCH,
                    volume=TTS_VOLUME,
                )

                await communicator.save(
                    str(output_path)
                )

                if (
                    output_path.exists()
                    and output_path.stat().st_size > 1000
                ):
                    edge_success = True
                    break

                raise RuntimeError(
                    "edge-tts создал пустой аудиофайл"
                )

            except Exception as error:
                last_edge_error = error

                logging.warning(
                    "Попытка edge-tts %s из 3 "
                    "завершилась ошибкой: %s",
                    attempt,
                    error,
                )

                if attempt < 3:
                    await asyncio.sleep(
                        attempt * 2
                    )

        # Запасная озвучка через gTTS
        if not edge_success:
            logging.warning(
                "edge-tts не сработал. "
                "Переключаюсь на gTTS. "
                "Последняя ошибка: %s",
                last_edge_error,
            )

            output_path.unlink(
                missing_ok=True
            )

            fallback_tts = gTTS(
                text=speech_text,
                lang="ru",
                slow=False,
            )

            await asyncio.to_thread(
                fallback_tts.save,
                str(output_path),
            )

        if (
            not output_path.exists()
            or output_path.stat().st_size <= 1000
        ):
            raise RuntimeError(
                "Не удалось создать аудиофайл"
            )

        try:
            # Сначала отправляем как настоящую голосовуху
            with output_path.open("rb") as audio_file:
                await message.reply_voice(
                    voice=audio_file,
                    filename="yayceslav.mp3",
                )

        except BadRequest as error:
            # Если пользователь запретил голосовые сообщения,
            # отправляем тот же файл как обычную аудиозапись.
            if (
                "Voice messages forbidden"
                not in str(error)
            ):
                raise

            with output_path.open("rb") as audio_file:
                await message.reply_audio(
                    audio=audio_file,
                    filename="yayceslav.mp3",
                    title="Ответ Яйцеслава",
                    performer="Яйцеслав",
                )

    finally:
        output_path.unlink(
            missing_ok=True
        )


async def send_answer(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
    text: str,
    force_voice: bool = False,
    show_buttons: bool = False,
    source_user_text: str | None = None,
    disable_voice: bool = False,
) -> None:
    """Отправляет voice/text; в групповой болтовне иногда делает ответ человечнее."""

    message = update.effective_message
    if not message:
        return

    answer_text = (text or "").strip()
    if not answer_text:
        answer_text = "Яйцеслав задумался и ничего не изрёк. Редкий анлак."

    # disable_voice is an explicit per-call override (e.g. a random 50/50
    # roll for video-circle replies) and wins over both force_voice and the
    # user's own voice_mode setting.
    use_voice = not disable_voice and (force_voice or voice_mode_enabled(context))
    if use_voice:
        try:
            await send_voice_answer(update, answer_text)
            return
        except Exception as error:
            logging.exception("Ошибка голосового ответа: %s", error)
            await message.reply_text("Голосовой тракт охрип. Держи ответ текстом.")

    trace = feedback_engine.get_current_trace()
    current_hostile_streak = 0
    if (
        trace is not None
        and getattr(trace, "conversation_mode", "") == "hostile"
        and update.effective_chat
        and update.effective_user
        and update.effective_chat.type in (ChatType.GROUP, ChatType.SUPERGROUP)
    ):
        current_hostile_streak = hostile_streak_engine.current(
            update.effective_chat.id,
            update.effective_user.id,
        )

    if source_user_text is None:
        plan = humanizer_engine.HumanizedReply((answer_text,), (0.0,))
    else:
        plan = humanizer_engine.humanize_reply(
            answer_text,
            user_text=source_user_text,
            trace=trace,
            hostile_streak=current_hostile_streak,
        )

    for message_index, planned_text in enumerate(plan.messages):
        delay = plan.delays[message_index] if message_index < len(plan.delays) else 0.0
        if delay > 0:
            await asyncio.sleep(delay)

        for position in range(0, len(planned_text), 4000):
            is_last_chunk = position + 4000 >= len(planned_text)
            is_last_planned = message_index == len(plan.messages) - 1
            reply_markup = None
            if (
                show_buttons
                and update.effective_chat
                and update.effective_chat.type == ChatType.PRIVATE
                and is_last_chunk
                and is_last_planned
            ):
                reply_markup = build_private_answer_keyboard()

            sent_message = await message.reply_text(
                planned_text[position:position + 4000],
                reply_markup=reply_markup,
            )

            is_typo_correction = (
                plan.effect == "typo_correction" and message_index == 1
            )
            if (
                trace is not None
                and update.effective_chat
                and update.effective_chat.type in (ChatType.GROUP, ChatType.SUPERGROUP)
                and not is_typo_correction
            ):
                await asyncio.to_thread(
                    store_bot_response_feedback_sync,
                    update.effective_chat.id,
                    sent_message.message_id,
                    trace,
                )


async def answer_button_callback(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """Обрабатывает кнопки под ответом в личном чате."""

    query = update.callback_query

    if not query:
        return

    if (
        not update.effective_chat
        or update.effective_chat.type
        != ChatType.PRIVATE
    ):
        await query.answer()
        return

    user_query = str(
        context.user_data.get(
            "last_user_query",
            "",
        )
    ).strip()

    last_answer = str(
        context.user_data.get(
            "last_answer",
            "",
        )
    ).strip()

    if not user_query or not last_answer:
        await query.answer(
            "Этот ответ уже устарел. "
            "Задай вопрос ещё раз.",
            show_alert=True,
        )
        return

    action = query.data or ""

    allowed_actions = (
        "answer_search",
        "answer_more",
        "answer_shorter",
        "answer_voice",
    )

    if action not in allowed_actions:
        await query.answer(
            "Неизвестная кнопка.",
            show_alert=True,
        )
        return

    await query.answer()

    # Поиск исходного вопроса в интернете
    if action == "answer_search":
        await perform_web_search(
            update=update,
            context=context,
            query=user_query,
        )
        return

    # Озвучка текущего ответа
    if action == "answer_voice":
        try:
            await send_voice_answer(
                update,
                last_answer,
            )
        except Exception as error:
            logging.exception(
                "Ошибка озвучки кнопкой: %s",
                error,
            )

            message = update.effective_message

            if message:
                await message.reply_text(
                    "Не удалось озвучить ответ."
                )

        return

    if action == "answer_more":
        prompt = f"""
Пользователь ранее задал вопрос:

{user_query}

Текущий ответ:

{last_answer}

Раскрой текущий ответ подробнее.
Добавь полезные объяснения, примеры и важные детали.
Не повторяй длинное вступление.
Не упоминай, что ты переписываешь предыдущий ответ.
""".strip()

        max_tokens = 650

    else:
        prompt = f"""
Пользователь ранее задал вопрос:

{user_query}

Текущий ответ:

{last_answer}

Сократи текущий ответ.
Оставь только главное и сохрани важные факты.
Ответ должен состоять примерно из двух–четырёх
коротких предложений.
Не упоминай, что ты сокращаешь предыдущий ответ.
""".strip()

        max_tokens = 220

    await context.bot.send_chat_action(
        chat_id=update.effective_chat.id,
        action=ChatAction.TYPING,
    )

    try:
        user_settings = None

        if update.effective_user:
            user_settings = await get_user_settings(
                update.effective_user.id
            )

        new_answer = await ask_gemini(
            contents=prompt,
            max_output_tokens=max_tokens,
            user_settings=user_settings,
            chat_id=(
                update.effective_chat.id
                if update.effective_chat
                else None
            ),
            chat_type=(
                str(update.effective_chat.type)
                if update.effective_chat
                else "private"
            ),
            user_id=(
                update.effective_user.id
                if update.effective_user
                else None
            ),
        )

        context.user_data[
            "last_answer"
        ] = new_answer

        await send_answer(
            update,
            context,
            new_answer,
            force_voice=bool(
                user_settings
                and user_settings.get(
                    "voice_enabled",
                    False,
                )
            ),
            show_buttons=True,
        )

        await increment_stat(
            "bot_answers"
        )

    except Exception as error:
        logging.exception(
            "Ошибка кнопки ответа: %s",
            error,
        )

        message = update.effective_message

        if message:
            await message.reply_text(
                "Не удалось изменить ответ. "
                "Нейронка опять поплыла."
            )
# ============================================================
# ТЕКСТОВЫЕ СООБЩЕНИЯ
# ============================================================

async def answer_text_message(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """Отвечает на текстовые сообщения."""

    if (
        not update.message
        or not update.effective_chat
    ):
        return

    user_text = await prepare_request_text(
        update=update,
        context=context,
        original_text=update.message.text,
        default_text=(
            "Пользователь просто позвал тебя. "
            "Коротко спроси, чего ему надо."
        ),
    )

    if user_text is None:
        return

    # Защита от слишком длинного текста
    if len(user_text) > MAX_USER_TEXT_CHARS:
        await update.message.reply_text(
            "Слишком много текста, гигант мысли. "
            "Сократи запрос до 3000 символов."
        )
        return

    # Лимит обычных текстовых запросов
    if not await enforce_rate_limit(
        update,
        "general",
    ):
        return

    # Проверяем просьбу ответить голосом
    force_voice = text_requests_voice(
        user_text
    )

    if force_voice:
        user_text = remove_voice_request(
            user_text
        )

        if not user_text:
            user_text = (
                "Коротко спроси пользователя, "
                "что именно ему озвучить."
            )

    # Загружаем персональные настройки
    user_settings = None

    if update.effective_user:
        user_settings = await get_user_settings(
            update.effective_user.id
        )

    search_mode = str(
        (
            user_settings
            or DEFAULT_USER_SETTINGS
        ).get(
            "search_mode",
            "button",
        )
    )

    # Проверяем явную просьбу выполнить поиск
    search_query = extract_search_query(
        user_text
    )

    # В автоматическом режиме сами включаем поиск — но только если
    # сообщение реально похоже на информационный запрос, а не на
    # реплику про самого бота (шутка/подкол/обещание/оскорбление).
    if (
        search_query is None
        and search_mode == "auto"
        and not is_conversation_about_bot(user_text)
        and should_auto_search(user_text)
    ):
        search_query = user_text

    if search_query is not None:
        await perform_web_search(
            update=update,
            context=context,
            query=search_query,
            force_voice=force_voice,
        )
        return
    # Учитываем обычный текстовый запрос
    await register_user_and_chat(
        update
    )

    await increment_stat(
        "total_requests"
    )

    await increment_stat(
        "text_requests"
    )
    use_voice_style = (
        force_voice
        or voice_mode_enabled(context)
    )

    await context.bot.send_chat_action(
        chat_id=update.effective_chat.id,
        action=ChatAction.TYPING,
    )

    try:
        request_for_gemini = user_text

        settings_voice_enabled = bool(
            user_settings
            and user_settings.get(
                "voice_enabled",
                False,
            )
        )

        use_voice_style = (
            use_voice_style
            or settings_voice_enabled
        )
        
        private_user_id: int | None = None
        group_chat_id: int | None = None
        group_author_name = ""
        request_user_name = ""
        recent_messages_list: list[str] = []

        # Личная переписка: память текущей задачи 15 минут
        if (
            update.effective_chat.type == ChatType.PRIVATE
            and update.effective_user
        ):
            private_user_id = update.effective_user.id

            request_user_name = (
                update.effective_user.full_name
                or update.effective_user.username
                or ""
            )

            record_last_user_message(
                update.effective_chat.id,
                private_user_id,
                user_text,
            )

            previous_context = build_memory_context(
                PRIVATE_MEMORY,
                private_user_id,
                PRIVATE_MEMORY_SECONDS,
            )

            if previous_context:
                recent_messages_list = previous_context.splitlines()

                request_for_gemini = (
                    "Ниже находится история текущей задачи "
                    "пользователя. Учитывай её при ответе, "
                    "но не пересказывай без необходимости.\n\n"
                    f"{previous_context}\n\n"
                    f"Новое сообщение пользователя:\n"
                    f"{user_text}"
                )

            # Текущее сообщение фиксируем ДО ожидания Gemini. Тогда
            # следующий concurrent update этого пользователя уже увидит его
            # в контексте, даже если первый ответ ещё генерируется.
            remember_message(
                PRIVATE_MEMORY,
                private_user_id,
                "user",
                user_text,
                PRIVATE_MEMORY_SECONDS,
                PRIVATE_MEMORY_MAX_MESSAGES,
            )

        # Группа: память разговора за последние пять минут
        elif (
            update.effective_chat.type in (
                ChatType.GROUP,
                ChatType.SUPERGROUP,
            )
            and update.effective_user
        ):
            group_chat_id = update.effective_chat.id

            # Кто-то напрямую обратился к боту — значит, предыдущие
            # случайные вмешательства не были полным игнором.
            register_group_engagement(group_chat_id)

            group_author_name = (
                update.effective_user.full_name
                or update.effective_user.username
                or "Участник"
            )
            request_user_name = group_author_name

            await touch_member_profile(
                group_chat_id,
                update.effective_user.id,
                str(update.effective_chat.type),
                group_author_name,
                update.effective_user.username,
            )

            await increment_chat_activity(
                group_chat_id,
                update.effective_user.id,
                str(update.effective_chat.type),
                current_msk_date_str(),
                **build_text_activity_deltas(
                    user_text, is_reply_to_bot=True
                ),
            )

            record_last_user_message(
                group_chat_id,
                update.effective_user.id,
                user_text,
            )

            previous_context = build_memory_context(
                GROUP_MEMORY,
                group_chat_id,
                GROUP_MEMORY_SECONDS,
            )

            if previous_context:
                recent_messages_list = previous_context.splitlines()

                request_for_gemini = (
                    "Ниже приведена переписка группы "
                    "за последние пять минут. "
                    "Используй её только как контекст. "
                    "Не пересказывай всю переписку и не утверждай, "
                    "что сообщения адресовались тебе.\n\n"
                    f"{previous_context}\n\n"
                    f"Новое обращение к тебе от "
                    f"{group_author_name}:\n"
                    f"{user_text}"
                )

            # Как и в личке: пользовательское сообщение входит в
            # память до сетевого await. Это убирает потерю контекста при
            # concurrent_updates(8).
            remember_message(
                GROUP_MEMORY,
                group_chat_id,
                "user",
                user_text,
                GROUP_MEMORY_SECONDS,
                GROUP_MEMORY_MAX_MESSAGES,
                group_author_name,
            )

        answer = await ask_gemini(
            contents=request_for_gemini,
            max_output_tokens=get_response_token_limit(
                user_settings,
                normal_tokens=360,
            ),
            voice_style=use_voice_style,
            user_settings=user_settings,
            chat_id=update.effective_chat.id,
            chat_type=str(update.effective_chat.type),
            user_name=request_user_name,
            recent_messages=recent_messages_list,
            bot_was_mentioned=True,
            thinking_level=thinking_engine.choose_thinking_level(
                user_text
            ),
            user_id=(
                update.effective_user.id
                if update.effective_user
                else None
            ),
        )

        # Сохраняем вопрос и ответ в памяти лички
        if private_user_id is not None:
            context.user_data[
                "last_user_query"
            ] = user_text

            context.user_data[
                "last_answer"
            ] = answer
            
            remember_message(
                PRIVATE_MEMORY,
                private_user_id,
                "assistant",
                answer,
                PRIVATE_MEMORY_SECONDS,
                PRIVATE_MEMORY_MAX_MESSAGES,
            )

        # Сохраняем обращение и ответ в памяти группы
        if group_chat_id is not None:
            remember_message(
                GROUP_MEMORY,
                group_chat_id,
                "assistant",
                answer,
                GROUP_MEMORY_SECONDS,
                GROUP_MEMORY_MAX_MESSAGES,
            )

        await send_answer(
            update,
            context,
            answer,
            force_voice=(
                force_voice
                or settings_voice_enabled
            ),
            show_buttons=True,
            source_user_text=user_text,
        )
        
        await increment_stat(
            "bot_answers"
        )    
    except Exception as error:
        logging.exception(
            "Ошибка текстового запроса: %s",
            error,
        )

        await update.message.reply_text(
            "Связь с нейросетью опять пала в бою 🥚\n"
            "Повтори позже, нищий."
        )

# ============================================================
# ФОТОГРАФИИ
# ============================================================

async def answer_photo(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """Анализирует фотографию."""

    if (
        not update.message
        or not update.message.photo
        or not update.effective_chat
    ):
        return

    prompt = await prepare_request_text(
        update=update,
        context=context,
        original_text=update.message.caption,
        default_text=(
            "Коротко опиши, что изображено "
            "на фотографии. "
            "Если это мем — объясни шутку. "
            "Если это задача или текст — "
            "прочитай и помоги."
        ),
    )

    if prompt is None:
        return

    if not await enforce_rate_limit(
        update,
        "media",
    ):
        return
    # Учитываем фотографию
    await register_user_and_chat(
        update
    )

    await increment_stat(
        "total_requests"
    )

    await increment_stat(
        "photo_requests"
    )
    force_voice = text_requests_voice(
        prompt
    )

    if force_voice:
        prompt = (
            remove_voice_request(prompt)
            or "Коротко опиши изображение."
        )

    user_settings = None

    if update.effective_user:
        user_settings = await get_user_settings(
            update.effective_user.id
        )

    settings_voice_enabled = bool(
        user_settings
        and user_settings.get(
            "voice_enabled",
            False,
        )
    )

    use_voice_style = (
        force_voice
        or voice_mode_enabled(context)
        or settings_voice_enabled
    )

    file_path = TEMP_DIR / (
        f"photo_"
        f"{update.effective_chat.id}_"
        f"{update.message.message_id}_"
        f"{uuid.uuid4().hex}.jpg"
    )

    await context.bot.send_chat_action(
        chat_id=update.effective_chat.id,
        action=ChatAction.TYPING,
    )

    try:
        photo = update.message.photo[-1]

        telegram_file = await photo.get_file()

        await telegram_file.download_to_drive(
            custom_path=str(file_path)
        )
        answer = await ask_gemini(
            contents=[
                types.Part.from_bytes(
                    data=file_path.read_bytes(),
                    mime_type="image/jpeg",
                ),
                prompt,
            ],
            max_output_tokens=get_response_token_limit(
                user_settings,
                normal_tokens=250,
            ),
            voice_style=use_voice_style,
            user_settings=user_settings,
            chat_id=(
                update.effective_chat.id
                if update.effective_chat
                else None
            ),
            chat_type=(
                str(update.effective_chat.type)
                if update.effective_chat
                else "private"
            ),
            user_id=(
                update.effective_user.id
                if update.effective_user
                else None
            ),
        )

        # Запоминаем обсуждение фотографии
        if update.effective_user:
            memory_text = (
                f"[Пользователь прислал фотографию] {prompt}"
            )

            if (
                update.effective_chat.type
                == ChatType.PRIVATE
            ):
                remember_message(
                    PRIVATE_MEMORY,
                    update.effective_user.id,
                    "user",
                    memory_text,
                    PRIVATE_MEMORY_SECONDS,
                    PRIVATE_MEMORY_MAX_MESSAGES,
                )

                remember_message(
                    PRIVATE_MEMORY,
                    update.effective_user.id,
                    "assistant",
                    answer,
                    PRIVATE_MEMORY_SECONDS,
                    PRIVATE_MEMORY_MAX_MESSAGES,
                )

            elif update.effective_chat.type in (
                ChatType.GROUP,
                ChatType.SUPERGROUP,
            ):
                author_name = (
                    update.effective_user.full_name
                    or update.effective_user.username
                    or "Участник"
                )

                remember_message(
                    GROUP_MEMORY,
                    update.effective_chat.id,
                    "user",
                    memory_text,
                    GROUP_MEMORY_SECONDS,
                    GROUP_MEMORY_MAX_MESSAGES,
                    author_name,
                )

                remember_message(
                    GROUP_MEMORY,
                    update.effective_chat.id,
                    "assistant",
                    answer,
                    GROUP_MEMORY_SECONDS,
                    GROUP_MEMORY_MAX_MESSAGES,
                )
       
        await send_answer(
            update,
            context,
            answer,
            force_voice=(
                force_voice
                or settings_voice_enabled
            ),
        )
        
        await increment_stat(
            "bot_answers"
        )
        
    except Exception as error:
        logging.exception(
            "Ошибка анализа фотографии: %s",
            error,
        )

        await update.message.reply_text(
            "Картинку не разобрал. "
            "Связь поплыла или фото кривое."
        )

    finally:
        file_path.unlink(
            missing_ok=True
        )


# ============================================================
# ДОКУМЕНТЫ
# ============================================================

async def answer_document(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """
    Анализирует PDF, DOCX, XLSX, CSV, TXT
    и изображения, отправленные как файл.
    """

    if (
        not update.message
        or not update.message.document
        or not update.effective_chat
    ):
        return

    document = update.message.document

    prompt = await prepare_request_text(
        update=update,
        context=context,
        original_text=update.message.caption,
        default_text=(
            "Коротко проанализируй файл. "
            "Скажи, о чём он, и выдели главное."
        ),
    )

    if prompt is None:
        return
    if not await enforce_rate_limit(
        update,
        "media",
    ):
        return
    caption_text = update.message.caption or ""

    user_settings = None

    if update.effective_user:
        user_settings = await get_user_settings(
            update.effective_user.id
        )

    settings_voice_enabled = bool(
        user_settings
        and user_settings.get(
            "voice_enabled",
            False,
        )
    )

    force_voice = (
        text_requests_voice(caption_text)
        or voice_mode_enabled(context)
        or settings_voice_enabled
    )

    use_voice_style = force_voice

    if force_voice:
        prompt = (
            remove_voice_request(prompt)
            or "Коротко проанализируй файл."
        )
   
    if (
        document.file_size
        and document.file_size > MAX_FILE_SIZE
    ):
        await update.message.reply_text(
            "Файл тяжелее 20 МБ. "
            "Этот кирпич я не потащу."
        )
        return
    # Учитываем документ
    await register_user_and_chat(
        update
    )

    await increment_stat(
        "total_requests"
    )

    await increment_stat(
        "document_requests"
    )
    original_filename = (
        document.file_name
        or f"document_{update.message.message_id}"
    )

    file_path = TEMP_DIR / make_safe_filename(
        original_filename,
        update.message.message_id,
        update.effective_chat.id,
    )

    await context.bot.send_chat_action(
        chat_id=update.effective_chat.id,
        action=ChatAction.TYPING,
    )

    user_settings = None

    if update.effective_user:
        user_settings = await get_user_settings(
            update.effective_user.id
        )

    try:
        telegram_file = await document.get_file()

        await telegram_file.download_to_drive(
            custom_path=str(file_path)
        )

        extension = file_path.suffix.lower()

        mime_type = (
            document.mime_type
            or mimetypes.guess_type(
                original_filename
            )[0]
            or "application/octet-stream"
        )

        # PDF
        if (
            extension == ".pdf"
            or mime_type == "application/pdf"
        ):
            answer = await ask_gemini(
                contents=[
                    types.Part.from_bytes(
                        data=file_path.read_bytes(),
                        mime_type="application/pdf",
                    ),
                    (
                        f"Название файла: "
                        f"{original_filename}\n"
                        f"Запрос пользователя: {prompt}"
                    ),
                ],
                max_output_tokens=get_response_token_limit(
                    user_settings,
                    normal_tokens=500,
                ),
                voice_style=use_voice_style,
                user_settings=user_settings,
                chat_id=(
                    update.effective_chat.id
                    if update.effective_chat
                    else None
                ),
                chat_type=(
                    str(update.effective_chat.type)
                    if update.effective_chat
                    else "private"
                ),
                user_id=(
                    update.effective_user.id
                    if update.effective_user
                    else None
                ),
            )

        # Изображение, отправленное как документ
        elif extension in {
            ".jpg",
            ".jpeg",
            ".png",
            ".webp",
            ".heic",
            ".heif",
        }:
            detected_mime = (
                mime_type
                if mime_type.startswith("image/")
                else mimetypes.guess_type(
                    original_filename
                )[0]
            ) or "image/jpeg"

            answer = await ask_gemini(
                contents=[
                    types.Part.from_bytes(
                        data=file_path.read_bytes(),
                        mime_type=detected_mime,
                    ),
                    (
                        f"Название файла: "
                        f"{original_filename}\n"
                        f"Запрос пользователя: {prompt}"
                    ),
                ],
                max_output_tokens=get_response_token_limit(
                    user_settings,
                    normal_tokens=300,
                ),
                voice_style=use_voice_style,
                user_settings=user_settings,
                chat_id=(
                    update.effective_chat.id
                    if update.effective_chat
                    else None
                ),
                chat_type=(
                    str(update.effective_chat.type)
                    if update.effective_chat
                    else "private"
                ),
                user_id=(
                    update.effective_user.id
                    if update.effective_user
                    else None
                ),
            )

        # Word
        elif extension == ".docx":
            extracted_text = read_docx_file(
                file_path
            )

            if not extracted_text:
                raise ValueError(
                    "В DOCX не найден текст"
                )

            answer = await ask_gemini(
                contents=(
                    f"Название файла: "
                    f"{original_filename}\n\n"
                    f"Запрос пользователя:\n"
                    f"{prompt}\n\n"
                    f"Содержимое DOCX:\n"
                    f"{extracted_text}"
                ),
                max_output_tokens=get_response_token_limit(
                    user_settings,
                    normal_tokens=500,
                ),
                voice_style=use_voice_style,
                user_settings=user_settings,
                chat_id=(
                    update.effective_chat.id
                    if update.effective_chat
                    else None
                ),
                chat_type=(
                    str(update.effective_chat.type)
                    if update.effective_chat
                    else "private"
                ),
                user_id=(
                    update.effective_user.id
                    if update.effective_user
                    else None
                ),
            )

        # Excel
        elif extension == ".xlsx":
            extracted_text = read_xlsx_file(
                file_path
            )

            if not extracted_text:
                raise ValueError(
                    "В XLSX не найдены данные"
                )

            answer = await ask_gemini(
                contents=(
                    f"Название файла: "
                    f"{original_filename}\n\n"
                    f"Запрос пользователя:\n"
                    f"{prompt}\n\n"
                    f"Данные XLSX:\n"
                    f"{extracted_text}"
                ),
                max_output_tokens=get_response_token_limit(
                    user_settings,
                    normal_tokens=500,
                ),
                voice_style=use_voice_style,
                user_settings=user_settings,
                chat_id=(
                    update.effective_chat.id
                    if update.effective_chat
                    else None
                ),
                chat_type=(
                    str(update.effective_chat.type)
                    if update.effective_chat
                    else "private"
                ),
                user_id=(
                    update.effective_user.id
                    if update.effective_user
                    else None
                ),
            )

        # CSV
        elif extension == ".csv":
            extracted_text = read_csv_file(
                file_path
            )

            answer = await ask_gemini(
                contents=(
                    f"Название файла: "
                    f"{original_filename}\n\n"
                    f"Запрос пользователя:\n"
                    f"{prompt}\n\n"
                    f"Данные CSV:\n"
                    f"{extracted_text}"
                ),
                max_output_tokens=get_response_token_limit(
                    user_settings,
                    normal_tokens=500,
                ),
                voice_style=use_voice_style,
                user_settings=user_settings,
                chat_id=(
                    update.effective_chat.id
                    if update.effective_chat
                    else None
                ),
                chat_type=(
                    str(update.effective_chat.type)
                    if update.effective_chat
                    else "private"
                ),
                user_id=(
                    update.effective_user.id
                    if update.effective_user
                    else None
                ),
            )

        # Текстовые файлы
        elif extension in {
            ".txt",
            ".md",
            ".json",
            ".log",
            ".py",
            ".xml",
            ".html",
            ".htm",
        }:
            extracted_text = read_text_file(
                file_path
            )

            answer = await ask_gemini(
                contents=(
                    f"Название файла: "
                    f"{original_filename}\n\n"
                    f"Запрос пользователя:\n"
                    f"{prompt}\n\n"
                    f"Содержимое файла:\n"
                    f"{extracted_text}"
                ),
                max_output_tokens=get_response_token_limit(
                    user_settings,
                    normal_tokens=500,
                ),
                voice_style=use_voice_style,
                user_settings=user_settings,
                chat_id=(
                    update.effective_chat.id
                    if update.effective_chat
                    else None
                ),
                chat_type=(
                    str(update.effective_chat.type)
                    if update.effective_chat
                    else "private"
                ),
                user_id=(
                    update.effective_user.id
                    if update.effective_user
                    else None
                ),
            )

        else:
            await update.message.reply_text(
                "Формат пока не поддерживаю. "
                "Кидай PDF, DOCX, XLSX, CSV, "
                "TXT или картинку."
            )
            return
        # Запоминаем обсуждение документа
        if update.effective_user:
            memory_text = (
                f"[Пользователь прислал файл "
                f"«{original_filename}»] {prompt}"
            )

            if (
                update.effective_chat.type
                == ChatType.PRIVATE
            ):
                remember_message(
                    PRIVATE_MEMORY,
                    update.effective_user.id,
                    "user",
                    memory_text,
                    PRIVATE_MEMORY_SECONDS,
                    PRIVATE_MEMORY_MAX_MESSAGES,
                )

                remember_message(
                    PRIVATE_MEMORY,
                    update.effective_user.id,
                    "assistant",
                    answer,
                    PRIVATE_MEMORY_SECONDS,
                    PRIVATE_MEMORY_MAX_MESSAGES,
                )

            elif update.effective_chat.type in (
                ChatType.GROUP,
                ChatType.SUPERGROUP,
            ):
                author_name = (
                    update.effective_user.full_name
                    or update.effective_user.username
                    or "Участник"
                )

                remember_message(
                    GROUP_MEMORY,
                    update.effective_chat.id,
                    "user",
                    memory_text,
                    GROUP_MEMORY_SECONDS,
                    GROUP_MEMORY_MAX_MESSAGES,
                    author_name,
                )

                remember_message(
                    GROUP_MEMORY,
                    update.effective_chat.id,
                    "assistant",
                    answer,
                    GROUP_MEMORY_SECONDS,
                    GROUP_MEMORY_MAX_MESSAGES,
                )
        await send_answer(
            update,
            context,
            answer,
            force_voice=force_voice,
        )
        
        await increment_stat(
            "bot_answers"
        )
        
    except Exception as error:
        logging.exception(
            "Ошибка анализа документа: %s",
            error,
        )

        await update.message.reply_text(
            "Файл не осилился. "
            "Формат кривой или связь обосралась."
        )

    finally:
        file_path.unlink(
            missing_ok=True
        )


# ============================================================
# ВХОДЯЩИЕ ГОЛОСОВЫЕ И АУДИО
# ============================================================

# Moderate, deliberately not tiny and not huge: a video circle nobody
# addressed to the bot still has a real but bounded chance of getting an
# unprompted reaction, so the bot can "butt into" a conversation without
# turning into spam. Shares the same chat-wide random-intervention budget
# (group_random_reply_allowed/record_group_random_reply) as the existing
# hard-mode text random-reply system, so the two don't compound.
VIDEO_NOTE_PROACTIVE_COMMENT_CHANCE = 0.20


async def answer_voice_or_audio(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """Понимает голосовые сообщения, аудиофайлы и видео-кружочки."""

    if (
        not update.message
        or not update.effective_chat
    ):
        return

    voice = update.message.voice
    audio = update.message.audio
    video_note = update.message.video_note

    media = voice or audio or video_note

    if media is None:
        return

    # В группе голосовое/аудио обрабатывается только как ответ на
    # сообщение Яйцеслава. Видео-кружок — исключение: даже не будучи
    # адресован боту, он может (редко, с оглядкой на общий лимит
    # случайных вмешательств в чате) получить проактивный комментарий.
    proactive_comment = False
    if update.effective_chat.type in (
        ChatType.GROUP,
        ChatType.SUPERGROUP,
    ):
        if not is_reply_to_bot(
            update,
            context,
        ):
            now = time.monotonic()
            if (
                video_note
                and group_random_reply_allowed(update.effective_chat.id, now)
                and random.random() < VIDEO_NOTE_PROACTIVE_COMMENT_CHANCE
            ):
                proactive_comment = True
                record_group_random_reply(update.effective_chat.id, now)
            else:
                return

    if not await enforce_rate_limit(
        update,
        "media",
    ):
        return

    file_size = getattr(
        media,
        "file_size",
        None,
    )

    if (
        file_size
        and file_size > MAX_FILE_SIZE
    ):
        await update.message.reply_text(
            "Это тяжелее 20 МБ. "
            "Это уже подкаст, а не голосовуха."
        )
        return
    # Учитываем голосовое сообщение, аудиофайл или видео-кружок
    await register_user_and_chat(
        update
    )

    await increment_stat(
        "total_requests"
    )

    await increment_stat(
        "voice_requests"
    )
    if voice:
        mime_type = (
            voice.mime_type
            or "audio/ogg"
        )
        suffix = ".ogg"

    elif audio:
        mime_type = (
            audio.mime_type
            or "audio/mpeg"
        )

        suffix = (
            Path(
                audio.file_name
                or "audio.mp3"
            ).suffix
            or ".mp3"
        )

    else:
        # Telegram video notes (video circles) have no mime_type field of
        # their own; Telegram always encodes them as mp4.
        mime_type = "video/mp4"
        suffix = ".mp4"

    file_path = TEMP_DIR / (
        f"audio_"
        f"{update.effective_chat.id}_"
        f"{update.message.message_id}_"
        f"{uuid.uuid4().hex}"
        f"{suffix}"
    )

    await context.bot.send_chat_action(
        chat_id=update.effective_chat.id,
        action=ChatAction.TYPING,
    )

    user_settings = None

    if update.effective_user:
        user_settings = await get_user_settings(
            update.effective_user.id
        )

    try:
        telegram_file = await media.get_file()

        await telegram_file.download_to_drive(
            custom_path=str(file_path)
        )

        if proactive_comment:
            prompt = (
                "Тебя никто не звал и не спрашивал — ты сам решил "
                "вклиниться в чужой разговор. Посмотри видео-кружок "
                "(там может рассказываться история, показываться что-то "
                "вроде котят и т.п.) и коротко вставь СВОЁ мнение или "
                "комментарий по содержанию, как будто сам встрял в беседу. "
                "Это не ответ на вопрос: не задавай встречных вопросов, "
                "не веди себя как ассистент. Одна короткая реплика "
                "в характере Яйцеслава."
            )
        else:
            prompt = (
                "Прослушай сообщение пользователя. "
                "Пойми вопрос и коротко ответь "
                "по существу в характере Яйцеслава. "
                "Полную расшифровку не делай, "
                "если её не просили."
            )

        answer = await ask_gemini(
            contents=[
                types.Part.from_bytes(
                    data=file_path.read_bytes(),
                    mime_type=mime_type,
                ),
                prompt,
            ],
            max_output_tokens=get_response_token_limit(
                user_settings,
                normal_tokens=320,
            ),
            voice_style=True,
            user_settings=user_settings,
            chat_id=(
                update.effective_chat.id
                if update.effective_chat
                else None
            ),
            chat_type=(
                str(update.effective_chat.type)
                if update.effective_chat
                else "private"
            ),
            user_id=(
                update.effective_user.id
                if update.effective_user
                else None
            ),
        )
        # Запоминаем обсуждение голосового сообщения или видео-кружка
        if update.effective_user:
            memory_text = (
                "[Пользователь отправил видео-кружок]"
                if video_note
                else "[Пользователь отправил голосовое сообщение]"
            )

            if (
                update.effective_chat.type
                == ChatType.PRIVATE
            ):
                remember_message(
                    PRIVATE_MEMORY,
                    update.effective_user.id,
                    "user",
                    memory_text,
                    PRIVATE_MEMORY_SECONDS,
                    PRIVATE_MEMORY_MAX_MESSAGES,
                )

                remember_message(
                    PRIVATE_MEMORY,
                    update.effective_user.id,
                    "assistant",
                    answer,
                    PRIVATE_MEMORY_SECONDS,
                    PRIVATE_MEMORY_MAX_MESSAGES,
                )

            elif update.effective_chat.type in (
                ChatType.GROUP,
                ChatType.SUPERGROUP,
            ):
                author_name = (
                    update.effective_user.full_name
                    or update.effective_user.username
                    or "Участник"
                )

                remember_message(
                    GROUP_MEMORY,
                    update.effective_chat.id,
                    "user",
                    memory_text,
                    GROUP_MEMORY_SECONDS,
                    GROUP_MEMORY_MAX_MESSAGES,
                    author_name,
                )

                remember_message(
                    GROUP_MEMORY,
                    update.effective_chat.id,
                    "assistant",
                    answer,
                    GROUP_MEMORY_SECONDS,
                    GROUP_MEMORY_MAX_MESSAGES,
                )
        if proactive_comment:
            # An unprompted comment stays text-only; a random voice
            # message interrupting someone else's conversation would be
            # more intrusive than a random text reply.
            await send_answer(
                update,
                context,
                answer,
                disable_voice=True,
            )
        else:
            # Voice, audio and video-circle replies are all a 50/50 coin
            # flip between voice and text.
            reply_as_voice = random.random() < 0.5

            await send_answer(
                update,
                context,
                answer,
                force_voice=reply_as_voice,
                disable_voice=not reply_as_voice,
            )

        await increment_stat(
            "bot_answers"
        )
        
    except Exception as error:
        logging.exception(
            "Ошибка обработки голосового/видео-кружка: %s",
            error,
        )

        await update.message.reply_text(
            "Кружок/голосовуху не разобрал. "
            "Либо связь поплыла, либо ты бубнишь."
        )

    finally:
        file_path.unlink(
            missing_ok=True
        )



async def message_reaction_feedback_handler(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """Учит вкус конкретного чата по реакциям людей на сообщения Яйцеслава."""

    del context
    reaction = update.message_reaction
    if reaction is None:
        return

    score_delta, count_delta = feedback_engine.reaction_delta(
        reaction.old_reaction,
        reaction.new_reaction,
    )
    if score_delta == 0 and count_delta == 0:
        return

    updated = await asyncio.to_thread(
        apply_bot_reaction_delta_sync,
        reaction.chat.id,
        reaction.message_id,
        score_delta,
        count_delta,
    )
    if updated:
        adaptation_cache.invalidate("feedback", reaction.chat.id)
        logging.info(
            "Reaction feedback matched chat=%s message=%s score_delta=%.2f count_delta=%s",
            reaction.chat.id,
            reaction.message_id,
            score_delta,
            count_delta,
        )
    else:
        logging.info(
            "Reaction feedback received for untracked message chat=%s message=%s",
            reaction.chat.id,
            reaction.message_id,
        )


async def chat_native_status_command(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """Показывает, насколько 13-й пакет уже освоил конкретный чат."""

    del context
    if not update.message or not update.effective_chat:
        return
    if update.effective_chat.type == ChatType.PRIVATE:
        await update.message.reply_text("chat_native существует только в группах.")
        return

    status = await asyncio.to_thread(
        get_chat_native_learning_status_sync,
        update.effective_chat.id,
    )
    terms = status.get("terms") or []
    if terms:
        learned = ", ".join(terms[:12])
        state = "готов и используется как отдельный 13-й voice pack"
    else:
        learned = "пока недостаточно устойчивых словечек"
        state = "ещё набирает выборку"

    await update.message.reply_text(
        "chat_native этого чата:\n"
        f"Статус: {state}\n"
        f"Участников в обучающей выборке: {status.get('observed_users', 0)}\n"
        f"Кандидатов-термов: {status.get('candidate_terms', 0)}\n"
        f"Отслеживаемых ответов Яйцеслава: {status.get('tracked_messages', 0)}\n"
        f"Ответов с реакционной обратной связью: {status.get('reacted_messages', 0)}\n"
        f"Освоено: {learned}"
    )


async def gemini_version_command(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """Проверяет реальную модель Gemini через Google API."""

    del context

    message = update.effective_message
    user = update.effective_user

    if not message or not user:
        return

    # Как и /stats: посторонним ничего не раскрываем.
    if (
        BOT_OWNER_ID == 0
        or user.id != BOT_OWNER_ID
    ):
        return

    try:
        model_info = await asyncio.wait_for(
            gemini_client.aio.models.get(
                model=MODEL_NAME,
            ),
            timeout=20,
        )

        api_name = getattr(model_info, "name", None) or "—"
        display_name = getattr(model_info, "display_name", None) or "—"
        version = getattr(model_info, "version", None) or "—"
        thinking = getattr(model_info, "thinking", None)
        input_limit = getattr(model_info, "input_token_limit", None)
        output_limit = getattr(model_info, "output_token_limit", None)

        await message.reply_text(
            "Gemini API: OK\n"
            f"Configured model: {MODEL_NAME}\n"
            f"API model: {api_name}\n"
            f"Display name: {display_name}\n"
            f"Version: {version}\n"
            f"Thinking supported: {thinking}\n"
            f"Input token limit: {input_limit}\n"
            f"Output token limit: {output_limit}"
        )

    except Exception as error:
        logging.exception(
            "Ошибка проверки версии Gemini: %s",
            error,
        )

        await message.reply_text(
            "Gemini API: ERROR\n"
            f"Configured model: {MODEL_NAME}\n"
            f"{type(error).__name__}: {error}"
        )


# ============================================================
# ЗАПУСК БОТА
# ============================================================
async def error_handler(
    update: object,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """Записывает необработанные ошибки в журнал."""

    error = context.error

    if error is None:
        return

    logging.error(
        "Необработанная ошибка Telegram: %s",
        error,
        exc_info=(
            type(error),
            error,
            error.__traceback__,
        ),
    )
def main() -> None:
    """Запускает Telegram-бота."""

    application = (
        Application.builder()
        .token(TELEGRAM_TOKEN)
        .connect_timeout(30)
        .read_timeout(30)
        .write_timeout(60)
        .get_updates_connect_timeout(30)
        .get_updates_read_timeout(60)
        .get_updates_write_timeout(30)
        .concurrent_updates(8)
        .post_init(on_application_startup)
        .post_shutdown(on_application_shutdown)
        .build()
    )

    application.add_handler(
        CommandHandler(
            "start",
            start,
        )
    )

    application.add_handler(
        CommandHandler(
            "help",
            help_command,
        )
    )
    application.add_handler(
        CommandHandler(
            "stats",
            stats_command,
        )
    )
    application.add_handler(
        CommandHandler(
            "geminiversion",
            gemini_version_command,
        )
    )
    application.add_handler(
        CommandHandler(
            "settings",
            settings_command,
        )
    )
    application.add_handler(
        CommandHandler(
            "forget",
            forget_command,
        )
    )
    application.add_handler(
        CommandHandler(
            "voice_on",
            voice_on,
        )
    )

    application.add_handler(
        CommandHandler(
            "voice_off",
            voice_off,
        )
    )
    application.add_handler(
        CommandHandler(
            "search",
           search_command,
       )
    )
    application.add_handler(
        CommandHandler(
            "roast",
           roast_command,
       )
    )

    application.add_handler(
       CommandHandler(
           "wisdom",
          wisdom_command,
       )
    )

    application.add_handler(
        CommandHandler(
            "mood",
           mood_command,
       )
    )

    application.add_handler(
        CommandHandler(
            "hard_on",
           hard_on_command,
       )
    )

    application.add_handler(
        CommandHandler(
            "hard_off",
           hard_off_command,
       )
    )

    application.add_handler(
        CommandHandler(
            "hard_status",
           hard_status_command,
       )
    )
    application.add_handler(
        CommandHandler(
            "hard_level",
            hard_level_command,
        )
    )
    application.add_handler(
        CommandHandler(
            "hard_stats",
            hard_stats_command,
        )
    )
    application.add_handler(
        CommandHandler(
            "profile",
            profile_command,
        )
    )
    application.add_handler(
        CommandHandler(
            "nickname",
            nickname_command,
        )
    )
    application.add_handler(
        CommandHandler(
            "nickname_off",
            nickname_off_command,
        )
    )
    application.add_handler(
        CommandHandler(
            "whoami",
            whoami_command,
        )
    )
    application.add_handler(
        CommandHandler(
            "remember_me",
            remember_me_command,
        )
    )
    application.add_handler(
        CommandHandler(
            "forget_me",
            forget_me_command,
        )
    )
    application.add_handler(
        CommandHandler(
            "people",
            people_command,
        )
    )
    application.add_handler(
        CommandHandler(
            "set_archetype",
            set_archetype_command,
        )
    )
    application.add_handler(
        CommandHandler(
            "title",
            title_command,
        )
    )
    application.add_handler(
        CommandHandler(
            "title_status",
            title_status_command,
        )
    )
    application.add_handler(
        CommandHandler(
            "chat_native_status",
            chat_native_status_command,
        )
    )
    application.add_handler(
        CommandHandler(
            "prophecy",
            prophecy_command,
        )
    )
    application.add_handler(
        CommandHandler(
            "judge",
            judge_command,
        )
    )
    application.add_handler(
        CommandHandler(
            "argument",
            argument_command,
        )
    )
    application.add_handler(
        CommandHandler(
            "debate",
            debate_command,
        )
    )
    application.add_handler(
        CommandHandler(
            "explain_like_skoof",
            explain_like_skoof_command,
        )
    )
    application.add_handler(
        CommandHandler(
            "explain_like_rus",
            explain_like_rus_command,
        )
    )
    application.add_handler(
        CommandHandler(
            "meme",
            meme_command,
        )
    )
    application.add_handler(
        CommandHandler(
            "recap",
            recap_command,
        )
    )
    application.add_handler(
        CommandHandler(
            "fact_or_bayan",
            fact_or_bayan_command,
        )
    )
    application.add_handler(
        CommandHandler(
            "anti_advice",
            anti_advice_command,
        )
    )
    application.add_handler(
        CommandHandler(
            "translate_yayceslav",
            translate_yayceslav_command,
        )
    )
    application.add_handler(
        CommandHandler(
            "duel",
            duel_command,
        )
    )
    application.add_handler(
        CommandHandler(
            "story",
            story_command,
        )
    )
    application.add_handler(
        CommandHandler(
            "week",
            week_command,
        )
    )
    application.add_handler(
        CommandHandler(
            "week_me",
            week_me_command,
        )
    )
    application.add_handler(
        CommandHandler(
            "leaderboard",
            leaderboard_command,
        )
    )
    application.add_handler(
        CommandHandler(
            "awards",
            awards_command,
        )
    )
    application.add_handler(
        CommandHandler(
            "week_auto_on",
            week_auto_on_command,
        )
    )
    application.add_handler(
        CommandHandler(
            "week_auto_off",
            week_auto_off_command,
        )
    )
    application.add_handler(
        CommandHandler(
            "week_auto_status",
            week_auto_status_command,
        )
    )
    application.add_handler(
        CommandHandler(
            "week_time",
            week_time_command,
        )
    )
    application.add_handler(
        CallbackQueryHandler(
            duel_accept_callback,
            pattern=r"^duel_accept_",
        )
    )
    application.add_handler(
        CallbackQueryHandler(
            answer_button_callback,
            pattern=r"^answer_",
        )
    )
    application.add_handler(
        CallbackQueryHandler(
            settings_button_callback,
            pattern=r"^settings_",
        )
    )
    application.add_handler(
        MessageReactionHandler(
            message_reaction_feedback_handler,
        )
    )
    application.add_handler(
        MessageHandler(
            filters.PHOTO,
            answer_photo,
        )
    )

    application.add_handler(
        MessageHandler(
            filters.Document.ALL,
            answer_document,
        )
    )

    application.add_handler(
        MessageHandler(
            filters.VOICE | filters.AUDIO | filters.VIDEO_NOTE,
            answer_voice_or_audio,
        )
    )

    application.add_handler(
        MessageHandler(
            filters.TEXT & ~filters.COMMAND,
            answer_text_message,
        )
    )
    application.add_handler(
        MessageHandler(
            filters.TEXT & ~filters.COMMAND,
            hard_mode_listener,
        ),
        group=1,
    )
    application.add_error_handler(
        error_handler
    )
    print(
        "Яйцеслав запущен.\n"
        "Текст, фото, документы, входящие голосовые, "
        "видео-кружочки и голосовые ответы подключены.\n"
        "Для остановки нажмите Ctrl+C."
    )

    application.run_polling(
        drop_pending_updates=True,
        allowed_updates=[
            UpdateType.MESSAGE,
            UpdateType.EDITED_MESSAGE,
            UpdateType.CALLBACK_QUERY,
            UpdateType.MESSAGE_REACTION,
        ],
    )


if __name__ == "__main__":
    main()
